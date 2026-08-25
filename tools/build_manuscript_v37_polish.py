#!/usr/bin/env python
"""Final narrative and table-width polish for the focused v36 manuscript."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import set_font
from build_manuscript_v33_slim import format_three_line

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v36_focused.docx"
DEST = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v37_submission.docx"
SUP_SOURCE = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v36.docx"
SUP_DEST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v37.docx"


def replace(p, value, indent=.35):
    p.clear(); p.style = "Normal"
    p.paragraph_format.first_line_indent = Inches(indent) if indent else None
    p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(value), size=10.5)


def find(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def remove_last_columns(table, count=2):
    for row in table.rows:
        for cell in row.cells[-count:]:
            row._tr.remove(cell._tc)
    format_three_line(table)


def main():
    doc = Document(SOURCE)
    sup = Document(SUP_SOURCE)

    # Abstract: positioning, one representative controlled result, and the conclusion.
    replace(find(doc, "We present AstroCFR,"),
        "We present AstroCFR, a modular framework that exposes candidate-recovery, measurement-precision, "
        "and computational-cost operating points in crowded stellar fields. In the registered CSST-like "
        "simulation configuration (four chips; approximately 4,000 reference sources), its conservative "
        "target-adapted RandomForest branch reaches 94.6% mean recall; the 100% precision is explicitly "
        "limited to supplied simulation-catalogue assembly. In a controlled common-protocol HST/ACS NGC 6752 "
        "test, AstroCFR ePSF-deblend recovers 87.6% (95% CI 84.0–90.4%) of dense V≤20 references, compared "
        "with 57.0% for DAO and 28.6% for SEP. Photutils provides lower positional and photometric RMS, "
        "whereas the recovery-oriented AstroCFR branch requires more computation. A simulation-developed "
        "front end can be recalibrated with a small spatially disjoint labelled target region, but zero-shot "
        "transfer fails. AstroCFR therefore provides a reproducible Pareto frontier for recovery, measurement "
        "precision, and cost rather than universal photometric dominance.")

    # The controlled scientific-comparison table must fit a two-column article.  Cost belongs in Table 9.
    controlled = next(t for t in doc.tables if t.cell(0, 0).text == "Cluster" and t.cell(0, 1).text == "Method")
    remove_last_columns(controlled)
    ablation = next(t for t in doc.tables if t.cell(0, 0).text == "Cluster" and t.cell(0, 1).text == "Branch")
    remove_last_columns(ablation)
    replace(find(doc, "Photutils PSFPhotometry uses the same DAOStarFinder"),
        "Photutils PSFPhotometry uses the same DAOStarFinder proposal frontend as DAO. Its proposal-level "
        "detection and artificial-star recovery are therefore identical to DAO by design; its independent "
        "role in this comparison is PSF-based astrometric and photometric measurement. All recovery and "
        "completeness intervals use the Wilson 95% interval. Runtime and peak RSS are reported separately "
        "in the operating-point table, so that the controlled table remains a scientific-performance comparison.")

    # Be explicit about the scope of the sensitivity audit, rather than overstating an unrun sweep.
    p = find(doc, "The dominant dense-field failure mode is structural blending")
    replace(p,
        "The dominant dense-field failure mode is structural blending. Candidate groups are processed by "
        "paired-PSF fitting, joint multi-source fitting, or residual deblending according to group complexity. "
        "All spatial scales and structural settings—including grouping radius, polynomial order, and calibration-"
        "tile size—are registered hyperparameters, not physical constants. Threshold and association-radius "
        "sensitivity are audited in Supplementary Tables S5–S6. Grouping and distortion-model settings must "
        "be revalidated on a new instrument's validation partition; the present work does not claim their "
        "instrument-independent optimality.")

    # One confident conclusion: the contribution is the measured Pareto frontier.
    replace(find(doc, "AstroCFR provides a reproducible comparison framework"),
        "The framework provides a reproducible Pareto frontier for crowded-field candidate recovery, "
        "single-image measurement, and computational cost. Its controlled HST evidence uses common image "
        "crops, association rules, spatial partitions, and injected scenes, while the CSST-like results remain "
        "a registered challenge configuration.")
    replace(find(doc, "Three bounded conclusions follow."),
        "The frontier has three practically useful operating points. In dense NGC 6752 and NGC 1851 subsets, "
        "the ePSF-based AstroCFR branches recover more reference and injected sources than the evaluated DAO, "
        "SEP, and RF branches. Photutils provides the lowest reported NGC 6752 positional and photometric RMS, "
        "while DOLPHOT/ALLFRAME-class multi-exposure measurement remains outside the present single-image "
        "comparison. The recovery branches incur a substantial CPU cost and should be used where their "
        "high-crowding recovery benefit justifies that cost. The spatial-ePSF/two-pass variant improves "
        "AstroCFR photometric RMS to 0.037 mag but does not overturn the Photutils positional result.")
    replace(find(doc, "Code availability:"),
        "Code availability: The manuscript-matched source is prepared for https://github.com/zxl1999/AstroCFR. "
        "The release contains reusable modules in src/wpdc, controlled HST and CSST experiment scripts, "
        "machine-readable summaries, environment locks, data provenance, and manuscript builders. The prepared "
        "version is v1.5.0; the public commit hash, tag, and release archive must be recorded after upload. "
        "No archival DOI is claimed before a public release exists.")

    for table in doc.tables:
        format_three_line(table)
    doc.save(DEST)
    sup.save(SUP_DEST)
    print(DEST)
    print(SUP_DEST)


if __name__ == "__main__":
    main()
