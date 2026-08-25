#!/usr/bin/env python
"""Remove residual duplicated prose and stale numbering from v37."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import set_font

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v37_submission.docx"
DEST = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v38_final.docx"
SUP_SOURCE = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v37.docx"
SUP_DEST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v38.docx"


def find(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def replace(p, value, indent=.35):
    p.clear(); p.style = "Normal"
    p.paragraph_format.first_line_indent = Inches(indent) if indent else None
    p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(value), size=10.5)


def remove(p):
    p._element.getparent().remove(p._element)


def insert_after(p, value):
    q = p._parent.add_paragraph(style="Normal")
    q._p.getparent().remove(q._p); p._p.addnext(q._p)
    replace(q, value)
    return q


def main():
    doc = Document(SOURCE)
    sup = Document(SUP_SOURCE)

    # Keep one background-model paragraph, including the registered threshold details.
    p = find(doc, "AstroCFR first estimates a two-dimensional background and RMS map using block-wise robust statistics")
    replace(p,
        "AstroCFR first estimates a two-dimensional background and RMS map using block-wise robust statistics "
        "with median filtering and interpolation. The background model tracks diffuse galaxy light and detector-"
        "scale variation, while the residual image exposes candidate point sources. The registered CSST-like "
        "configuration uses a nominal 2.5-sigma proposal threshold; chip 18 uses 2.39 sigma after validation "
        "because of its higher background RMS. These are dataset-specific operating parameters, not universal "
        "constants. The public-HST benchmark uses a common 3-sigma front end; threshold and matching-radius "
        "sensitivity are reported in Supplementary Tables S5–S6.")
    duplicate = find(doc, "AstroCFR first estimates a two-dimensional background and RMS map using block-wise robust statistics")
    if duplicate is not p:
        remove(duplicate)
    else:
        # The second duplicated paragraph has the CSST-like challenge wording.
        q = next(x for x in doc.paragraphs if x.text.startswith("AstroCFR first estimates a two-dimensional background") and x is not p)
        remove(q)

    # The matching rule is defined once in Sec. 4.1; later prose just uses it.
    rule = find(doc, "Detected sources are matched to reference stars using a greedy nearest-neighbor")
    replace(rule,
        "Detected sources are matched to reference stars using the greedy one-to-one two-pixel rule. In simulations, "
        "recall and precision use the complete supplied challenge catalogue. In HST and external-survey tests, the "
        "association fraction is reported only as a catalogue-match lower bound because deeper real sources can be "
        "unmatched; astrometric and photometric RMS use matched held-out residuals.")
    duplicate_rule = find(doc, "Detected sources are associated with reference stars through a greedy one-to-one two-pixel rule")
    remove(duplicate_rule)

    # Keep the calibrated SExtractor description in Sec. 4.3 and give 4.4 its actual method content.
    for prefix in (
        "A calibrated SExtractor baseline is run on the same simulated exposures after correcting",
        "The SExtractor comparison is diagnostic rather than a universal ranking",
    ):
        remove(find(doc, prefix))
    routing_heading = find(doc, "4.4 Density-adaptive branch-routing protocol")
    insert_after(routing_heading,
        "The density-adaptive policy is evaluated as a fixed, pre-registered routing rule over local density strata: "
        "a fast branch is used in lower-density regions and a recovery-oriented ePSF branch in higher-density regions. "
        "It is an explicitly measured operating point, not an automatic image-only gate; the latter remains a negative "
        "control in the Supplementary Materials.")

    # Correct v35 residue in the runtime protocol paragraph.
    runtime = find(doc, "To avoid treating one point estimate as a deployment guarantee")
    replace(runtime, runtime.text.replace("Tables 15–18 and Table 9", "Tables 4–7 and Table 9").replace(
        "The CNN classifier is profiled separately in Table 22 on the available GPU.",
        "The optional CNN CPU/GPU profile is archived in Supplementary Section S7."
    ))

    # State the three deployment choices directly in the conclusion.
    conclusion = find(doc, "The frontier has three practically useful operating points")
    replace(conclusion,
        "The Pareto frontier exposes three practically useful operating points. DAOStarFinder and AstroCFR-RF are "
        "suitable for low-latency screening. Photutils is the preferred choice when positional or photometric accuracy "
        "is paramount. The ePSF-based AstroCFR branches should be used only in high-density regions where the recovery "
        "gain justifies their substantial CPU cost. In dense NGC 6752 and NGC 1851 subsets, these recovery branches "
        "recover more reference and injected sources than the evaluated DAO, SEP, and RF branches; the spatial-ePSF/"
        "two-pass variant improves AstroCFR photometric RMS to 0.037 mag but does not overturn the Photutils positional result.")

    doc.save(DEST)
    sup.save(SUP_DEST)
    print(DEST)
    print(SUP_DEST)


if __name__ == "__main__":
    main()
