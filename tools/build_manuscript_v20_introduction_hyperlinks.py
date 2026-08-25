#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Replace the introduction and hyperlink author-year citations to References."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[1]
SOURCE=ROOT/'supplementary'/'WPDC_Multimedia_Systems_SCI_manuscript_v19_reference_order.docx'
DEST=ROOT/'supplementary'/'WPDC_Multimedia_Systems_SCI_manuscript_v20_introduction_hyperlinks.docx'

INTRO=[
"Scientific imaging pipelines are increasingly analogous to multimedia systems: they ingest high-resolution visual streams, fuse pixel-level and feature-level representations, perform detection and classification under domain constraints, and output structured data products for downstream retrieval and analysis. Dense stellar-field photometry offers a representative high-stress case. In regions such as the Galactic bulge, globular clusters, and nearby resolved galaxies, overlapping point-spread functions (PSFs), variable sky backgrounds, and bright-star artifacts degrade both the completeness and purity of automated source extraction.",
"The China Space Station Telescope (CSST) will provide high-resolution wide-field imaging, and simulated CSST dense-field data already reveal the main processing difficulties expected in crowded fields. The M31 field is particularly challenging because reference stars can be separated by only a few detector pixels. A practical system must therefore do more than detect local maxima: it must combine background modeling, multi-branch candidate generation, artifact rejection, deblending, astrometric correction, and photometric calibration as a coupled system.",
"Our positioning follows established measurement and deployment practice: empirical ACS ePSFs and crowded-field catalogues (Anderson & King, 2006; Anderson et al., 2008; Bellini et al., 2011), survey-scale reference systems (Chambers et al., 2016; Dey et al., 2019; Gaia Collaboration et al., 2016, 2023), and reproducible scientific-computing components (Astropy Collaboration, 2022; Harris et al., 2020; Hunter, 2007; Virtanen et al., 2020; Barbary, 2016; Bradley et al., 2024; Pedregosa et al., 2011). Deep-learning foundations and residual networks provide the image-model context (Goodfellow et al., 2016; He et al., 2016). Recent context includes CSST photometric preparation (Shi et al., 2024), deep-feature point-source detection (Long et al., 2025), Rubin crowded-field photometry (Wainer et al., 2025), PSF-fitting uncertainty analysis (Espinosa et al., 2025), image-based PSF recovery (Centofanti et al., 2026), and recent HST/Euclid cluster studies (Libralato et al., 2024; Salaris et al., 2024). Transfer-learning claims are framed using domain-shift theory and adaptation methods (Ben-David et al., 2010; Pan & Yang, 2010; Ganin et al., 2016; Long et al., 2015), while the classifier baselines are standard Random Forest and gradient-boosting references (Breiman, 2001; Chen & Guestrin, 2016).",
"Classical systems such as DAOPHOT (Stetson, 1987) and SExtractor (Bertin & Arnouts, 1996) established the foundations of stellar detection and photometry. Later crowded-field tools improved PSF modeling and joint fitting (Dolphin, 2000; Schlafly et al., 2018), while modern deblenders added new ways to separate overlapping objects (Melchior et al., 2018). CSST-focused studies have also investigated dense-field astrometry, photometry, and star-galaxy classification (Wang et al., 2024; Zhang et al., 2024a; Zhang et al., 2024b). However, many studies evaluate a single module outside a complete processing chain. For a production multimedia system, the key question is not only which classifier has the best isolated metric, but how classifier outputs interact with downstream calibration and catalog assembly.",
"We present WPDC as a deployable system for candidate generation, domain adaptation, and crowded-field recovery in large-scale astronomical survey images. Its practical contribution is not one universally optimal catalogue: it exposes two reproducible operating points. The target-adapted RandomForest (WPDC-RF) provides a conservative catalogue branch, whereas the ePSF plus residual-deblending branch pursues higher crowded-field candidate recovery at a measured computational cost. The contribution is evaluated with identical HST images, association rule, spatial test partition, and artificial-star scenes across classical and WPDC baselines.",
"The manuscript is organized as follows. Section 2 describes the CSST-like dataset and WPDC architecture. Section 3 details the detection, deblending, classification, astrometric, and photometric modules. Section 4 presents the experimental protocol, including classifier benchmarking and SExtractor comparison. Section 5 reports results. Section 6 discusses system-level implications and limitations, and Section 7 concludes."
]

def bookmark(p,name,ident):
    s=OxmlElement('w:bookmarkStart');s.set(qn('w:id'),str(ident));s.set(qn('w:name'),name)
    e=OxmlElement('w:bookmarkEnd');e.set(qn('w:id'),str(ident));p._p.insert(0,s);p._p.append(e)

def font_run(r):
    r.font.name='Times New Roman';r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体');r.font.size=Pt(10.5);r.font.color.rgb=RGBColor(0,0,0)

def link(p,label,anchor):
    h=OxmlElement('w:hyperlink');h.set(qn('w:anchor'),anchor);h.set(qn('w:history'),'1')
    r=OxmlElement('w:r');rp=OxmlElement('w:rPr');col=OxmlElement('w:color');col.set(qn('w:val'),'0563C1');u=OxmlElement('w:u');u.set(qn('w:val'),'single');rp.append(col);rp.append(u);r.append(rp);t=OxmlElement('w:t');t.text=label;r.append(t);h.append(r);p._p.append(h)

def main():
    doc=Document(SOURCE)
    ridx=next(i for i,p in enumerate(doc.paragraphs) if p.text.strip()=='References')
    refs=[p for p in doc.paragraphs[ridx+1:] if p.text.strip()]
    # Build reference anchors by unambiguous bibliography prefixes.
    prefixes=['Anderson, J., & King','Anderson, J., et al.','Bellini,','Chambers,','Dey,','Gaia Collaboration, et al. (2016)','Gaia Collaboration, et al. (2023)','Astropy Collaboration,','Harris,','Hunter,','Virtanen,','Barbary,','Bradley,','Pedregosa,','Goodfellow,','He,','Shi,','Long, M., Xin','Wainer,','Espinosa,','Centofanti,','Libralato,','Salaris,','Ben-David,','Pan,','Ganin,','Long, M., Cao','Breiman,','Chen,','Stetson,','Bertin,','Dolphin,','Schlafly,','Melchior,','Wang,','Zhang, Y.,','Zhang, S.']
    anchor={}
    for n,p in enumerate(refs,1):
        for pref in prefixes:
            if p.text.startswith(pref):anchor[pref]=f'ref_{n}'
        bookmark(p,f'ref_{n}',500+n)
    cite={
      'Anderson & King, 2006':'Anderson, J., & King','Anderson et al., 2008':'Anderson, J., et al.','Bellini et al., 2011':'Bellini,','Chambers et al., 2016':'Chambers,','Dey et al., 2019':'Dey,','Gaia Collaboration et al., 2016, 2023':'SPECIAL_GAIA','Astropy Collaboration, 2022':'Astropy Collaboration,','Harris et al., 2020':'Harris,','Hunter, 2007':'Hunter,','Virtanen et al., 2020':'Virtanen,','Barbary, 2016':'Barbary,','Bradley et al., 2024':'Bradley,','Pedregosa et al., 2011':'Pedregosa,','Goodfellow et al., 2016':'Goodfellow,','He et al., 2016':'He,','Shi et al., 2024':'Shi,','Long et al., 2025':'Long, M., Xin','Wainer et al., 2025':'Wainer,','Espinosa et al., 2025':'Espinosa,','Centofanti et al., 2026':'Centofanti,','Libralato et al., 2024':'Libralato,','Salaris et al., 2024':'Salaris,','Ben-David et al., 2010':'Ben-David,','Pan & Yang, 2010':'Pan,','Ganin et al., 2016':'Ganin,','Long et al., 2015':'Long, M., Cao','Breiman, 2001':'Breiman,','Chen & Guestrin, 2016':'Chen,','Stetson, 1987':'Stetson,','Bertin & Arnouts, 1996':'Bertin,','Dolphin, 2000':'Dolphin,','Schlafly et al., 2018':'Schlafly,','Melchior et al., 2018':'Melchior,','Wang et al., 2024':'Wang,','Zhang et al., 2024a':'Zhang, Y.,','Zhang et al., 2024b':'Zhang, S.'}
    intro_index=next(i for i,p in enumerate(doc.paragraphs) if p.text.strip()=='1 Introduction')
    targets=[]
    for p in doc.paragraphs[intro_index+1:]:
        if p.text.startswith('The corresponding visualization is shown in Fig. 1.'):break
        if p.text.strip():targets.append(p)
    if len(targets)!=6:raise RuntimeError(f'Expected 6 introduction paragraphs, found {len(targets)}')
    rx=re.compile('|'.join(re.escape(x) for x in sorted(cite,key=len,reverse=True)))
    for i,(p,text) in enumerate(zip(targets,INTRO)):
        p.clear();p.style='Normal';p.paragraph_format.first_line_indent=None if i==0 else Inches(.35);p.paragraph_format.line_spacing=1.08
        at=0
        for m in rx.finditer(text):
            if m.start()>at:r=p.add_run(text[at:m.start()]);font_run(r)
            label=m.group(0)
            if cite[label]=='SPECIAL_GAIA':
                link(p,'Gaia Collaboration et al., 2016',anchor['Gaia Collaboration, et al. (2016)'])
                r=p.add_run(', ');font_run(r)
                link(p,'2023',anchor['Gaia Collaboration, et al. (2023)'])
            else:
                link(p,label,anchor[cite[label]])
            at=m.end()
        if at<len(text):r=p.add_run(text[at:]);font_run(r)
    doc.save(DEST);print(DEST,'intro citations',sum(len(rx.findall(t)) for t in INTRO))
if __name__=='__main__':main()
