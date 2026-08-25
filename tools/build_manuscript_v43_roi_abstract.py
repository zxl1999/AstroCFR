#!/usr/bin/env python
"""Build v43 with a quantitative closed-book, adaptation-ROI, and cost abstract."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import set_font


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v42_closed_book_scope.docx"
DEST = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v43_roi_abstract.docx"
SUP_SOURCE = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v42.docx"
SUP_DEST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v43.docx"


def find(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def rewrite(paragraph, text, indent=.35):
    paragraph.clear()
    paragraph.style = "Normal"
    paragraph.paragraph_format.first_line_indent = Inches(indent) if indent else None
    paragraph.paragraph_format.line_spacing = 1.08
    set_font(paragraph.add_run(text), size=10.5)
    return paragraph


def insert_after(paragraph, text):
    new = paragraph._parent.add_paragraph(style="Normal")
    new._p.getparent().remove(new._p)
    paragraph._p.addnext(new._p)
    return rewrite(new, text)


def main():
    doc = Document(SOURCE)
    sup = Document(SUP_SOURCE)

    rewrite(find(doc, "We present AstroCFR,"),
        "We present AstroCFR, a modular framework that exposes candidate-recovery, measurement-precision, and "
        "computational-cost operating points in crowded stellar fields. In four CSST-like simulation chips "
        "containing approximately 4,000 supplied references, the registered branch reaches 94.6% mean reference "
        "recovery; reference-aware catalogue assembly is not treated as blind precision. A closed-book image-only "
        "injection pilot writes candidate catalogues before truth is opened and recovers 18.3% in high-density, "
        "peak-SNR-10 scenes and 89.4% in low-density, peak-SNR-30 scenes; because the supplied top-1000 catalogue is "
        "not certified exhaustive, this pilot does not establish blind purity. On the common-protocol HST/ACS "
        "NGC 6752 test, AstroCFR ePSF-deblend recovers 87.6% (95% CI 84.0-90.4%) of dense V<=20 references versus "
        "57.0% for DAO and 28.6% for SEP, but requires 24.24 rather than 0.11 s/MPix on the same CPU workstation "
        "and does not match Photutils measurement RMS. Adaptation cost is bounded explicitly: one spatially isolated "
        "200 x 200-pixel tile (2.8% of the crop; approximately 101 positive and 15 negative candidates) raises "
        "held-out NGC 6752 RF recall from 5.7% to 64.8%, whereas zero-shot transfer fails. AstroCFR therefore "
        "provides a reproducible recovery-precision-cost frontier with measurable calibration overhead rather than "
        "universal photometric or cross-instrument dominance.", indent=0)

    roi = find(doc, "Thus the defensible deployment statement is not that a fixed number of real images")
    insert_after(roi,
        "Calibration effort is reported in labelled area and candidate counts, not person-minutes. The labels in "
        "this experiment were derived from the official catalogue rather than collected in a timed manual-annotation "
        "study; assigning a five-minute human cost would therefore be unsupported. For NGC 6752, the measured return "
        "is a 59.1-percentage-point held-out RF-recall gain from one 2.8%-area tile containing approximately 116 "
        "labelled candidates. The separate 87.6% ePSF-deblend recovery result is image-driven and must not be "
        "attributed to that RF calibration tile.")

    doc.save(DEST)
    sup.save(SUP_DEST)
    print(DEST)
    print(SUP_DEST)


if __name__ == "__main__":
    main()
