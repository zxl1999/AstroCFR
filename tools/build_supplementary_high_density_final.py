#!/usr/bin/env python
"""Synchronize supplementary material with the 15-field high-density manuscript."""
from __future__ import annotations

from pathlib import Path
import csv

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

from build_manuscript_v30_submission_fixes import set_cell, set_font
from build_manuscript_high_density_final import (
    BASE, DST as MAIN_DST, FIELDS, FIELD_LABEL, LIT, load_rows,
    apply_three_line_table, rebuild_cross_field_table, runtime_density_summary,
)

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v45_submission_ready.docx"
DST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v45_high_density_final.docx"
CSST_UNIFIED = ROOT / "results" / "csst_unified_five_methods" / "csst_unified_five_methods.csv"


def pfind(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix) or prefix in p.text)


def rewrite(paragraph, text):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_font(run, size=10.5, bold=paragraph.style.name.startswith("Heading"))


def insert_after(paragraph, style="Normal"):
    new = paragraph._parent.add_paragraph(style=style)
    new._p.getparent().remove(new._p)
    paragraph._p.addnext(new._p)
    return new


def equation_after(doc, anchor, linear_equation, number):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table._tbl.getparent().remove(table._tbl)
    anchor.addnext(table._tbl)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}"); node.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "nil"); borders.append(node)
    table._tbl.tblPr.append(borders)
    for cell, width in zip(table.rows[0].cells, (1200, 7200, 1200)):
        cell.width = width
    eq = table.rows[0].cells[1].paragraphs[0]
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.first_line_indent = None
    omath_para = OxmlElement("m:oMathPara")
    omath = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    text.text = linear_equation
    run.append(text); omath.append(run); omath_para.append(omath)
    eq._p.append(omath_para)
    number_p = table.rows[0].cells[2].paragraphs[0]
    number_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_p.paragraph_format.first_line_indent = None
    set_font(number_p.add_run(f"({number})"), size=10.5)
    return table


def rebuild_inventory(table):
    headers = ["Evidence tier", "Unit", "Scene", "Status in this manuscript"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, size=7.4)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    records = [
        ("CSST-like full-frame simulation", "Chip 12", "supplied-catalogue dense chip", "registered AstroCFR ePSF-deblend and SExtractor audit"),
        ("CSST-like full-frame simulation", "Chip 13", "supplied-catalogue dense chip", "registered AstroCFR ePSF-deblend and SExtractor audit"),
        ("CSST-like full-frame simulation", "Chip 17", "supplied-catalogue dense chip", "registered AstroCFR ePSF-deblend and SExtractor audit"),
        ("CSST-like full-frame simulation", "Chip 18", "supplied-catalogue dense chip", "registered AstroCFR ePSF-deblend and SExtractor audit"),
    ]
    records += [("HST/ACS single-F606W controlled", FIELD_LABEL[field], "ACSGGCT globular cluster", "five common image-only branches plus two literature-mapped PSF controls") for field in FIELDS]
    for values in records:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell(cell, value, size=7.0)
    for ridx, row in enumerate(table.rows):
        trpr = row._tr.get_or_add_trPr()
        trpr.append(OxmlElement("w:cantSplit"))
        if ridx == 0:
            trpr.append(OxmlElement("w:tblHeader"))


def add_csst_unified_audit(doc, anchor):
    """Append the controlled-crop CSST five-branch audit after S26."""
    if not CSST_UNIFIED.exists():
        return
    rows = list(csv.DictReader(CSST_UNIFIED.open(encoding="utf-8")))
    cap = insert_after(anchor, style="Caption")
    cap.paragraph_format.keep_with_next = True
    lead = cap.add_run("Table S27.")
    set_font(lead, size=10.0, bold=True)
    rest = cap.add_run(" CSST unified five-branch controlled-crop audit. The four 1200 x 1200 image-only crops (chips 12, 13, 17, and 18) use a common 200-pixel spatial partition, a two-pixel one-to-one association rule, and reference catalogues only for final scoring. This table is a feasibility audit and does not replace the registered full-frame CSST benchmark or the 11-field HST primary comparison.")
    set_font(rest, size=10.0)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table._tbl.getparent().remove(table._tbl)
    cap._p.addnext(table._tbl)
    headers = ["Chip", "Crop origin (px)", "Method", "Test refs", "Recovered", "Recovery", "Position RMS (mas)", "Magnitude RMS (mag)"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, size=7.0)
    for r in rows:
        vals = [r["chip"], f"({r['crop_x0']}, {r['crop_y0']})", r["method"], r["test_references"], r["test_recovered"], f"{100*float(r['test_recovery']):.1f}%", f"{float(r['test_position_rms_mas']):.2f}", f"{float(r['test_magnitude_rms_mag']):.4f}"]
        cells = table.add_row().cells
        for cell, value in zip(cells, vals):
            set_cell(cell, value, size=6.7)
    for ridx, row in enumerate(table.rows):
        trpr = row._tr.get_or_add_trPr()
        trpr.append(OxmlElement("w:cantSplit"))
        if ridx == 0:
            trpr.append(OxmlElement("w:tblHeader"))
    note = insert_after(cap, style="Normal")
    note.add_run("Interpretation. Across the four crops, the median test recovery is 63.3% for DAOStarFinder, 90.7% for SEP/SExtractor-style extraction, 63.3% for Photutils, 70.5% for AstroCFR ePSF-deblend, and 70.5% for AstroCFR spatial-ePSF joint fitting. The corresponding median position RMS values are 6.53, 25.86, 9.35, 5.65, and 6.41 mas. Chips 17 and 18 contain only 8 and 10 test references, respectively; therefore these values support implementation feasibility and protocol alignment, not a full-frame CSST SOTA claim.")
    set_font(note.runs[0], size=9.5)


def main():
    doc = Document(SRC)
    base, lit = load_rows()
    inventory_table = doc.tables[21]
    full_table = doc.tables[25]
    rewrite(pfind(doc, "This supplement contains the detailed classifier analyses"),
            "This supplement contains classifier, sensitivity, reproducibility, and high-density field-level material supporting the revised manuscript. The primary observational evidence is the completed 11-field ACSGGCT HST/ACS F606W suite, evaluated with common image crops, held-out spatial partitions, a two-pixel association rule, and a fixed high-density definition. CSST full-frame and controlled-crop results remain a separate simulation-to-reality reference tier and are not pooled with the HST evidence.")
    rewrite(pfind(doc, "For isolated classifier benchmarking"),
            "The archived nine-method classifier screen uses complete pre-classification candidate tables with features, cutout references, coordinates, and reference-match labels. It reports ten repeated stratified trials, per-chip normalization, 3:1 negative-to-positive undersampling, validation-selected thresholds, and held-out test scores. The original per-trial feature cache and complete parameter files for this legacy screen are not distributed in the upload package; it is therefore retained as a contextual diagnostic, not as part of the primary 11-field HST claim. The fully reproducible classifier protocol, including exact partitions, seed, labels, and hyperparameters, is the controlled ablation specified in Section S2 below.")
    rewrite(pfind(doc, "To test whether a newer attention-based classifier is necessary"),
            "To test whether a newer attention-based classifier is necessary, we performed a controlled simulation-domain ablation. Candidate generation, the association rule, and the four CSST-like chips were fixed. Candidates within 2 pixels of a simulated reference are positive; candidates farther than 8 pixels are negative; the 2-8 pixel annulus is excluded. The resulting cache has 11,615 candidates (3,615 positive and 8,000 negative) and uses a stratified 60/20/20 train/validation/test split with seed 20260806. The feature standardizer is fitted only on the 60% training split. The validation partition selects the highest threshold retaining at least 90% recall; all reported metrics are computed once on the untouched 20% test split. RandomForest uses 400 trees, maximum depth 15, minimum leaf size 2, sqrt feature sampling, class weights 1:6 for negative:positive, and the same seed. The original AstroCFR StarBogusNet and the patch Transformer use Adam with learning rate 1e-3, weight decay 1e-4, cosine annealing for 30 epochs, training batches of 64, and inference batches of 256. Both use the same 25 x 25 normalized cutout and 17 handcrafted features; the Transformer uses 25 non-overlapping 5 x 5 patches, 64-dimensional tokens, two four-head encoder layers, and a small feature-fusion head (74,193 parameters).")
    s8 = pfind(doc, "The six controlled branches were rerun")
    rewrite(s8,
            "The completed high-density audit reruns DAOStarFinder, SEP/SExtractor-style extraction, Photutils PSFPhotometry, AstroCFR ePSF-deblend, and AstroCFR spatial-ePSF joint fitting on the same 1200 x 1200 crops for 11 ACSGGCT fields. Global empirical-ePSF and three-Gaussian dPSF are diagnostic PSF-representation controls: they share the image, candidate front end, spatial hold-out, and neighbour-fitting harness, but their archived output uses the dense held-out reference denominator rather than the full baseline metric set. They are consequently reported with their own denominator and are not pooled into the Table 8 medians. Recovery uses a two-pixel one-to-one association and the baseline high-density subset contains V<=20 references with at least three neighbours within 10 pixels.")
    eq1 = equation_after(doc, s8._p, "D = {i : Vᵢ ≤ 20, n₁₀(i) ≥ 3};   R_dense = N_dense⁻¹∑ᵢ∈D 𝟙[minⱼ||x̂ⱼ−xᵢ||₂≤2 px]", "S1")
    equation_after(doc, eq1._tbl, "ΔR = R_dense(spatial-ePSF) − R_dense(Photutils)", "S2")
    rewrite(pfind(doc, "S11 Completed empirical evidence inventory"), "S11 Completed 15-field high-density evidence inventory and reproducibility package")
    rewrite(pfind(doc, "The completed manuscript evidence inventory has four CSST-like"),
            "The completed evidence inventory contains four CSST-like full-frame simulations and 11 public ACSGGCT HST/ACS F606W fields. The HST tier comprises NGC 2808, 5286, 6388, 6441, 0104, 0362, 6093, 6624, 6397, 6752, and 1851. The reproducibility package contains the unified baseline JSON, literature-control JSON, 15-field all-method matrix, pairwise spatial-ePSF versus Photutils table, and availability registry. HST and CSST values are never combined into one numerical average.")
    rewrite(pfind(doc, "Table S22. Completed empirical evidence inventory"),
            "Table S22. Completed 15-field evidence inventory. The HST tier is the primary common single-F606W, held-out spatial-partition evaluation. The CSST tier is a separate registered simulation-to-reality reference comprising the full-frame supplied-catalogue audit and the controlled-crop five-branch feasibility audit. Input-incompatible pipelines, protocol-excluded HST RF, and partial hybrid coverage are recorded in the machine-readable availability registry rather than assigned numerical scores.")
    rewrite(pfind(doc, "Table S23. Completed real HST expansion fields"),
            "Table S23. Archived non-ACSGGCT public-field inventory. These fields are retained for provenance but are not included in the primary 11-field high-density ACSGGCT median.")
    rewrite(pfind(doc, "Table S24. All admitted real-expansion comparisons"),
            "Table S24. Archived non-ACSGGCT comparison results. They use distinct filters and external catalogues and are not pooled with the primary ACSGGCT result.")
    rewrite(pfind(doc, "Table S25. GR8 real multi-exposure"),
            "Table S25. Archived GR8 multi-exposure benchmark. This separate image domain is retained for audit only and is not ranked against the single-image ACSGGCT methods.")
    rewrite(s8,
            "The completed high-density audit reruns DAOStarFinder, SEP/SExtractor-style extraction, Photutils PSFPhotometry, AstroCFR ePSF-deblend, and AstroCFR spatial-ePSF joint fitting on the same 1200 x 1200 crops for 11 ACSGGCT fields. A separate controlled-crop CSST audit runs the same five branches on chips 12, 13, 17, and 18; its 20 method-chip rows are reported in Table S27 and are not pooled with the registered full-frame CSST integration audit. Global empirical-ePSF and three-Gaussian dPSF are diagnostic HST PSF-representation controls: they share the image, candidate front end, spatial hold-out, and neighbour-fitting harness, but their archived output uses the dense held-out reference denominator rather than the full baseline metric set. They are consequently reported with their own denominator and are not pooled with the Table 8 medians. Recovery uses a two-pixel one-to-one association and the baseline high-density subset contains V<=20 references with at least three neighbours within 10 pixels.")
    disjoint = insert_after(s8, style="Normal")
    rewrite(disjoint,
            "Partition-isolated ePSF sensitivity audit. The full-image candidate list was held fixed, but only image-only DAO proposals in spatial cells 0/1 supplied global and quadrant ePSF stamps; partition 2 remained the held-out scoring region. Across the 11 ACSGGCT fields, the dense-recovery delta relative to the primary transductive branch had median -0.13 percentage points (range -2.33 to +0.45 pp), position-RMS delta median +0.09 mas (range -0.05 to +0.34 mas), and magnitude-RMS delta median +0.0033 mag (range -0.0014 to +0.0087 mag). The 11/11 spatial-ePSF-versus-Photutils dense-recovery ordering was unchanged. This is a PSF-construction robustness check, not an additional leaderboard; complete field-level deltas are in results/partition_isolated_spatial_epsf_sensitivity_all11/summary.json and comparison.csv.")
    threshold_audit = insert_after(disjoint, style="Normal")
    rewrite(threshold_audit,
            "Threshold and association robustness audit (NGC 6752). The image-only DAO proposal threshold was varied from 2.5 to 4.0 sigma without using partition 2 for selection. Quality-reference recovery changed from 77.13% to 74.99%, and the catalogue-match lower bound from 93.68% to 94.62%. With all produced catalogues held fixed, changing only the one-to-one association radius from 1 to 3 pixels retained the spatial-ePSF ordering over Photutils (93.19–93.45% versus 76.17–76.25% all-quality recovery). The audit verifies that the result is not a two-pixel matching artefact; it does not replace the fixed, preregistered two-pixel metric in the primary table. Machine-readable output: results/hst_parameter_sensitivity/parameter_sensitivity_audit.json.")
    runtime_audit = insert_after(threshold_audit, style="Normal")
    runtime = runtime_density_summary(base)
    rewrite(runtime_audit,
            ("Runtime and magnitude-stratified summary. On the common 1,200 x 1,200 CPU crops, field-median costs are "
             f"{runtime['dao']['runtime_s_per_mpix']:.2f}, {runtime['sep']['runtime_s_per_mpix']:.2f}, {runtime['photutils_psf']['runtime_s_per_mpix']:.2f}, "
             f"{runtime['wpdc_epsf_deblend']['runtime_s_per_mpix']:.2f}, and {runtime['wpdc_spatial_epsf_joint']['runtime_s_per_mpix']:.2f} s MPix⁻¹ "
             "for DAOStarFinder, SEP, Photutils, ePSF-deblend, and spatial-ePSF joint fitting, respectively. "
             f"The spatial branch is {runtime['wpdc_spatial_epsf_joint']['runtime_s_per_mpix']/runtime['photutils_psf']['runtime_s_per_mpix']:.1f}x the Photutils cost in this implementation. "
             f"Its field-median cumulative recovery is {100*runtime['wpdc_spatial_epsf_joint']['recall_v_le_18']:.1f}%, {100*runtime['wpdc_spatial_epsf_joint']['recall_v_le_20']:.1f}%, and {100*runtime['wpdc_spatial_epsf_joint']['recall_v_le_22']:.1f}% at V<=18, 20, and 22, respectively, versus "
             f"{100*runtime['photutils_psf']['recall_v_le_18']:.1f}%, {100*runtime['photutils_psf']['recall_v_le_20']:.1f}%, and {100*runtime['photutils_psf']['recall_v_le_22']:.1f}% for Photutils. "
             "These are crop-level CPU measurements; no full-detector or GPU throughput claim is made."))
    dolphot_boundary = insert_after(runtime_audit, style="Normal")
    rewrite(dolphot_boundary,
            "DOLPHOT gate interpretation. NGC 1851 was selected prospectively from the prespecified 11-field ACSGGCT sample because six native F606W FLC exposures were available. The retry was stopped by a declared registration gate requiring stable sub-pixel inter-exposure residuals for the common 2-pixel association and milliarcsecond-scale held-out scoring. The observed 0.698-pixel and 1.805-pixel warnings do not imply that DOLPHOT is generally unusable, nor do they justify a universal comparison claim; they show that this particular FLC set did not produce an input-registered catalogue suitable for the present protocol. The failed run is therefore retained as an audit boundary, not a DOLPHOT performance measurement.")
    cap = pfind(doc, "Table S26. Full cross-field branch comparison")
    cap.clear()
    cap.style = "Caption"
    cap.paragraph_format.keep_with_next = True
    lead = cap.add_run("Table S26.")
    set_font(lead, size=10.0, bold=True)
    rest = cap.add_run(" Full 11-field ACSGGCT high-density branch comparison with global empirical-ePSF and three-Gaussian dPSF controls. Baseline high-density recovery uses V<=20 and at least three neighbours within 10 pixels; PSF-representation control rows use their disclosed dense held-out denominator and are not pooled with baseline medians.")
    set_font(rest, size=10.0)
    cap.paragraph_format.keep_with_next = True
    rewrite(pfind(doc, "Method-definition boundary"),
            "Method-definition boundary. ‘AstroCFR spatial-ePSF + joint fit’ uses quadrant image-derived empirical PSFs and two neighbour-aware coordinate/flux fitting passes. ‘Global empirical ePSF + neighbour joint’ and ‘three-Gaussian dPSF + neighbour joint’ are literature-mapped controls run in the same image-only harness. ‘AstroCFR+Photutils hybrid’ is an archived three-field ablation and is not a complete 11-field method. RF is unavailable because the frozen simulation-training feature cache is not distributed in this workspace.")
    rewrite(pfind(doc, "Method-definition boundary"),
            "Method-definition boundary. AstroCFR spatial-ePSF + joint fit uses quadrant image-derived empirical PSFs and two neighbour-aware coordinate/flux fitting passes. Global empirical ePSF + neighbour joint and three-Gaussian dPSF + neighbour joint are diagnostic PSF-representation controls run in the same image-only harness, reported with a distinct dense held-out denominator. AstroCFR+Photutils hybrid is an archived three-field ablation and is not a complete 11-field method. HST RF is protocol-excluded: transferring the CSST-trained RF would be unvalidated domain transfer, while training on the ACSGGCT evaluation catalogue would define a separate supervised target-adaptation experiment.")
    rewrite(pfind(doc, "Scope boundary: the reported empirical evidence excludes"),
            "Scope boundary: the primary claim is limited to the 11 matched ACSGGCT F606W fields and the separate CSST registered audit. DOLPHOT/ALLFRAME, crowdsource, Euclid/VVV workflows, and CSST-PSFNet are not numerically ranked because their input requirements are not matched. The package reports their incompatibility explicitly rather than treating missing runs as failed methods.")
    rebuild_inventory(inventory_table)
    rebuild_cross_field_table(doc, base, lit, table=full_table)
    add_csst_unified_audit(doc, pfind(doc, "Table S26."))
    doc.save(DST)
    print(DST)


if __name__ == "__main__":
    main()
