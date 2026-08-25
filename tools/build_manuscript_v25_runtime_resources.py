#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Build a new manuscript revision with auditable CPU/GPU resource accounting."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v23_2026_refs_fig1.docx"
DEST = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v25_runtime_resources.docx"


def font_run(run, size=10.5, bold=False, color=(0, 0, 0)):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def clear_add(paragraph, text, first_indent=0.35, size=10.5):
    paragraph.clear()
    paragraph.style = "Normal"
    paragraph.paragraph_format.first_line_indent = Inches(first_indent) if first_indent else None
    paragraph.paragraph_format.line_spacing = 1.08
    font_run(paragraph.add_run(text), size=size)


def bookmark(paragraph, name, ident):
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(ident)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(ident))
    paragraph._p.insert(0, start); paragraph._p.append(end)


def hyperlink(paragraph, label, anchor):
    link = OxmlElement("w:hyperlink"); link.set(qn("w:anchor"), anchor); link.set(qn("w:history"), "1")
    run = OxmlElement("w:r"); props = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    props.append(color); props.append(underline); run.append(props)
    text = OxmlElement("w:t"); text.text = label; run.append(text); link.append(run)
    paragraph._p.append(link)


def set_cell(cell, text, bold=False, size=8.6, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]; p.alignment = align; p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    font_run(p.add_run(str(text)), size=size, bold=bold)


def border(cell, edge, val="single", sz="8", color="000000"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None: borders = OxmlElement("w:tcBorders"); tcPr.append(borders)
    tag = "w:" + edge; element = borders.find(qn(tag))
    if element is None: element = OxmlElement(tag); borders.append(element)
    element.set(qn("w:val"), val); element.set(qn("w:sz"), sz); element.set(qn("w:color"), color)


def three_line_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            border(cell, "top", "single", "14" if row_i == 0 else "8")
            border(cell, "bottom", "single", "14" if row_i == len(table.rows) - 1 else "0")
            border(cell, "left", "nil", "0"); border(cell, "right", "nil", "0")


def insert_after(cursor, element):
    cursor.addnext(element)
    return element


def main():
    doc = Document(SOURCE)

    # Keep the abstract quantitative but explicitly bounded by the disclosed protocol.
    abstract = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Abstract")
    abstract_text = ("We present WPDC, a deployable dense stellar-field processing system integrating adaptive background modelling, candidate generation, target-domain RandomForest adaptation, empirical-PSF modelling, and residual deblending. On four CSST-like simulated chips (approximately 4,000 reference sources), WPDC achieves mean recall of 94.6% (90.8–96.9%) with 100% simulated-catalogue precision under the disclosed matching and per-chip threshold protocol, astrometric RMS of 10.9 mas (8.2–17.2 mas), and photometric RMS of 0.066 mag (0.058–0.080 mag). Under the same simulation protocol, calibrated SExtractor attains 13.2–22.8% precision and 2–6 times larger measurement RMS. On real HST/ACS crowded-field benchmarks, the recovery advantage is replicated: on dense NGC 6752 V≤20 references, WPDC ePSF+deblend recovers 87.6% (95% CI: 84.0–90.4%) versus 57.0% for DAO and 28.6% for SEP; on NGC 1851, the spatial-ePSF branch reaches 86.3% (84.4–88.0%) versus 34.5% for DAO/Photutils. Target adaptation is sample-efficient: one 200 × 200 pixel calibration tile increases held-out recall on NGC 6752 from 0.057 to 0.648. WPDC therefore provides a bounded crowded-field recovery benefit at disclosed computational cost (28 s/MPix for the ePSF branch versus 0.12 s/MPix for DAO), whereas Photutils provides lower position and magnitude RMS. These results support a modular system with explicit operating-point trade-offs, not a universally superior method.")
    clear_add(doc.paragraphs[abstract + 1], abstract_text, first_indent=0)

    # Add reproducibility details to the existing uncertainty/runtime subsection.
    runtime_heading = next(p for p in doc.paragraphs if p.text.strip() == "5.14 Measurement uncertainty and runtime repeatability")
    runtime_para = runtime_heading._p.getnext()
    while runtime_para is not None and runtime_para.tag != qn("w:p"):
        runtime_para = runtime_para.getnext()
    runtime_text = ("To avoid treating one point estimate as a deployment guarantee, we added two uncertainty analyses. For position and magnitude RMS, the held-out matches, affine registration, robust residual mask, and photometric zero point are fixed, and the retained test residuals are resampled 1,000 times. These are conditional residual-bootstrap intervals: they quantify measurement scatter conditional on the detected catalogue, not the uncertainty of the whole detection pipeline. Runtime and RSS are measured with five repeated warm runs after one warm-up run, using the same 1200 × 1200 crop and the same method-stage timing wrapper. All controlled runs used one Python process on an Intel Core Ultra 9 275HX host (24 logical cores, 50.9 GB RAM), Python 3.12.7, NumPy 2.5.1, SciPy 1.18.0, Astropy 7.2.0, Photutils 3.0.0, and scikit-learn 1.9.0. The methods were executed on CPU only; no CUDA or GPU kernels were invoked. Although an NVIDIA RTX 5060 Laptop GPU was present, GPU memory is therefore not applicable (N/A). Runtime is reported in seconds per million processed pixels and RSS as the peak process-memory increase relative to the warm baseline.")
    if runtime_para is None: raise RuntimeError("Runtime paragraph not found")
    p_runtime = next(p for p in doc.paragraphs if p._p is runtime_para)
    clear_add(p_runtime, runtime_text, first_indent=0)

    # Insert a compact deployment operating-point table before Discussion.
    discussion = next(p for p in doc.paragraphs if p.text.strip() == "6 Discussion")
    heading = doc.add_paragraph("5.18 Accuracy–cost operating points", style="Heading 1")
    prose = doc.add_paragraph(style="Normal")
    clear_add(prose, "The controlled measurements expose distinct deployment operating points rather than a single winner. DAOStarFinder and WPDC-RF are appropriate for low-latency first-pass screening, Photutils is preferable when positional and photometric RMS dominate, and the WPDC ePSF branches trade CPU time for higher recovery in dense regions. The table reports the NGC 6752 controlled test under identical crop and matching rules; the GPU column is N/A because all branches were executed on CPU.")
    refp = doc.add_paragraph(style="Normal"); refp.paragraph_format.first_line_indent = Inches(0.35); refp.paragraph_format.line_spacing = 1.08
    font_run(refp.add_run("The corresponding quantitative results are summarized in "))
    hyperlink(refp, "Table 22", "table_22")
    font_run(refp.add_run("."))
    caption = doc.add_paragraph("Table 22. Accuracy–cost operating points on the NGC 6752 controlled test. Runtime and RSS are CPU measurements; GPU memory is not applicable because no GPU path was used.", style="Caption")
    bookmark(caption, "table_22", 2022)
    table = doc.add_table(rows=1, cols=7)
    headers = ["Method", "Operating point", "Dense V≤20 recovery", "Pos. RMS / mas", "Mag. RMS / mag", "CPU / s MPix⁻¹", "Peak RSS / MB"]
    for cell, value in zip(table.rows[0].cells, headers): set_cell(cell, value, bold=True, size=8.2)
    rows = [
        ("DAOStarFinder", "Fast first-pass screening", "57.0% [52.1–61.7]", "2.13", "0.058", "0.12", "39.9"),
        ("Photutils PSFPhotometry", "Precision-oriented measurement", "57.0% [52.1–61.7]", "1.05", "0.039", "8.72", "124.8"),
        ("WPDC-RF", "Fast conservative catalogue", "52.7% [47.9–57.6]", "1.99", "0.153", "0.51", "40.6"),
        ("WPDC ePSF+deblend", "Crowded-field recovery", "87.6% [84.0–90.4]", "1.35", "0.042", "28.0", "53.7"),
        ("WPDC spatial-ePSF+joint", "Recovery–measurement balance", "88.3% [84.8–91.1]", "1.27", "0.037", "33.1", "53.0"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row): set_cell(cell, value, size=8.2, align=WD_ALIGN_PARAGRAPH.LEFT if cell is cells[0] or cell is cells[1] else WD_ALIGN_PARAGRAPH.CENTER)
    three_line_table(table)

    # Put the new block immediately before section 6, preserving document order.
    cursor = discussion._p.getprevious()
    # Remove the newly created nodes from their end position, then insert in order.
    nodes = [heading._p, prose._p, refp._p, caption._p, table._element]
    for node in nodes: node.getparent().remove(node)
    for node in nodes: cursor = insert_after(cursor, node)

    # Explicitly document the negative group-wise prototype in the limitations.
    limitations = next(p for p in doc.paragraphs if p.text.strip() == "6.3 Limitations")
    next_p = limitations._p.getnext()
    while next_p is not None and next_p.tag != qn("w:p"): next_p = next_p.getnext()
    if next_p is not None:
        p_lim = next(p for p in doc.paragraphs if p._p is next_p)
        old_text = p_lim.text
        addition = " An exploratory group-wise joint-background prototype slightly improved NGC 6752 dense-field recovery and position RMS but increased CPU cost and worsened magnitude RMS; it is retained as a supplementary ablation rather than a claimed replacement."
        clear_add(p_lim, old_text + addition)

    doc.save(DEST)
    print(DEST)


if __name__ == "__main__": main()
