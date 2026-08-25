#!/usr/bin/env python
"""Build v30 with the final pre-submission audit corrections.

This leaves v29 unchanged.  It corrects the three audited bibliography
metadata/DOI faults, restores all figure/table/reference internal links, adds
external DOI/data/code hyperlinks, and uses the Table 18 repeated-run medians
as the sole NGC 6752 runtime/RSS values wherever those operating points recur.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v29_cross_task_audit.docx"
DEST = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v30_submission_fixes.docx"


def set_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def append_text(paragraph, text, size=10.5, bold=False):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    return run


def remove_bookmarks(paragraph, name):
    for node in list(paragraph._p.findall(qn("w:bookmarkStart"))):
        if node.get(qn("w:name")) == name:
            ident = node.get(qn("w:id"))
            paragraph._p.remove(node)
            for end in list(paragraph._p.findall(qn("w:bookmarkEnd"))):
                if end.get(qn("w:id")) == ident:
                    paragraph._p.remove(end)


def bookmark(paragraph, name, ident):
    remove_bookmarks(paragraph, name)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(ident))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(ident))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def internal_link(paragraph, label, anchor, size=10.5):
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)
    link.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts"); fonts.set(qn("w:ascii"), "Times New Roman"); fonts.set(qn("w:hAnsi"), "Times New Roman")
    props.extend([fonts, color, underline])
    run.append(props)
    text = OxmlElement("w:t"); text.text = label; run.append(text)
    link.append(run); paragraph._p.append(link)


def external_link(paragraph, label, url, size=10.5):
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    link.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts"); fonts.set(qn("w:ascii"), "Times New Roman"); fonts.set(qn("w:hAnsi"), "Times New Roman")
    props.extend([fonts, color, underline])
    run.append(props)
    text = OxmlElement("w:t"); text.text = label; run.append(text)
    link.append(run); paragraph._p.append(link)


def rewrite_plain(paragraph, text, indent=0.35):
    paragraph.clear()
    paragraph.style = "Normal"
    paragraph.paragraph_format.first_line_indent = Inches(indent) if indent else None
    paragraph.paragraph_format.line_spacing = 1.08
    append_text(paragraph, text)


def rewrite_reference(paragraph, text, number):
    paragraph.clear()
    paragraph.style = "Normal"
    paragraph.paragraph_format.left_indent = Inches(0.20)
    paragraph.paragraph_format.first_line_indent = Inches(-0.20)
    paragraph.paragraph_format.line_spacing = 1.0
    bookmark(paragraph, f"ref_{number}", 4000 + number)
    cursor = 0
    for match in re.finditer(r"https?://\S+", text):
        if match.start() > cursor:
            append_text(paragraph, text[cursor:match.start()])
        url = match.group(0).rstrip(".")
        external_link(paragraph, url, url)
        cursor = match.end()
    if cursor < len(text):
        append_text(paragraph, text[cursor:])


def set_cell(cell, text, size=8.2):
    cell.text = str(text)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        set_font(run, size=size)


def replace_linked_phrase(paragraph, old, label, anchor):
    text = paragraph.text
    if old not in text:
        raise RuntimeError(f"Expected phrase not found: {old}")
    before, after = text.split(old, 1)
    paragraph.clear(); paragraph.style = "Normal"; paragraph.paragraph_format.first_line_indent = Inches(0.35); paragraph.paragraph_format.line_spacing = 1.08
    append_text(paragraph, before)
    internal_link(paragraph, label, anchor)
    append_text(paragraph, after)


def replace_text_nodes(paragraph, replacements):
    """Replace visible text without discarding existing hyperlink elements."""
    for node in paragraph._p.iter(qn("w:t")):
        if node.text:
            for old, new in replacements.items():
                node.text = node.text.replace(old, new)


def main():
    doc = Document(SOURCE)

    # Canonical repeated-run runtime/RSS values: Table 18 is the registered
    # five-repeat source, so Tables 15/17/22 and all prose now agree with it.
    t15 = doc.tables[14]
    canonical_t15 = [("0.11", "39.9"), ("5.18", "5.5"), ("8.64", "133.2"),
                     ("0.48", "39.3"), ("24.24", "50.8")]
    for row, (runtime, rss) in zip(t15.rows[6:11], canonical_t15):
        set_cell(row.cells[7], runtime); set_cell(row.cells[8], rss)
    t17 = doc.tables[16]
    set_cell(t17.rows[1].cells[6], "24.24"); set_cell(t17.rows[1].cells[7], "50.8")
    set_cell(t17.rows[2].cells[6], "32.23"); set_cell(t17.rows[2].cells[7], "50.8")
    t22 = doc.tables[21]
    canonical_t22 = [("0.11", "39.9"), ("8.64", "133.2"), ("0.48", "39.3"),
                     ("24.24", "50.8"), ("32.23", "50.8")]
    for row, (runtime, rss) in zip(t22.rows[1:], canonical_t22):
        set_cell(row.cells[5], runtime); set_cell(row.cells[6], rss)

    abstract = next(p for p in doc.paragraphs if p.text.startswith("We present WPDC,"))
    rewrite_plain(abstract, abstract.text.replace("28 s/MPix for the ePSF branch versus 0.12 s/MPix for DAO", "24.24 s/MPix for the ePSF branch versus 0.11 s/MPix for DAO"), indent=0)
    runtime_discussion = next(p for p in doc.paragraphs if p.text.startswith("The controlled HST evaluation now reports"))
    rewrite_plain(runtime_discussion, runtime_discussion.text.replace("approximately 230 times slower than DAOStarFinder on NGC 6752 (28.0 versus 0.12 s/MPix)", "approximately 220 times slower than DAOStarFinder on NGC 6752 (24.24 versus 0.11 s/MPix)"))

    # v26 reset this caption, losing the v25 bookmark; v28 never created a
    # Table 24 bookmark/link.  Restore both without changing caption wording.
    table22_caption = next(p for p in doc.paragraphs if p.text.startswith("Table 22."))
    bookmark(table22_caption, "table_22", 3022)
    table24_caption = next(p for p in doc.paragraphs if p.text.startswith("Table 24."))
    bookmark(table24_caption, "table_24", 3024)
    table24_prose = next(p for p in doc.paragraphs if p.text.startswith("The results are reported in Table 24."))
    replace_linked_phrase(table24_prose, "Table 24", "Table 24", "table_24")

    # MNRAS volume 527(4) was formally published on 23 December 2023.
    # Keep the linked in-text years synchronized with the corrected reference.
    csst_literature = next(p for p in doc.paragraphs if "CSST-focused studies have also investigated" in p.text)
    replace_text_nodes(csst_literature, {
        "Zhang et al., 2024a": "Zhang et al., 2023",
        "Zhang et al., 2024b": "Zhang et al., 2024",
    })

    # Correct reference metadata identified by DOI/Crossref audit.
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "References")
    refs = [p for p in paragraphs[start + 1:] if p.text.strip()]
    if len(refs) != 41:
        raise RuntimeError(f"Expected 41 references, found {len(refs)}")
    corrected = [p.text.strip() for p in refs]
    corrected[0] = ("Anderson, J., & King, I. R. (2006). PSFs, photometry, and astronomy for the ACS/WFC. "
                    "ACS Instrument Science Report ACS 2006-01, Space Telescope Science Institute. "
                    "https://www.stsci.edu/hst/instrumentation/acs/documentation/instrument-science-reports-isrs")
    corrected[2] = ("Bellini, A., Anderson, J., & Bedin, L. R. (2011). Astrometry and photometry with HST WFC3. II. "
                    "Improved geometric-distortion corrections for 10 filters of the UVIS channel. Publications of the "
                    "Astronomical Society of the Pacific, 123, 622-637. https://doi.org/10.1086/659878")
    corrected[4] = corrected[4].replace("10.3847/1538-3881/aafc93", "10.3847/1538-3881/ab089d")
    corrected[39] = corrected[39].replace("(2024).", "(2023).", 1)
    for number, (paragraph, text) in enumerate(zip(refs, corrected), 1):
        rewrite_reference(paragraph, text, number)

    # Make external data/code destinations actual Word hyperlinks and remove the
    # obsolete claim that the exact release is only local.
    data_p = next(p for p in doc.paragraphs if p.text.startswith("Data availability:"))
    data_p.clear(); data_p.style = "Normal"; data_p.paragraph_format.first_line_indent = None; data_p.paragraph_format.line_spacing = 1.08
    append_text(data_p, "Data availability: The CSST challenge data are distributed through the ")
    external_link(data_p, "National Astronomical Data Center CSST Data Challenge 2026 page", "https://nadc.china-vo.org/events/CSSTdatachallenge2026/info/challenge_11th")
    append_text(data_p, " and are not redistributed in this repository. The HST/ACS evaluation uses public ")
    external_link(data_p, "MAST HLSP ACSGGCT v2", "https://archive.stsci.edu/prepds/acsggct/")
    append_text(data_p, " F606W images and official catalogues for NGC 6397, NGC 6752, and NGC 1851. Download URLs, byte counts, and SHA-256 hashes are recorded in data/manifest.csv. Users must follow the NADC, MAST, ACSGGCT, survey, and Gaia redistribution terms.")
    code_p = next(p for p in doc.paragraphs if p.text.startswith("Code availability:"))
    code_p.clear(); code_p.style = "Normal"; code_p.paragraph_format.first_line_indent = None; code_p.paragraph_format.line_spacing = 1.08
    append_text(code_p, "Code availability: The designated public repository for the manuscript-matched source release is ")
    external_link(code_p, "github.com/zxl1999/WPDC", "https://github.com/zxl1999/WPDC")
    append_text(code_p, ". The local manuscript-matched release package contains src/wpdc for reusable modules, experiments/hst for controlled HST baselines and artificial-star runs, experiments/csst for CSST adapters, data for provenance instructions, environment for package locks, results for machine-readable summaries, supplementary material, and tools for manuscript generation. It is identified locally by tag v1.0.0 and is prepared for publication at the stated GitHub URL; the public tag and archive URL must be cited after repository credentials are available. An archival DOI may be added through Zenodo after repository release, but no DOI is claimed here before it exists.")

    doc.save(DEST)
    print(DEST)


if __name__ == "__main__":
    main()
