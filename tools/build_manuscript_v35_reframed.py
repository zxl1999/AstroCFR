#!/usr/bin/env python
"""Reframe v34 as a crowded-field methods/system paper, not multimedia work."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import internal_link, set_font
from build_manuscript_v33_slim import format_three_line

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "AstroCFR_Multimedia_Systems_SCI_manuscript_v34_scalability.docx"
DEST = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v35_reframed.docx"
SUP_SOURCE = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v34.docx"
SUP_DEST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v35.docx"
SENSITIVITY = ROOT / "results" / "hst_parameter_sensitivity" / "parameter_sensitivity_audit.json"
FIG_DIR = ROOT / "results" / "astrocfr_supplementary_figures"


def clear(paragraph, indent=.35):
    paragraph.clear(); paragraph.style = "Normal"
    paragraph.paragraph_format.first_line_indent = Inches(indent) if indent else None
    paragraph.paragraph_format.line_spacing = 1.08


def text(paragraph, value, size=10.5, bold=False):
    run = paragraph.add_run(value); set_font(run, size=size, bold=bold); return run


def cite(paragraph, label, anchor):
    internal_link(paragraph, label, anchor)


def linked(paragraph, pieces, indent=.35):
    clear(paragraph, indent)
    for piece in pieces:
        if isinstance(piece, tuple): cite(paragraph, *piece)
        else: text(paragraph, piece)


def replace_plain(paragraph, value, indent=.35):
    clear(paragraph, indent); text(paragraph, value)


def find(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def remove_elements_between(doc, start_prefix, end_prefix):
    start = find(doc, start_prefix)._p
    end = find(doc, end_prefix)._p
    node = start
    while node is not end:
        following = node.getnext()
        node.getparent().remove(node)
        node = following


def insert_after(paragraph, value, indent=.35):
    new = paragraph._parent.add_paragraph(style="Normal")
    new._p.getparent().remove(new._p); paragraph._p.addnext(new._p)
    replace_plain(new, value, indent)
    return new


def renumber_figures(doc):
    # Figs 9--12 move to the supplement; old Fig. 13 becomes new Fig. 9.
    mapping = {number: number if number <= 8 else number - 4 for number in range(13, 24)}
    rx = re.compile(r"\b(Fig(?:ure)?\.?)\s+(\d+)\b")
    for node in doc._element.body.iter(qn("w:t")):
        if node.text:
            node.text = rx.sub(lambda m: f"{m.group(1)} {mapping.get(int(m.group(2)), int(m.group(2)))}", node.text)
    for bookmark in doc._element.body.iter(qn("w:bookmarkStart")):
        name = bookmark.get(qn("w:name")) or ""
        match = re.fullmatch(r"figure_(\d+)", name)
        if match and int(match.group(1)) in mapping:
            bookmark.set(qn("w:name"), f"figure_{mapping[int(match.group(1))]}")
    for hyperlink in doc._element.body.iter(qn("w:hyperlink")):
        anchor = hyperlink.get(qn("w:anchor")) or ""
        match = re.fullmatch(r"figure_(\d+)", anchor)
        if match and int(match.group(1)) in mapping:
            hyperlink.set(qn("w:anchor"), f"figure_{mapping[int(match.group(1))]}")


def standardize_branch_names(doc):
    """Map historical result labels to one manuscript-facing nomenclature."""
    replacements = (
        ("WPDC original (target-adapted RF)", "AstroCFR-RF (target-adapted)"),
        ("WPDC ePSF + residual deblend", "AstroCFR ePSF-deblend"),
        ("WPDC spatial ePSF + joint fit", "AstroCFR spatial-ePSF joint fit"),
        ("AstroCFR ePSF + residual deblend", "AstroCFR ePSF-deblend"),
        ("AstroCFR spatial ePSF + joint fit", "AstroCFR spatial-ePSF joint fit"),
        ("spatial-ePSF+joint fit", "spatial-ePSF joint fit"),
        ("ePSF+deblend", "ePSF-deblend"),
    )
    for node in doc._element.body.iter(qn("w:t")):
        if node.text:
            for old, new in replacements:
                node.text = node.text.replace(old, new)


def main_manuscript():
    doc = Document(SOURCE)
    doc.paragraphs[0].clear(); text(doc.paragraphs[0],
        "AstroCFR: A Modular Candidate-Recovery and Measurement Framework for Crowded Stellar Fields", size=15, bold=True)
    replace_plain(doc.paragraphs[6],
        "Keywords: crowded-field photometry; source detection; candidate recovery; empirical PSF; target-domain adaptation; CSST", 0)
    replace_plain(doc.paragraphs[5],
        "We present AstroCFR, a modular candidate-recovery and measurement framework for crowded stellar fields. The system combines adaptive background estimation, multi-branch candidate generation, target-domain RandomForest screening, empirical-PSF fitting, residual deblending, and validation-gated calibration. On four CSST-like simulated chips containing approximately 4,000 reference sources, the registered configuration reaches a mean recall of 94.6% and 100% simulation-catalogue precision; this precision is specific to the supplied simulation catalogue and assembly protocol. On three public HST/ACS fields, the recovery-oriented ePSF branches improve crowded-field recovery under a common 2-pixel association rule, while Photutils PSFPhotometry provides lower positional and photometric RMS on the controlled NGC 6752 test. A fixed density-stratified policy recovers 50.00% of 800 identical artificial-star injections at 20.44 s/MPix, between Photutils-only (39.25%, 8.64 s/MPix) and always-on spatial-ePSF processing (60.38%, 32.23 s/MPix). One spatially isolated 200 x 200 pixel target-calibration tile raises held-out NGC 6752 recall from 0.057 to 0.648, but the external zero-shot screen fails on Pan-STARRS1 and Legacy Survey images. AstroCFR is therefore presented as an auditable set of recovery–precision–cost operating points, not as a universal replacement for Photutils or multi-exposure DOLPHOT/ALLFRAME-class photometry.")
    replace_plain(doc.paragraphs[8],
        "Crowded stellar-field photometry is a difficult source-extraction and measurement problem. In the Galactic bulge, globular clusters, and nearby resolved galaxies, overlapping point-spread functions (PSFs), spatially varying backgrounds, and bright-star artifacts jointly reduce catalogue completeness and distort positions and fluxes. A useful pipeline must therefore make its candidate-generation, deblending, measurement, and catalogue-release decisions explicit rather than treating a local-maximum detector as a final scientific product.")
    linked(doc.paragraphs[10], [
        "The measurement layer is grounded in empirical HST PSF and crowded-field catalogue practice (",
        ("Anderson & King, 2006", "ref_1"), "; ", ("Anderson et al., 2008", "ref_2"), "; ",
        ("Bellini et al., 2011", "ref_3"), "). Survey reference systems define external astrometric context (",
        ("Chambers et al., 2016", "ref_4"), "; ", ("Dey et al., 2019", "ref_5"), "; ",
        ("Gaia Collaboration et al., 2016", "ref_6"), "; ", ("Gaia Collaboration et al., 2023", "ref_7"), "). "
        "Astropy, Photutils, SEP, SciPy, and scikit-learn provide the reproducible implementation components (",
        ("Astropy Collaboration, 2022", "ref_8"), "; ", ("Harris et al., 2020", "ref_9"), "; ",
        ("Hunter, 2007", "ref_10"), "; ", ("Virtanen et al., 2020", "ref_11"), "; ",
        ("Barbary, 2016", "ref_12"), "; ", ("Bradley et al., 2024", "ref_13"), "; ",
        ("Pedregosa et al., 2011", "ref_14"), "). Deep-learning references provide context for the optional image classifier, not the central claim (",
        ("Goodfellow et al., 2016", "ref_15"), "; ", ("He et al., 2016", "ref_16"), "). Recent studies address CSST preparation, source detection, PSF uncertainty and reconstruction, and crowded-field measurement (",
        ("Shi et al., 2024", "ref_17"), "; ", ("Long et al., 2025", "ref_18"), "; ", ("Han et al., 2026", "ref_19"), "; ",
        ("Wainer et al., 2025", "ref_20"), "; ", ("Espinosa et al., 2025", "ref_21"), "; ", ("De Alba et al., 2026", "ref_22"), "; ",
        ("Centofanti et al., 2026", "ref_23"), "; ", ("Wang et al., 2026", "ref_24"), "; ", ("Zhang et al., 2026", "ref_25"), "; ",
        ("Libralato et al., 2024", "ref_26"), "; ", ("Salaris et al., 2024", "ref_27"), "). Transfer terminology and classifier baselines follow (",
        ("Ben-David et al., 2010", "ref_28"), "; ", ("Pan & Yang, 2010", "ref_29"), "; ", ("Ganin et al., 2016", "ref_30"), "; ",
        ("Long et al., 2015", "ref_31"), "; ", ("Breiman, 2001", "ref_32"), "; ", ("Chen & Guestrin, 2016", "ref_33"), ")."
    ])
    linked(doc.paragraphs[11], [
        "The relevant methodological map has four parts. Classical source detection and PSF photometry are represented by DAOPHOT (",
        ("Stetson, 1987", "ref_34"), ") and SExtractor (", ("Bertin & Arnouts, 1996", "ref_35"), "). "
        "DOLPHOT (", ("Dolphin, 2000", "ref_36"), ") and survey-scale crowded-field pipelines (", ("Schlafly et al., 2018", "ref_37"), ") provide stronger joint-fitting or multi-exposure measurement baselines than the single-image experiment studied here. Modern deblenders address blend separation (", ("Melchior et al., 2018", "ref_38"), "). "
        "Accordingly, AstroCFR does not claim to replace DOLPHOT/ALLFRAME-class multi-exposure measurement; it studies candidate recovery and single-image operating points under one common protocol. CSST-focused work has separately addressed dense-field astrometry, photometry, and classification (",
        ("Wang et al., 2024", "ref_39"), "; ", ("Zhang et al., 2023", "ref_40"), "; ", ("Zhang et al., 2024", "ref_41"), ")."
    ])
    linked(doc.paragraphs[12], [
        "Recent detector-level work on blends, density classification, distortion, PSF fitting, and segmentation (",
        ("Yan et al., 2026", "ref_42"), "; ", ("Lai et al., 2026", "ref_43"), "; ", ("Yan et al., 2026", "ref_44"), "; ",
        ("Nie et al., 2025", "ref_45"), "; ", ("Shaw et al., 2025", "ref_46"), "; ", ("Burke et al., 2019", "ref_47"), ") motivates individual modules. AstroCFR's limited contribution is their reproducible comparison as candidate-recovery, measurement, and cost operating points."
    ])
    replace_plain(doc.paragraphs[13],
        "AstroCFR is therefore a modular crowded-field processing framework, not a universal replacement for specialised photometric software and not a new detector family. Its contribution is a common protocol that makes candidate recovery, measurement precision, target-specific calibration, and computational cost visible together. The resulting branches are intended as explicit operating points: AstroCFR-RF for conservative screening, and ePSF-based branches for recovery-oriented analysis in sufficiently crowded regions.")
    for prefix in ("The manuscript is organized as follows.", "The corresponding visualization is shown in Fig. 1"):
        paragraph = find(doc, prefix); paragraph._element.getparent().remove(paragraph._element)
    replace_plain(find(doc, "Fig. 1."), "Fig. 1. AstroCFR architecture and decision flow for crowded-field candidate recovery and measurement.", 0)
    # Section 2.2 is no longer a multimedia rationale.
    heading = find(doc, "2.2 Multimedia system perspective")
    heading.clear(); text(heading, "2.2 System scope and decision boundary", size=10.5, bold=True)
    replace_plain(doc.paragraphs[35],
        "AstroCFR is treated as a modular astronomical processing chain. It produces candidate lists, local image cutouts, empirical PSFs, residual diagnostics, and calibrated catalogues, but these intermediate products are not claimed as a separate conceptual contribution. The system boundary is deliberately narrow: this paper evaluates single-image candidate recovery and measurement operating points, not multi-epoch catalogue production or universal survey throughput.")
    redundant_scope = find(doc, "AstroCFR treats astronomical image processing as a structured multimedia")
    redundant_scope._element.getparent().remove(redundant_scope._element)
    replace_plain(doc.paragraphs[39],
        "AstroCFR first estimates a two-dimensional background and RMS map using block-wise robust statistics with median filtering and interpolation. The CSST-like challenge configuration uses a nominal 2.5-sigma proposal threshold; chip 18 uses 2.39 sigma after validation because of its higher background RMS. These are registered dataset-specific operating parameters, not universal constants. The public-HST benchmark uses a common 3-sigma front end, and Supplementary Tables S5–S6 report threshold and matching-radius sensitivity separately from final catalogue selection.")
    replace_plain(doc.paragraphs[47],
        "The dominant dense-field failure mode is structural blending rather than isolated-source sensitivity. AstroCFR constructs candidate groups within five pixels in the registered CSST-like configuration and escalates recovery through three levels: paired-PSF fitting for two-source groups, joint multi-PSF least-squares fitting for larger groups, and Richardson-Lucy deconvolution only for severe or saturated groups. The five-pixel grouping scale is an implementation parameter tied to the simulated PSF sampling, not a claimed instrument-independent optimum.")
    replace_plain(doc.paragraphs[59],
        "Astrometric refinement uses a three-stage cascade: a third-order polynomial removes large-scale WCS distortion, a look-up table interpolates residual spatial distortion, and Gaussian-process regression is considered only when validation residuals retain spatial structure. The polynomial degree and grid settings are registered challenge-configuration choices rather than universal optima. Local centroid refinement is accepted only when the validation set improves, which prevents flexible correction layers from being silently retained on training residuals alone.")
    replace_plain(doc.paragraphs[66],
        "The quality gates have an explicit order rather than acting as an unspecified vote. Broad image-only proposals are first screened by morphology and SNR; candidate probabilities then define the conservative RF catalogue; residual-improvement tests govern whether a blended companion is retained; and astrometric or photometric correction layers are accepted only after validation improvement. When a classifier probability and a physical residual test disagree, the residual test governs companion acceptance and the conservative catalogue threshold governs release. Reference catalogues are used for the disclosed simulation evaluation and target-adaptation calibration partitions, not as an inference-time gate in the HST baseline branches.")
    replace_plain(doc.paragraphs[67],
        "The registered gates are adaptive proposal thresholding from background RMS, morphology limits before classification, isolated high-SNR bright-source exceptions, residual-improvement acceptance for deblending, validation-selected classifier thresholds, validation-selected astrometric corrections, and K-fold-validated photometric refinements. Each can be disabled or changed in a new instrument configuration; the present results establish only the disclosed configuration.")
    replace_plain(doc.paragraphs[71],
        "Detected sources are associated with reference stars through a greedy one-to-one two-pixel rule in the primary protocol. In simulations, recall and precision are computed against the complete supplied challenge reference catalogue. In HST and external-survey tests, the official catalogue association fraction is reported only as a catalogue-match lower bound because deeper or unmatched real sources can exist; it is not called blind purity. Astrometric and photometric RMS are computed from matched held-out residuals.")
    replace_plain(doc.paragraphs[76],
        "A calibrated SExtractor baseline is run on the same simulated exposures after correcting the saturation level to the 16-bit ADU ceiling. It is evaluated as a general-purpose source-extraction baseline under the same association rule, not as a star-galaxy classification experiment. Its threshold is tuned per chip by the same validation principle, while its reported precision remains simulation-catalogue precision and should not be conflated with the HST catalogue-match lower bound.")
    replace_plain(doc.paragraphs[77],
        "The SExtractor comparison is diagnostic rather than a universal ranking. AstroCFR is specialised for the registered CSST-like dense fields and contains extra deblending and candidate-screening stages. The comparison quantifies how the selected operating points differ under one shared truth catalogue; it does not imply that an unmodified SExtractor workflow is the appropriate scientific baseline for every crowded-field measurement task.")
    # Remove the repetitive lead-in sentences and the four internal-report figures.
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith("The corresponding quantitative results are summarized in Table") or paragraph.text.startswith("The corresponding visualization is shown in Fig.") or paragraph.text.startswith("The corresponding failure-case visualization") or paragraph.text.startswith("The corresponding stratified-recovery visualization"):
            paragraph._element.getparent().remove(paragraph._element)
    remove_elements_between(doc, "The original diagnostic panels are retained", "5.2 Simulation-to-real target adaptation")
    # Retain a short, unambiguous negative-control statement.
    replace_plain(find(doc, "The external test exposes a strong domain shift"),
        "The frozen simulation-trained screen fails on the external domains: retained recall is 2.8% on Pan-STARRS1 M31 and 0.0% on Legacy Survey M13 under the stated Gaia-limited evaluation. This negative control is reported briefly to delimit the method, not as evidence of zero-shot generalization. The following target-adapted analysis is supervised few-shot domain adaptation.")
    replace_plain(find(doc, "Calibrated SExtractor reaches substantial recall"),
        "On the CSST-like simulations, calibrated SExtractor reaches 83.7–91.6% recall and 13.2–22.8% simulation-catalogue precision under the disclosed settings. AstroCFR's 100% figure is a competition-catalogue result after its full reference-aware assembly protocol; it should not be read as a directly transferable blind-purity advantage over SExtractor. The relevant conclusion is that the two pipelines occupy different recovery–false-positive operating points on this registered simulation task.")
    replace_plain(find(doc, "The candidate-count trajectory reveals"),
        "The candidate-count trajectory illustrates a candidate-filtering trade-off. AstroCFR begins with hundreds of thousands to more than one million proposals per chip and reduces them to approximately one thousand high-confidence simulation-catalogue entries. This compression supports the conservative branch but does not establish blind purity for a deeper real field.")
    # Insert the registered sensitivity result after the SExtractor comparison.
    marker = find(doc, "On the CSST-like simulations, calibrated SExtractor")
    insert_after(marker,
        "Parameter sensitivity is reported as an audit rather than an optimisation claim. On the NGC 6752 HST test partition, changing the common proposal threshold from 2.5 to 4.0 sigma changes DAO-style proposal completeness from 0.771 to 0.750 while its catalogue-match lower bound changes from 0.937 to 0.946. Holding catalogues fixed and varying the association radius from 1 to 3 pixels preserves the qualitative ordering: ePSF+deblend completeness is 0.926–0.932, spatial-ePSF+joint fit is 0.932–0.935, DAO/Photutils is about 0.762, and SEP is 0.491–0.526. Full values are in Supplementary Tables S5–S6.")
    replace_plain(find(doc, "The controlled HST evaluation now reports wall-clock"),
        "The controlled HST evaluation reports wall-clock seconds per megapixel and peak process RSS for every branch. These are single-machine relative costs, not hardware-independent throughput guarantees. In particular, AstroCFR ePSF+deblend is approximately 220 times slower than DAOStarFinder on NGC 6752 (24.24 versus 0.11 s/MPix). It is therefore a scientific recovery/measurement branch for selected crowded regions, not a survey-wide fast-processing branch; AstroCFR-RF and DAO-style proposals are the appropriate fast front ends.")
    replace_plain(find(doc, "The threshold experiment shows that evaluation protocol"),
        "The threshold experiment shows that evaluation protocol can be a hidden confound. A global classifier threshold assumes that all chips share a probability calibration, whereas the faint chip requires a different recovery–false-positive balance. Per-chip thresholds are therefore validation-selected deployment parameters in the registered simulation experiment, not evidence that one universal classifier transfers unchanged across detectors.")
    broader = find(doc, "This point is broader than astronomy")
    broader._element.getparent().remove(broader._element)
    replace_plain(find(doc, "AstroCFR demonstrates a deployable decomposition"),
        "AstroCFR provides a reproducible comparison framework for crowded-field candidate recovery and single-image measurement, not a universal photometric replacement. The simulation results describe a registered challenge configuration; the controlled HST comparison is the central evidence because methods share image crops, associations, spatial partitions, and injected scenes.")
    replace_plain(find(doc, "The controlled results support three limited conclusions"),
        "Three bounded conclusions follow. First, in dense NGC 6752 and NGC 1851 subsets, the ePSF-based AstroCFR branches recover more reference and injected sources than the evaluated DAO, SEP, and RF operating points. Second, this is not a universal measurement win: Photutils provides the lowest reported NGC 6752 positional and photometric RMS, while DOLPHOT/ALLFRAME-class multi-exposure measurement remains outside the present comparison. Third, the recovery branch incurs a large CPU cost and should be deployed only where its high-crowding recovery benefit justifies that cost. The spatial-ePSF/two-pass variant improves AstroCFR photometric RMS to 0.037 mag but does not overturn the Photutils positional result.")
    replace_plain(find(doc, "Code availability:"),
        "Code availability: The manuscript-matched source is prepared for https://github.com/zxl1999/AstroCFR. The release contains reusable modules in src/wpdc, controlled HST and CSST experiment scripts, machine-readable summaries, environment locks, data provenance, and manuscript builders. The version is v1.3.0; the public commit hash, tag, and release archive must be recorded after upload. No archival DOI is claimed before a public release exists.")
    renumber_figures(doc)
    standardize_branch_names(doc)
    for table in doc.tables: format_three_line(table)
    doc.save(DEST)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers): cell.text = value
    for values in rows:
        for cell, value in zip(table.add_row().cells, values): cell.text = str(value)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs: set_font(run, size=7.5, bold=row_index == 0)
    format_three_line(table)


def picture_blob(source, index, destination):
    shape = source.inline_shapes[index]
    rid = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
    destination.write_bytes(source.part.related_parts[rid].blob)


def supplement():
    source = Document(SOURCE)
    doc = Document(SUP_SOURCE)
    sensitivity = json.loads(SENSITIVITY.read_text(encoding="utf-8"))
    h = doc.add_paragraph("S5 Parameter-sensitivity audit", style="Heading 1")
    for run in h.runs: set_font(run, size=12, bold=True)
    p = doc.add_paragraph("Detector threshold and association radius are reported as post-hoc robustness audits, not method-selection parameters. The threshold scan uses the common image-only DAO-style proposal front end with the primary two-pixel evaluation rule. The radius scan holds every produced catalogue fixed before changing the evaluation association radius. Catalogue-match values remain lower bounds, not blind purity.")
    for run in p.runs: set_font(run, size=10.5)
    caption = doc.add_paragraph("Table S5. Common proposal-threshold sensitivity on the NGC 6752 spatial test partition (fixed two-pixel association).", style="Caption")
    for run in caption.runs: set_font(run, size=9)
    add_table(doc, ["Threshold (sigma)", "Test candidates", "Completeness", "Catalogue-match lower bound"],
              [[f"{r['threshold_sigma']:.1f}", r["test_candidates"], f"{r['quality_completeness']:.3f}", f"{r['catalogue_match_lower_bound']:.3f}"]
               for r in sensitivity["detection_threshold_scan"]])
    caption = doc.add_paragraph("Table S6. Association-radius sensitivity with fixed NGC 6752 catalogues. Completeness is measured against the same held-out quality references.", style="Caption")
    for run in caption.runs: set_font(run, size=9)
    rows = []
    for method in ("dao", "sep", "photutils_psf", "wpdc_rf", "wpdc_epsf_deblend", "wpdc_spatial_epsf_joint"):
        values = [r for r in sensitivity["association_radius_scan"] if r["method"] == method]
        rows.append([values[0]["label"], *[f"{r['quality_completeness']:.3f}" for r in values]])
    add_table(doc, ["Fixed catalogue", "1.0 px", "1.5 px", "2.0 px", "2.5 px", "3.0 px"], rows)
    h = doc.add_paragraph("S6 Archived simulation diagnostics", style="Heading 1")
    for run in h.runs: set_font(run, size=12, bold=True)
    p = doc.add_paragraph("The following diagnostic panels were removed from the main text to maintain focus on controlled recovery, measurement, and cost results. They document calibration internals and remain available for audit; they are not additional efficacy claims.")
    for run in p.runs: set_font(run, size=10.5)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    captions = [
        "Fig. S2. WCS astrometric diagnostics for the four CSST-like chips.",
        "Fig. S3. K-fold cross-validation diagnostics for the ML magnitude refinement.",
        "Fig. S4. Photometric calibration diagnostics for the four CSST-like chips.",
        "Fig. S5. Final simulation-catalogue output diagnostics for the four chips.",
    ]
    for offset, caption_text in enumerate(captions):
        output = FIG_DIR / f"figS{offset + 2}_archived_simulation_diagnostic.png"
        picture_blob(source, offset + 8, output)
        image = doc.add_picture(str(output), width=Inches(6.25))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph(caption_text, style="Caption")
        for run in caption.runs: set_font(run, size=9)
    standardize_branch_names(doc)
    doc.save(SUP_DEST)


if __name__ == "__main__":
    main_manuscript(); supplement(); print(DEST); print(SUP_DEST)
