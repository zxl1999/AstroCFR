#!/usr/bin/env python
"""Build v42 with the closed-book pilot and multi-exposure scope audit.

This revision deliberately separates three evidential levels:

* the reference-aware CSST challenge assembly audit;
* an image-only, closed-book artificial-source proposal pilot; and
* the controlled single-stacked-image HST comparisons.

It does not promote the failed FLC registration pilot to a science result.
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import set_cell, set_font


ROOT = Path(__file__).resolve().parents[1]
MAIN_SOURCE = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v41_stratified_quality.docx"
MAIN_DEST = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v42_closed_book_scope.docx"
SUP_SOURCE = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v41.docx"
SUP_DEST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v42.docx"
FIGURE_DEST = ROOT / "results" / "csst_blind_proposal_injections_n40_crop3000" / "csst_registered_measurement_audit.png"


def find(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def rewrite(paragraph, text, indent=.35, style="Normal"):
    paragraph.clear()
    paragraph.style = style
    paragraph.paragraph_format.first_line_indent = Inches(indent) if indent else None
    paragraph.paragraph_format.line_spacing = 1.08
    set_font(paragraph.add_run(text), size=10.5, bold=style.startswith("Heading"))
    return paragraph


def after(paragraph, text, indent=.35, style="Normal"):
    new = paragraph._parent.add_paragraph(style=style)
    new._p.getparent().remove(new._p)
    paragraph._p.addnext(new._p)
    return rewrite(new, text, indent=indent, style=style)


def append(doc, text, indent=.35, style="Normal"):
    p = doc.add_paragraph(style=style)
    return rewrite(p, text, indent=indent, style=style)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, attrs in edges.items():
        tag = "w:" + edge
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        for key, value in attrs.items():
            node.set(qn("w:" + key), str(value))


def three_line_table(table):
    table.style = "Table Grid"
    nil = {"val": "nil"}
    thick = {"val": "single", "sz": "12", "color": "000000"}
    mid = {"val": "single", "sz": "8", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil, insideH=nil, insideV=nil)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=thick, bottom=mid)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=thick)


def bookmark(paragraph, name, ident):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(ident))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(ident))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def drop_table_column(table, index):
    for row in table.rows:
        row._tr.remove(row.cells[index]._tc)
    grid = table._tbl.tblGrid
    grid.remove(grid.gridCol_lst[index])


def draw_registered_audit():
    chips = np.array([12, 13, 17, 18])
    sex_recall = np.array([91.6, 87.6, 89.1, 83.7])
    astro_recall = np.array([96.9, 93.1, 96.4, 91.8])
    sex_pos = np.array([32.8, 25.5, 20.2, 23.0])
    astro_pos = np.array([17.2, 8.2, 9.1, 8.8])
    sex_mag = np.array([0.406, 0.497, 0.250, 0.392])
    astro_mag = np.array([0.0596, 0.0575, 0.0802, 0.0684])
    x = np.arange(len(chips)); width = .36
    fig, axes = plt.subplots(1, 3, figsize=(10.8, 3.25), constrained_layout=True)
    panels = [
        (sex_recall, astro_recall, "Recall (%)", (0, 105)),
        (sex_pos, astro_pos, "Position RMS (mas)", None),
        (sex_mag, astro_mag, "Magnitude RMS (mag)", None),
    ]
    for ax, (a, b, ylabel, ylim) in zip(axes, panels):
        ax.bar(x - width / 2, a, width, label="SExtractor", color="#8C8C8C")
        ax.bar(x + width / 2, b, width, label="AstroCFR", color="#2878B5")
        ax.set_xticks(x, [str(c) for c in chips])
        ax.set_xlabel("CSST-like chip")
        ax.set_ylabel(ylabel)
        ax.grid(axis="y", alpha=.22, linewidth=.6)
        if ylim:
            ax.set_ylim(*ylim)
    fig.legend(["SExtractor", "AstroCFR"], loc="upper center", ncol=2, frameon=False,
               bbox_to_anchor=(.5, 1.04))
    fig.savefig(FIGURE_DEST, dpi=300, bbox_inches="tight")
    plt.close(fig)


def main():
    draw_registered_audit()
    doc = Document(MAIN_SOURCE)
    sup = Document(SUP_SOURCE)

    # Abstract: remove the oracle-conditioned 100% value from the headline chain.
    rewrite(find(doc, "We present AstroCFR,"),
        "We present AstroCFR, a modular framework that exposes candidate-recovery, measurement-precision, and "
        "computational-cost operating points in crowded stellar fields. In the registered CSST-like simulation "
        "configuration (four chips; approximately 4,000 supplied references), its conservative target-adapted "
        "RandomForest branch reaches 94.6% mean reference recovery. In a controlled common-protocol HST/ACS "
        "NGC 6752 test, AstroCFR ePSF-deblend recovers 87.6% (95% CI 84.0-90.4%) of dense V<=20 references, "
        "compared with 57.0% for DAO and 28.6% for SEP. Photutils provides lower positional and photometric RMS, "
        "whereas the recovery-oriented AstroCFR branch requires more computation. A simulation-developed front end "
        "can be recalibrated with a small spatially disjoint labelled target region, but zero-shot transfer fails. "
        "AstroCFR therefore provides a reproducible Pareto frontier for recovery, measurement precision, and cost "
        "rather than universal photometric dominance.", indent=0)

    # Remove residual wording that implies an exhaustive simulation truth set or blind purity.
    rewrite(find(doc, "The problem is framed as dense-field source detection and calibration."),
        "The problem is framed as dense-field source detection and calibration. Given a noisy wide-field image and "
        "a supplied evaluation catalogue, the system produces point-source positions and magnitudes. Reference "
        "recovery measures how many supplied entries are associated; reference-scoped precision measures the "
        "fraction of accepted entries associated with that supplied catalogue; astrometric and photometric RMS use "
        "matched residuals. Reference-scoped precision is an audit quantity and becomes oracle-conditioned when the "
        "same catalogue participates in final assembly; it is not blind purity unless scene truth is exhaustive.")
    rewrite(find(doc, "Several safeguards are included to prevent the recovery passes"),
        "Several safeguards prevent the recovery passes from admitting obvious artifacts. Saturated bright-star "
        "recovery is constrained by local isolation and peak significance. Halo recovery is limited to neighbourhoods "
        "in which a bright source can plausibly hide a companion. Forced photometry after WCS alignment is used only "
        "when the predicted location has sufficient local SNR and does not duplicate an existing detection. These "
        "safeguards define a conservative candidate mode; they do not by themselves establish blind catalogue purity.")
    rewrite(find(doc, "A key methodological decision is per-chip threshold calibration."),
        "A key methodological decision is per-chip threshold calibration. Candidate distributions differ strongly "
        "across chips, so a global threshold can be unfairly conservative for faint chips. AstroCFR selects thresholds "
        "on validation candidates using a reference-scoped recovery/false-acceptance constraint and never uses the "
        "held-out test partition for threshold selection. The resulting validation metric remains catalogue-scoped; "
        "it is not a claim that probability calibration transfers unchanged to a blind detector.")
    rewrite(find(doc, "In a production CSST setting, the same framework could support multiple catalog modes."),
        "In a production CSST setting, the same framework could expose multiple candidate modes. A conservative mode "
        "would keep strict image-derived quality gates; a recovery mode could expose low-probability or deblend-"
        "ambiguous sources for follow-up; and a training-data mode could export rejected candidates and uncertain "
        "zones. Their blind reliability must be calibrated on exhaustive truth or deeper independent data rather than "
        "inferred from the reference-aware challenge assembly.")
    rewrite(find(doc, "Detected sources are matched to reference stars using the greedy one-to-one two-pixel rule."),
        "Detected sources are matched to supplied references using a greedy one-to-one two-pixel rule. In simulations, "
        "the resulting values are reference-scoped challenge metrics; the supplied top-1000 catalogue is not assumed "
        "to be exhaustive scene truth. If the catalogue participates in assembly, its precision-like value is oracle-"
        "conditioned. In HST and external-survey tests, association fraction is reported only as a catalogue-match "
        "lower bound because deeper real sources can be unmatched; measurement RMS uses matched held-out residuals.")
    rewrite(find(doc, "The SExtractor comparison is intentionally not framed as a universal ranking of tools."),
        "The SExtractor comparison is intentionally not framed as a universal ranking. SExtractor is a mature "
        "general-purpose extractor, whereas AstroCFR is specialised for the registered dense-field simulation and "
        "uses additional deblending and screening stages. The main comparison therefore retains reference recovery "
        "and conditional measurement residuals only; the reference-aware AstroCFR assembly value is excluded from "
        "the scientific-performance ranking.")

    # Main CSST section: separate challenge assembly from the closed-book proposal pilot.
    p = find(doc, "On the CSST-like simulations, calibrated SExtractor reaches")
    rewrite(p,
        "On the CSST-like simulations, calibrated SExtractor reaches 83.7-91.6% reference recovery under the "
        "disclosed settings. The former AstroCFR 100% value is not used as a headline precision result because final "
        "challenge assembly is reference-aware. It is retained only as an oracle-conditioned implementation audit "
        "in Supplementary Section S7 and cannot establish blind purity or false-discovery rate.")
    after(p,
        "A separate closed-book pilot is reported in Supplementary Section S9. Each candidate catalogue was written "
        "before the withheld test-chip catalogue was opened, and injected positions were unavailable to candidate "
        "generation. The image-only proposal front end recovers 63.8% and 89.4% of low-density injections at peak "
        "SNR 10 and 30, respectively, versus 18.3% and 75.8% in high-density positions. Because this pilot uses only "
        "40 injections per available chip and density-SNR cell, a central crop, and a supplied top-1000 catalogue "
        "that is not certified exhaustive, it is a proposal stress test rather than a full AstroCFR blind-purity result.")

    # Remove the misleading precision column from the main simulation table.
    t2 = doc.tables[1]
    drop_table_column(t2, 3)
    table2_caption = rewrite(find(doc, "Table 2."),
        "Table 2. Registered CSST-like measurement audit. The oracle-conditioned assembly value is excluded; "
        "closed-book proposal recovery is reported separately in Supplementary Section S9.", indent=0, style="Caption")
    bookmark(table2_caption, "table_2", 1)

    # Replace the former recall/precision plot with recall and physical residuals only.
    image_p = next(p for p in doc.paragraphs if p._p.xpath('.//w:drawing') and
                   p._p.getprevious() is not None and
                   "Table 2" in find(doc, "Table 2.").text)
    # The first drawing after Table 2 is the old Figure 2.
    started = False
    for candidate in doc.paragraphs:
        if candidate is table2_caption:
            started = True
        elif started and candidate._p.xpath('.//w:drawing'):
            image_p = candidate
            break
    image_p.clear()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.add_run().add_picture(str(FIGURE_DEST), width=Inches(6.55))
    figure2_caption = rewrite(find(doc, "Fig. 2."),
        "Fig. 2. Registered CSST-like reference recovery and conditional measurement residuals. Precision is not "
        "plotted because AstroCFR challenge assembly is reference-aware; the closed-book proposal pilot is reported "
        "in Supplementary Section S9.", indent=0, style="Caption")
    bookmark(figure2_caption, "figure_2", 2)

    # Limitations: use unambiguous oracle and multi-exposure language.
    rewrite(find(doc, "Another limitation is that the final competition configuration"),
        "The final competition configuration uses reference-aware catalogue assembly. Its reported 100% "
        "simulation-catalogue precision is therefore oracle-conditioned and is not a blind deployment metric. The "
        "released image-derived quality flags support catalogue screening but do not estimate blind purity. A valid "
        "false-discovery claim requires exhaustive scene truth or an independently deeper reference catalogue.")
    flc = find(doc, "The primary development experiments use CSST-like simulations")
    after(flc,
        "A multi-exposure FLC comparison is not reported because the available sequence did not satisfy the "
        "predeclared sub-pixel registration and PSF-quality acceptance gate. Failed alignment diagnostics are "
        "recorded in Supplementary Section S9 for reproducibility but are not entered into any DOLPHOT/HST1PASS "
        "performance table.")
    rewrite(find(doc, "The next evaluation should use multiple HST/ACS FLC exposures"),
        "The next evaluation should use a homogeneous HST/ACS FLC sequence that first passes sub-pixel "
        "inter-exposure registration. Source coordinates would then be shared across exposures with exposure-specific "
        "spatial ePSFs, backgrounds, and Poisson/read-noise weights. The same physical backend must be compared with "
        "and without AstroCFR candidate and neighbour-group priors under identical exposure lists, spatial hold-outs, "
        "and artificial-star scenes, with DOLPHOT or HST1PASS as an external reference. This is the required route to "
        "a multi-exposure claim, not a result of the present manuscript.")

    # Keep repository version consistent with the current package metadata.
    rewrite(find(doc, "Code availability:"),
        "Code availability: The manuscript-matched source is prepared for https://github.com/zxl1999/AstroCFR. "
        "The release contains reusable modules in src/wpdc, controlled HST and CSST experiment scripts, "
        "machine-readable summaries, environment locks, data provenance, and manuscript builders. The prepared "
        "package version is 1.6.2; the public commit hash, tag, and release archive must be recorded after upload. "
        "No archival DOI is claimed before a public release exists.", indent=0)

    # Supplement S9: enough detail for audit, deliberately not a positive multi-exposure result.
    append(sup, "S9 Closed-book proposal pilot and multi-exposure readiness audit", indent=0, style="Heading 1")
    append(sup,
        "For the closed-book CSST-like pilot, the test-chip candidate ECSV was written before opening that chip's "
        "supplied top-1000 catalogue. Candidate generation used an image-only proposal branch; no injected position "
        "was provided to detection or recovery. Processing was restricted to a fixed central 3000 x 3000-pixel crop "
        "to control memory. These controls remove direct test-truth leakage from proposal generation, but the "
        "top-1000 catalogue is not known to be exhaustive, so catalogue-match fraction is not blind purity.")
    append(sup,
        "Aggregated over the available chips, recovery is 102/160 at peak SNR 10 and 143/160 at peak SNR 30 in "
        "low-density positions; the corresponding high-density counts are 22/120 and 91/120. The different "
        "denominators arise because chip 12 did not contain valid high-density injection locations in the fixed crop. "
        "The pilot therefore establishes the expected density- and SNR-dependent proposal failure, but its scale and "
        "image-only frontend do not support a full AstroCFR efficacy, purity, or survey-throughput claim.")
    append(sup, "Table S20. Closed-book image-only proposal recovery. Wilson 95% intervals quantify binomial "
                "recovery only; injected positions were withheld from processing.", indent=0, style="Caption")
    table = sup.add_table(rows=1, cols=6)
    headers = ["Density", "Peak SNR", "Injected", "Recovered", "Recovery", "Wilson 95% CI"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, size=8.2)
    rows = [
        ("Low", "10", "160", "102", "0.638", "[0.561, 0.708]"),
        ("Low", "30", "160", "143", "0.894", "[0.836, 0.933]"),
        ("High", "10", "120", "22", "0.183", "[0.124, 0.262]"),
        ("High", "30", "120", "91", "0.758", "[0.674, 0.826]"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell(cell, value, size=8.2)
    three_line_table(table)
    append(sup,
        "A stricter leave-one-chip-out RF screen was also attempted on the cropped data, but retention was unstable "
        "and collapsed to zero on some withheld-chip settings. It is recorded as a negative control in the machine-"
        "readable audit and is not used as evidence of cross-chip generalisation.")
    append(sup,
        "For the multi-exposure readiness audit, four homogeneous deep ACS/WFC F606W FLC exposures of NGC 6752 "
        "were selected. WCS-derived detector-shift initialisation recovered thousands of alignment candidates, but "
        "the best tested transformation retained 0.817-pixel inter-exposure scatter, which failed the predeclared "
        "sub-pixel gate. Processing was stopped and no NGC 1851 comparison was launched under the unresolved "
        "condition. These values diagnose dataset/registration readiness only; they are not DOLPHOT photometric or "
        "astrometric performance measurements.")
    append(sup,
        "Accordingly, no failed DOLPHOT output is included in the comparative tables. A valid experiment must first "
        "pass registration and PSF-star acceptance, then run identical FLC lists, spatial test regions, and "
        "artificial-star scenes for the physical backend with and without AstroCFR candidate/group priors.")

    doc.save(MAIN_DEST)
    sup.save(SUP_DEST)
    print(MAIN_DEST)
    print(SUP_DEST)


if __name__ == "__main__":
    main()
