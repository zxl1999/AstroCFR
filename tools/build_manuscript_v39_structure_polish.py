#!/usr/bin/env python
"""Merge the residual one-sentence method subsection and restore a direct citation link."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import internal_link, set_font

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v38_final.docx"
DEST = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v39_final.docx"
SUP_SOURCE = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v38.docx"
SUP_DEST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v39.docx"


def find(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def remove(p):
    p._element.getparent().remove(p._element)


def new_after(p, value, bold=False, indent=.35):
    q = p._parent.add_paragraph(style="Normal")
    q._p.getparent().remove(q._p); p._p.addnext(q._p)
    q.paragraph_format.first_line_indent = Inches(indent) if indent else None
    q.paragraph_format.line_spacing = 1.08
    set_font(q.add_run(value), size=10.5, bold=bold)
    return q


def main():
    doc = Document(SOURCE)
    sup = Document(SUP_SOURCE)

    # Restore a real 4.3 heading, fold routing into it, and remove the singleton 4.4.
    classifier_end = find(doc, "RandomForest, XGBoost, and hybrid stacking classifiers are re-inserted")
    h = new_after(classifier_end, "4.3 Full-pipeline integration and calibrated baselines", bold=True, indent=0)
    s_extractor = find(doc, "A calibrated SExtractor baseline is also run")
    add = new_after(s_extractor,
        "An additional density-adaptive policy is evaluated in the same framework; its definition and "
        "controlled operating-point results are reported in Section 5.4.1.")
    remove(find(doc, "4.4 Density-adaptive branch-routing protocol"))
    remove(find(doc, "The density-adaptive policy is evaluated as a fixed, pre-registered routing rule"))

    # Replace the literal author-year mention with a visible, clickable reference identifier.
    p = find(doc, "The primary reference set requires a valid F606W measurement")
    p.clear(); p.style = "Normal"
    p.paragraph_format.first_line_indent = Inches(.35); p.paragraph_format.line_spacing = 1.08
    before = (
        "The primary reference set requires a valid F606W measurement, reported magnitude error below "
        "0.10 mag, qfitV below 0.30, neighbour-light fraction below 1, and at least one F606W measurement. "
        "The same three-way 200-pixel spatial partition is retained: partitions 0 and 1 support target "
        "fitting and threshold selection, and partition 2 is untouched until final evaluation. Positions are "
        "associated within 2 pixels. Astrometric precision is evaluated after fitting the six-parameter affine "
        "transformation recommended by "
    )
    set_font(p.add_run(before), size=10.5)
    internal_link(p, "Anderson et al. (2008; Ref. 2)", "ref_2")
    set_font(p.add_run(" on non-test matches and applying it to the test partition only."), size=10.5)

    doc.save(DEST)
    sup.save(SUP_DEST)
    print(DEST)
    print(SUP_DEST)


if __name__ == "__main__":
    main()
