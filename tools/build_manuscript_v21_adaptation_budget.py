#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Append the target-adaptation budget experiment to the v20 manuscript."""
from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v20_introduction_hyperlinks.docx"
DEST = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v21_adaptation_budget.docx"
CSV = ROOT / "results" / "hst_target_adaptation_budget" / "adaptation_budget_summary.csv"
FIGURE = ROOT / "results" / "hst_target_adaptation_budget" / "adaptation_budget_curve.png"


def set_run(run, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold


def text_paragraph(doc, text, first=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = None if first else Inches(.35)
    p.paragraph_format.line_spacing = 1.08
    set_run(p.add_run(text))
    return p


def caption(doc, text, bookmark_name, bookmark_id):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    set_run(p.add_run(text))
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(bookmark_id)); start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(bookmark_id))
    p._p.insert(0, start); p._p.append(end)
    return p


def internal_link(p, label, anchor):
    h = OxmlElement("w:hyperlink"); h.set(qn("w:anchor"), anchor); h.set(qn("w:history"), "1")
    r = OxmlElement("w:r"); rp = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    rp.append(color); rp.append(underline); r.append(rp)
    t = OxmlElement("w:t"); t.text = label; r.append(t); h.append(r); p._p.append(h)


def linked_sentence(doc, before, label, anchor, after):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Inches(.35)
    p.paragraph_format.line_spacing = 1.08
    set_run(p.add_run(before)); internal_link(p, label, anchor); set_run(p.add_run(after))
    return p


def three_line(table):
    borders = OxmlElement("w:tblBorders")
    for edge, value, size in (("top", "single", "12"), ("bottom", "single", "12"), ("insideH", "single", "6")):
        node = OxmlElement(f"w:{edge}"); node.set(qn("w:val"), value); node.set(qn("w:sz"), size); node.set(qn("w:color"), "000000")
        borders.append(node)
    table._tbl.tblPr.append(borders)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs: set_run(run, bold=(row == table.rows[0]))


def move_before(element, marker):
    marker._p.addprevious(element._p if hasattr(element, "_p") else element)


def main():
    doc = Document(SOURCE)
    rows = list(csv.DictReader(CSV.open(encoding="utf-8")))
    marker = next(p for p in doc.paragraphs if p.text.strip() == "6 Discussion")
    created = []
    heading = doc.add_paragraph("5.17 Target-adaptation budget learning curve", style="Heading 1"); created.append(heading)
    created.append(text_paragraph(doc, "WPDC demonstrates a two-stage deployment strategy. Stage 1 develops the candidate screen on CSST-like simulations. Stage 2 uses the target image to estimate the PSF and normalize candidate features without catalogue labels, then uses a small, spatially disjoint labelled target region to refit the candidate classifier and select its operating threshold. The production run is performed only after this lightweight calibration. This separates simulation development from the bounded target-specific action that is required by the observed sensor and crowding shift.", first=True))
    created.append(text_paragraph(doc, "We quantified the Stage-2 budget on three public HST/ACS F606W fields (NGC 6397, NGC 6752, and NGC 1851). Each 1200 x 1200 crop was divided into 200 x 200 pixel tiles. Sixty percent of the tiles formed the target-fit pool, 20% were reserved for threshold validation, and 20% were untouched for final testing. A budget of 1, 3, or 6 tiles corresponds to 2.8%, 8.3%, or 16.7% of one crop. The non-zero budgets use five deterministic tile selections; their 95% intervals describe selection variability. These tiles are within-image calibration units, not independent telescope visits or independent images."))
    created.append(linked_sentence(doc, "The corresponding quantitative results are summarized in ", "Table 21", "table_21", "."))
    cap_table = caption(doc, "Table 21. Simulation-to-HST target-adaptation budget curve. Recall and catalogue match-rate lower bound are measured on spatially untouched tiles. Intervals are 95% selection intervals across five calibration-tile draws, not population-level confidence intervals.", "table_21", 78); created.append(cap_table)
    table = doc.add_table(rows=1, cols=7); created.append(table._tbl)
    headers = ["Cluster", "Tiles", "Field fraction", "Positive / negative labels", "Test recall (95% interval)", "V<=20 recall (95% interval)", "Match-rate lower bound (95% interval)"]
    for cell, value in zip(table.rows[0].cells, headers): cell.text = value
    for row in rows:
        cells = table.add_row().cells
        values = [row["cluster"].upper(), row["budget_tiles"], f"{float(row['field_fraction']):.3f}",
                  f"{float(row['n_train_positive_mean']):.0f} / {float(row['n_train_negative_mean']):.0f}",
                  f"{float(row['test_recall_mean']):.3f} +/- {float(row['test_recall_ci95']):.3f}",
                  f"{float(row['test_bright_recall_v_le_20_mean']):.3f} +/- {float(row['test_bright_recall_v_le_20_ci95']):.3f}",
                  f"{float(row['test_match_rate_lower_bound_mean']):.3f} +/- {float(row['test_match_rate_lower_bound_ci95']):.3f}"]
        for cell, value in zip(cells, values): cell.text = value
    three_line(table)
    created.append(linked_sentence(doc, "The corresponding visualization is shown in ", "Fig. 22", "figure_22", "."))
    fig_p = doc.add_paragraph(); fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER; fig_p.add_run().add_picture(str(FIGURE), width=Inches(6.2)); created.append(fig_p)
    cap_figure = caption(doc, "Fig. 22. Learning curves for target-adaptation calibration budget. A single tile is sufficient for a large recovery gain in NGC 6397 and NGC 6752; the difficult NGC 1851 field requires several tiles containing enough conservative negative labels for a stable update.", "figure_22", 79); created.append(cap_figure)
    created.append(text_paragraph(doc, "The effect is substantial but not universal. On NGC 6397, one calibration tile containing on average 38 positive and 8 conservative negative examples increases held-out all-quality recall from 0.183 to 0.828 +/- 0.036 and V<=20 recall from 0.231 to 0.915 +/- 0.021. On NGC 6752, one tile with approximately 101 positive and 15 negative examples increases all-quality recall from 0.057 to 0.648 +/- 0.021 and V<=20 recall from 0.079 to 0.782 +/- 0.006. The difficult NGC 1851 case provides the deployment boundary: one and three tiles are unstable because negative labels are scarce, whereas six tiles (about 364 positive and 19 negative examples) produce a stable 0.166 +/- 0.002 all-quality recall, close to the 0.173 proposal-stage ceiling of this single image."))
    created.append(text_paragraph(doc, "Thus the defensible deployment statement is not that a fixed number of real images always suffices. Rather, a simulation-trained WPDC front end can be made operational with a small, labelled and spatially disjoint target calibration region when that region contains adequate positive and negative candidate examples. For the two better-sampled fields, one 2.8%-area tile is already effective; the third field shows why a label-count and stability check is required before releasing an adapted catalogue."))
    for element in created:
        move_before(element, marker)
    doc.save(DEST)
    print(DEST)


if __name__ == "__main__":
    main()
