#!/usr/bin/env python
"""Add a separate, accurately scoped CNN CPU/GPU resource profile to v25."""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v25_runtime_resources.docx"
DEST = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v26_cnn_gpu_resources.docx"

def font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size); run.font.bold = bold; run.font.color.rgb = RGBColor(0, 0, 0)

def rewrite(p, text, indent=0.35):
    p.clear(); p.style = "Normal"; p.paragraph_format.first_line_indent = Inches(indent) if indent else None; p.paragraph_format.line_spacing = 1.08
    font(p.add_run(text))

def bookmark(p, name, ident):
    a = OxmlElement("w:bookmarkStart"); a.set(qn("w:id"), str(ident)); a.set(qn("w:name"), name)
    b = OxmlElement("w:bookmarkEnd"); b.set(qn("w:id"), str(ident)); p._p.insert(0, a); p._p.append(b)

def link(p, label, anchor):
    h = OxmlElement("w:hyperlink"); h.set(qn("w:anchor"), anchor); h.set(qn("w:history"), "1")
    r = OxmlElement("w:r"); rp = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
    rp.append(c); rp.append(u); r.append(rp); t = OxmlElement("w:t"); t.text = label; r.append(t); h.append(r); p._p.append(h)

def cell_text(cell, text, bold=False):
    cell.text = ""; cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    font(p.add_run(str(text)), size=8.4, bold=bold)

def border(cell, edge, val, sz):
    pr = cell._tc.get_or_add_tcPr(); borders = pr.first_child_found_in("w:tcBorders")
    if borders is None: borders = OxmlElement("w:tcBorders"); pr.append(borders)
    item = borders.find(qn("w:" + edge))
    if item is None: item = OxmlElement("w:" + edge); borders.append(item)
    item.set(qn("w:val"), val); item.set(qn("w:sz"), sz); item.set(qn("w:color"), "000000")

def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            border(cell, "top", "single", "14" if i == 0 else "0")
            border(cell, "bottom", "single", "14" if i == len(table.rows)-1 else "0")
            border(cell, "left", "nil", "0"); border(cell, "right", "nil", "0")

def main():
    doc = Document(SOURCE)
    # Correct the scope of the CPU-only claim in v25.
    heading = next(p for p in doc.paragraphs if p.text.strip() == "5.14 Measurement uncertainty and runtime repeatability")
    runtime = next(p for p in doc.paragraphs if p._p is heading._p.getnext())
    rewrite(runtime, "To avoid treating one point estimate as a deployment guarantee, we added two uncertainty analyses. For position and magnitude RMS, the held-out matches, affine registration, robust residual mask, and photometric zero point are fixed, and the retained test residuals are resampled 1,000 times. These are conditional residual-bootstrap intervals: they quantify measurement scatter conditional on the detected catalogue, not the uncertainty of the whole detection pipeline. Runtime and RSS are measured with five repeated warm runs after one warm-up run, using the same 1200 × 1200 crop and the same method-stage timing wrapper. The controlled HST/ACS baselines in Tables 15–18 and Table 22 used one Python process on an Intel Core Ultra 9 275HX host (24 logical cores, 50.9 GB RAM), Python 3.12.7, NumPy 2.5.1, SciPy 1.18.0, Astropy 7.2.0, Photutils 3.0.0, and scikit-learn 1.9.0. Those DAO, SEP, Photutils, RF, and ePSF branches were executed on CPU only; GPU memory is therefore not applicable to those rows. The CNN classifier is profiled separately in Table 23 on the available GPU. Runtime is reported in seconds per million processed pixels and RSS as the peak process-memory increase relative to the warm baseline.", indent=0)
    caption22 = next(p for p in doc.paragraphs if p.text.startswith("Table 22."))
    caption22.text = "Table 22. Accuracy–cost operating points on the NGC 6752 controlled HST/ACS test. Runtime and RSS are CPU measurements for the methods in this table; GPU memory is not applicable to these rows."
    caption22.style = "Caption"
    for r in caption22.runs: font(r)

    discussion = next(p for p in doc.paragraphs if p.text.strip() == "6 Discussion")
    h = doc.add_paragraph("5.19 CNN accelerator resource profile", style="Heading 1")
    p = doc.add_paragraph(style="Normal")
    rewrite(p, "The CNN classifier is a separate simulation-domain operating point and should not be forced into the CPU-only HST baseline ranking. We therefore measured the original unmodified StarBogusNet architecture (25 × 25 normalized cutout plus 17 handcrafted features; 28,081 trainable parameters) on deterministic representative tensors with the source-code defaults of 30 epochs, training batch size 64, and inference batch size 256. The measurement includes model-stage tensor transfer and inference, but excludes candidate generation, cutout construction, disk I/O, and catalogue calibration. It is consequently a reproducible accelerator-resource profile, not an end-to-end survey throughput or accuracy comparison.")
    ref = doc.add_paragraph(style="Normal"); ref.paragraph_format.first_line_indent = Inches(.35); ref.paragraph_format.line_spacing = 1.08
    font(ref.add_run("The corresponding quantitative results are summarized in ")); link(ref, "Table 23", "table_23"); font(ref.add_run("."))
    cap = doc.add_paragraph("Table 23. Original WPDC StarBogusNet resource profile. The CPU and GPU rows use identical input geometry, sample count, seeds, epochs, and batch sizes. GPU: NVIDIA GeForce RTX 5060 Laptop GPU; PyTorch 2.10.0.dev+cu128; CUDA 12.8.", style="Caption")
    bookmark(cap, "table_23", 2023)
    table = doc.add_table(rows=1, cols=8)
    headers = ["Device", "Training / s", "Training / epoch s", "Inference / ms per 1k", "Process RSS / MB", "GPU allocated / MB", "GPU reserved / MB", "Scope"]
    for c, v in zip(table.rows[0].cells, headers): cell_text(c, v, True)
    rows = [
        ("CPU", "8.738", "0.291", "45.285", "147.7", "N/A", "N/A", "model-stage"),
        ("RTX 5060 GPU", "3.674", "0.122", "8.984", "655.0", "58.3", "86.0", "model-stage"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for c, v in zip(cells, row): cell_text(c, v)
    style_table(table)
    cursor = discussion._p.getprevious(); nodes = [h._p, p._p, ref._p, cap._p, table._element]
    for node in nodes: node.getparent().remove(node)
    for node in nodes: cursor.addnext(node); cursor = node
    doc.save(DEST); print(DEST)

if __name__ == "__main__": main()
