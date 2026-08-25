#!/usr/bin/env python
"""Build v44 with independent Anderson-PSF injection evidence.

The script leaves v43 untouched.  It adds a compact primary-result paragraph
to the main text and moves the full protocol/table/renderer audit to S10.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import set_cell, set_font
from build_manuscript_v42_closed_book_scope import three_line_table


ROOT = Path(__file__).resolve().parents[1]
MAIN_SOURCE = ROOT / "supplementary/AstroCFR_Crowded_Field_Manuscript_v43_roi_abstract.docx"
MAIN_DEST = ROOT / "supplementary/AstroCFR_Crowded_Field_Manuscript_v44_independent_psf.docx"
SUP_SOURCE = ROOT / "supplementary/AstroCFR_Supplementary_Materials_v43.docx"
SUP_DEST = ROOT / "supplementary/AstroCFR_Supplementary_Materials_v44.docx"
RESULTS = ROOT / "results/non_globular_runs/m31_b21_f15/matched_coordinate_scene/independent_psf_validation"
AUDIT = ROOT / "results/non_globular_runs/m31_b21_f15/matched_coordinate_scene/anderson_psf_audit/anderson_drz_psf_audit.json"
FIGURE = RESULTS / "independent_anderson_psf_validation.png"


def find(doc: Document, prefix: str):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def rewrite(paragraph, text: str, *, style: str = "Normal", indent: float = 0.35):
    paragraph.clear(); paragraph.style = style
    paragraph.paragraph_format.first_line_indent = Inches(indent) if indent else None
    paragraph.paragraph_format.line_spacing = 1.08
    set_font(paragraph.add_run(text), size=10.5, bold=style.startswith("Heading"))
    return paragraph


def after(paragraph, text: str, *, style: str = "Normal", indent: float = 0.35):
    new = paragraph._parent.add_paragraph(style=style)
    new._p.getparent().remove(new._p); paragraph._p.addnext(new._p)
    return rewrite(new, text, style=style, indent=indent)


def append(doc: Document, text: str, *, style: str = "Normal", indent: float = 0.35):
    return rewrite(doc.add_paragraph(style=style), text, style=style, indent=indent)


def replace_text(paragraph, old: str, new: str):
    if old not in paragraph.text:
        raise ValueError(f"replacement source not found: {old}")
    rewrite(paragraph, paragraph.text.replace(old, new), style=paragraph.style.name,
            indent=0 if paragraph.style.name != "Normal" else 0.35)


def bookmark(paragraph, name: str, ident: int):
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(ident)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(ident))
    paragraph._p.insert(0, start); paragraph._p.append(end)


def move_after(element, paragraph):
    element.getparent().remove(element); paragraph._p.addnext(element)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, size=7.5)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell(cell, value, size=7.5)
    three_line_table(table)
    return table


def main() -> None:
    result = json.loads((RESULTS / "independent_psf_validation.json").read_text(encoding="utf-8"))
    audit = json.loads(AUDIT.read_text(encoding="utf-8"))
    independent = [r for r in result["recovery"] if r["injection_psf"] == "anderson"]
    tests = [r for r in result["between_method_paired_tests"] if r["injection_psf"] == "anderson"]

    main = Document(MAIN_SOURCE)
    paragraph = find(main, "The initial artificial-star experiment used approximately 40 injections")
    replace_text(
        paragraph,
        "The experiment therefore measures proposal recovery under a controlled injected scene, not catalogue purity.",
        "The experiment therefore measures proposal recovery under a controlled injected scene, not catalogue purity. Because these NGC 6752 stars were rendered with an image-derived empirical PSF from the same modelling family used by the AstroCFR ePSF branch, the result is treated as model-matched controlled evidence rather than an independent PSF validation.",
    )
    inserted = after(
        find(main, "Fig. 7. Expanded NGC 6752 artificial-star recovery."),
        "We therefore ran an independent-PSF stress test on the public PHAT M31 B21-F15 F475W DRZ image. The official Anderson ACS/WFC F475W standard PSF, rather than AstroCFR's image-derived ePSF, was evaluated at each source location through the WCS solutions of three registered 370-s FLC exposures, broadened with a fixed 0.55-pixel Gaussian output kernel to match the measured 2.2-pixel DRZ core, and normalized on the DRZ grid. Each star was injected separately, a position entered the denominator only when all four methods had no pre-existing detection within 2 pixels, and recovery required a new detection within 2 pixels. On the common denominator, AstroCFR ePSF recovered 39/41 (95.1%, 95% CI 83.9-98.7%) low-density and 35/42 (83.3%, 69.4-91.7%) high-density F475W=26.5 stars, versus 28/41 (68.3%) and 24/42 (57.1%) for Photutils. The paired differences were +26.8 percentage points (95% bootstrap CI 14.6-41.5; exact McNemar p=0.00098) and +26.2 points (14.3-40.5; p=0.00098), respectively. This confirms that the faint-star proposal gain is not explained solely by reusing the recovery ePSF as the injector. The renderer approximates the local dithered DRZ PSF but does not reproduce the full AstroDrizzle kernel or correlated noise, so the result supports candidate recovery rather than final photometric superiority; full tables and audits are in Supplementary Section S10.",
    )
    bookmark(inserted, "sec_independent_psf_result", 9901)
    limitation = find(main, "The primary development experiments use CSST-like simulations")
    after(
        limitation,
        "The primary NGC 6752 fixed-scene artificial-star curves use an image-derived empirical injection PSF and are therefore model-matched controlled evidence. The independent PHAT M31 test reduces this circularity with the official Anderson F475W library and retains a statistically supported faint-star recovery advantage after a fixed output-grid broadening matches the measured DRZ core scale. It remains a single-DRZ approximation: WCS projection through three FLC exposures does not reproduce the complete AstroDrizzle kernel or its correlated-noise statistics, and the test does not establish blind catalogue purity or final photometric dominance.",
    )
    conclusion = find(main, "The Pareto frontier exposes three practically useful operating points.")
    replace_text(
        conclusion,
        "these recovery branches recover more reference and injected sources than the evaluated DAO, SEP, and RF branches;",
        "these recovery branches recover more reference and injected sources than the evaluated DAO, SEP, and RF branches; an independent Anderson-PSF injection on PHAT M31 preserves the faint-star proposal advantage over DAO/Photutils;",
    )
    main.save(MAIN_DEST)

    sup = Document(SUP_SOURCE)
    heading = append(sup, "S10 Independent Anderson-PSF injection validation", style="Heading 1", indent=0)
    bookmark(heading, "supp_s10_anderson", 9902)
    append(
        sup,
        "The earlier NGC 6752 artificial-star experiment is a controlled, method-identical scene test, but its artificial sources are generated from an image-derived empirical PSF belonging to the same modelling family as AstroCFR's recovery ePSF. To test whether this model match accounts for the recovery gain, we used the official Anderson ACS/WFC F475W standard PSF file (9 x 10 detector grid, 101 x 101 samples per PSF, four-times oversampled; 3,677,760 bytes; SHA-256 9B3A98844020581FFDF1EEEBB5D4488F03011BC4563C578D19042A76A81B5C82). The injector does not use AstroCFR's recovered PSF.",
    )
    append(
        sup,
        "For each PHAT DRZ output-pixel centre, celestial WCS maps the pixel and requested star location back to each of three registered F475W FLC exposures (jbex18u6q, jbex18u9q, and jbex18ucq). The corresponding spatial Anderson PSF is evaluated in the appropriate ACS chip coordinates; the three exposure contributions are averaged, convolved with a fixed 0.55-pixel Gaussian output kernel chosen before the final run to match the 2.2-pixel DRZ core, and normalized in a 25 x 25 DRZ stamp. A registered geometry audit required unit flux, no negative/non-finite values, and less than 0.20-pixel radial centroid offset. All four extension-density test positions passed; the maximum offset was %.3f pixel. This calculation represents local geometry and dithering but not the full AstroDrizzle kernel or correlated-noise process." % audit["maximum_centroid_radial_offset_px"],
    )
    append(
        sup,
        "The deterministic input contains 25 positions per logical extension x density x magnitude stratum (200 attempts). Every position is injected and processed separately. Within an injection model, the paired denominator is the intersection of positions for which DAO, SEP, Photutils, and AstroCFR all have no pre-existing detection within 2 pixels; 174 positions meet this criterion. Recovery is a new detection within 2 pixels. Individual intervals are Wilson 95% intervals. Recovery-difference intervals are paired 20,000-replicate bootstraps with seed 20260812, and p values are exact two-sided McNemar tests.",
    )
    caption = append(
        sup,
        "Table S21. Strict single-star recovery under independent official Anderson F475W PSF injection. RMS is the radial distance to the nearest recovered detection, not a calibrated multi-exposure astrometric residual. Times are medians per local baseline-plus-injected trial.",
        style="Caption", indent=0,
    )
    rows = []
    labels = {"dao": "DAO", "sep": "SEP", "photutils": "Photutils", "astrocfr_epsf": "AstroCFR ePSF"}
    for r in independent:
        rows.append([
            r["density_band"], f"{r['input_vegamag_f475w']:.1f}", labels[r["method"]],
            f"{r['recovered']}/{r['common_eligible']}",
            f"{100*r['recovery']:.1f} [{100*r['ci95_low_wilson']:.1f}, {100*r['ci95_high_wilson']:.1f}]",
            f"{r['nearest_detection_radial_rms_px']:.3f}", f"{r['runtime_per_trial_median_s']:.3f}",
        ])
    table = add_table(sup, ["Density", "F475W", "Method", "Recovered", "Recovery / % [95% CI]", "Nearest RMS / px", "Time / s"], rows)
    move_after(table._tbl, caption)
    bookmark(caption, "table_s21", 9903)
    main_result = next(r for r in tests if r["density_band"] == "high" and r["input_vegamag_f475w"] == 26.5 and r["method_b"] == "photutils")
    append(
        sup,
        "For F475W=26.5, AstroCFR exceeds DAO/Photutils by 26.8 percentage points in the low-density stratum (paired-bootstrap 95%% CI 14.6-41.5; 11 AstroCFR-only recoveries, 0 baseline-only; exact McNemar p=0.00098) and by %.1f points in the high-density stratum (%.1f-%.1f; %d versus %d; p=%.4f). Relative to the empirical-PSF injection on the same 174 positions, AstroCFR recovery changes by +4.9 points at low density and 0.0 points at high density; neither injection-model difference is significant by exact McNemar testing (p=0.50 and 1.00)." % (100*main_result["recovery_difference_a_minus_b"], 100*main_result["paired_bootstrap_ci95_low"], 100*main_result["paired_bootstrap_ci95_high"], main_result["a_only"], main_result["b_only"], main_result["mcnemar_exact_p"]),
    )
    image_p = append(sup, "", indent=0)
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.add_run().add_picture(str(FIGURE), width=Inches(6.25))
    fig_caption = append(
        sup,
        "Fig. S17. Independent Anderson-PSF validation. (a) Strict paired recovery with Wilson 95% intervals. (b) AstroCFR's F475W=26.5 recovery change when the injector is changed from the image-derived empirical PSF to the official Anderson library; error bars are paired-bootstrap 95% intervals.",
        style="Caption", indent=0,
    )
    bookmark(fig_caption, "figure_s17", 9904)
    append(
        sup,
        "Interpretation boundary. This experiment removes exact reuse of AstroCFR's image-derived ePSF as the artificial-star generator, but it remains a single-stack candidate-recovery experiment. It does not establish blind catalogue purity, a pixel-identical comparison with DOLPHOT native FakeStars, a complete AstroDrizzle simulation, or superior final astrometry/photometry. DOLPHOT's three-FLC FakeStars measurements remain a separate physical-backend image domain and are not pooled into this leaderboard.",
    )
    sup.save(SUP_DEST)
    print(MAIN_DEST); print(SUP_DEST)


if __name__ == "__main__":
    main()
