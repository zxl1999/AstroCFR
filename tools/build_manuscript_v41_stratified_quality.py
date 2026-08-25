#!/usr/bin/env python
"""Add the compact stratified and blind-quality release audit to v40."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import set_font

ROOT = Path(__file__).resolve().parents[1]
MAIN_SOURCE = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v40_methodology.docx"
MAIN_DEST = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v41_stratified_quality.docx"
SUP_SOURCE = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v40.docx"
SUP_DEST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v41.docx"
FIGURES = ROOT / "results" / "hst_stratified_quality"


def find(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def after(paragraph, text, heading=False, indent=.35):
    new = paragraph._parent.add_paragraph(style="Heading 1" if heading else "Normal")
    new._p.getparent().remove(new._p)
    paragraph._p.addnext(new._p)
    new.paragraph_format.first_line_indent = Inches(indent) if indent else None
    new.paragraph_format.line_spacing = 1.08
    set_font(new.add_run(text), size=10.5, bold=heading)
    return new


def append(doc, text, style="Normal", indent=.35):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Inches(indent) if indent else None
    p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(text), size=10.5, bold=style.startswith("Heading"))
    return p


def main():
    main_doc = Document(MAIN_SOURCE)
    sup = Document(SUP_SOURCE)
    p = find(main_doc, "To avoid treating one point estimate as a deployment guarantee")
    after(p,
          "To expose where the operating points differ, Supplementary Section S8 releases common-protocol curves "
          "for both NGC 6752 and NGC 1851. They stratify held-out completeness by reference magnitude and local "
          "density, catalogue-match lower bounds by detection SNR, and conditional positional and magnitude RMS. "
          "The associated AstroCFR spatial-ePSF catalogues include reference-free SNR, residual-improvement, local-" 
          "crowding, deblend, bright-core, classifier-probability, and PSF-fit flags. These fields make blind catalogue "
          "screening auditable, but do not convert the HST catalogue-match lower bound into blind purity.")
    main_doc.save(MAIN_DEST)

    append(sup, "S8 Common-protocol stratification and blind-quality release", style="Heading 1", indent=0)
    append(sup,
           "The six controlled branches were rerun on the same 1200 × 1200 HST/ACS test crops for NGC 6752 and "
           "NGC 1851. The released CSV/JSON separates three quantities: reference-conditioned held-out completeness "
           "by magnitude and local density; conditional position and magnitude RMS on matched held-out references; "
           "and the detection-SNR catalogue-match lower bound. The latter is not blind purity. Completeness and "
           "lower-bound intervals are Wilson 95% intervals. Wall times in the figure are single-run relative values; "
           "the five-repeat runtime intervals in the main text remain the canonical resource measurements.")
    append(sup,
           "For operational use, the spatial-ePSF catalogue exports only image/candidate-derived metadata: SNR, local "
           "residual improvement, neighbour count within 10 pixels, residual-deblend status, bright-core proximity, "
           "available classifier probability, local PSF-fit quality, and a uint16 quality bitmask. Bit 1 denotes "
           "SNR < 5; bit 2, local PSF residual > 3 RMS units; bit 4, at least three neighbours; bit 8, bright-core "
           "proximity; bit 16, residual deblending; and bit 32, a classifier probability in (0.2, 0.8). Local PSF-fit "
           "quality is a screening diagnostic, not an absolute goodness-of-fit in unresolved blends.")
    for cluster in ("ngc6752", "ngc1851"):
        p = sup.add_paragraph(style="Normal"); p.paragraph_format.first_line_indent = Inches(0)
        p.alignment = 1
        p.add_run().add_picture(str(FIGURES / f"{cluster}_stratified_recovery_precision.png"), width=Inches(6.25))
        append(sup,
               f"Fig. S{15 if cluster == 'ngc6752' else 16}. {cluster.upper()} common-protocol stratification. "
               "Recovery curves use held-out quality references; position and magnitude RMS are conditional on "
               "matched held-out references; the SNR panel is a catalogue-match lower bound rather than blind purity.",
               style="Caption", indent=0)
    sup.save(SUP_DEST)
    print(MAIN_DEST)
    print(SUP_DEST)


if __name__ == "__main__":
    main()
