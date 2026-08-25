#!/usr/bin/env python
"""Embed failure-case and stratified-recovery diagnostics in a new revision."""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v26_cnn_gpu_resources.docx"
DEST = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v27_diagnostics.docx"
FAIL = ROOT / "results" / "hst_failure_cases" / "fig_failure_cases.png"
STRAT = ROOT / "results" / "hst_failure_cases" / "fig_density_magnitude_recovery.png"

def font(run, size=10.5):
    run.font.name = "Times New Roman"; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体"); run.font.size = Pt(size); run.font.color.rgb = RGBColor(0, 0, 0)

def bookmark(p, name, ident):
    a = OxmlElement("w:bookmarkStart"); a.set(qn("w:id"), str(ident)); a.set(qn("w:name"), name)
    b = OxmlElement("w:bookmarkEnd"); b.set(qn("w:id"), str(ident)); p._p.insert(0, a); p._p.append(b)

def link(p, label, anchor):
    h = OxmlElement("w:hyperlink"); h.set(qn("w:anchor"), anchor); h.set(qn("w:history"), "1")
    r = OxmlElement("w:r"); rp = OxmlElement("w:rPr"); c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rp.append(c); rp.append(u); r.append(rp); t = OxmlElement("w:t"); t.text = label; r.append(t); h.append(r); p._p.append(h)

def para(doc, text, style="Normal", indent=.35):
    p = doc.add_paragraph(style=style); p.paragraph_format.first_line_indent = Inches(indent) if indent else None; p.paragraph_format.line_spacing = 1.08; font(p.add_run(text)); return p

def main():
    if not FAIL.exists() or not STRAT.exists(): raise FileNotFoundError("Diagnostic figures must be rendered first")
    doc = Document(SOURCE)
    discussion = next(p for p in doc.paragraphs if p.text.strip() == "6 Discussion")
    h = doc.add_paragraph("5.20 Failure analysis and stratified recovery", style="Heading 1")
    p1 = para(doc, "The system-level metrics are complemented by diagnostics selected after the candidate run. Figure 23 shows three representative failure modes: a high-density NGC 6752 miss, a missed reference near a bright-star core, and a difficult held-out NGC 1851 region after target adaptation. Cyan circles denote quality-filtered reference stars, red crosses denote references without a candidate within 2 pixels, and yellow points denote WPDC candidates. These panels are explanatory examples, not additional tuned test results.")
    rp1 = para(doc, "The corresponding failure-case visualization is shown in ", indent=.35); link(rp1, "Fig. 23", "figure_23"); font(rp1.add_run("."))
    cap1 = para(doc, "Fig. 23. Representative WPDC failure cases: (a) high-crowding miss in NGC 6752, (b) bright-star artifact region in NGC 6752, and (c) difficult NGC 1851 held-out domain-adaptation region. All panels use the same public HST/ACS crop and 2-pixel association rule.", style="Caption", indent=0); bookmark(cap1, "figure_23", 2023)
    doc.add_picture(str(FAIL), width=Inches(6.4)); pic1 = doc.paragraphs[-1]
    p2 = para(doc, "The fixed artificial-star experiment also exposes the interaction between magnitude and local density. Figure 24 separates low-density scenes (0–1 quality references within 10 pixels) from high-density scenes (at least 3 neighbours) and reports recovery within 2 pixels at V=20 and V=22 with Wilson 95% intervals. The separation is largest for the recovery-oriented WPDC branches, which supports the bounded claim that their benefit is concentrated in crowded, relatively bright regimes rather than being uniform over all magnitudes and densities.")
    rp2 = para(doc, "The corresponding stratified-recovery visualization is shown in ", indent=.35); link(rp2, "Fig. 24", "figure_24"); font(rp2.add_run("."))
    cap2 = para(doc, "Fig. 24. Artificial-star recovery stratified by local density and injected V magnitude. Shaded regions are Wilson 95% intervals; identical fixed scenes are supplied to all methods.", style="Caption", indent=0); bookmark(cap2, "figure_24", 2024)
    doc.add_picture(str(STRAT), width=Inches(6.4)); pic2 = doc.paragraphs[-1]
    nodes = [h._p, p1._p, rp1._p, cap1._p, pic1._p, p2._p, rp2._p, cap2._p, pic2._p]
    cursor = discussion._p.getprevious()
    for node in nodes: node.getparent().remove(node)
    for node in nodes: cursor.addnext(node); cursor = node
    doc.save(DEST); print(DEST)

if __name__ == "__main__": main()
