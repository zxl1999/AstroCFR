#!/usr/bin/env python
"""Build v45 with broader morphology and non-globular scope evidence.

The script leaves v44 untouched. It adds the fixed-truth Galactic-centre-like,
thin-disk-like, and dwarf-galaxy-like benchmark plus the independent ANGST
M81/NGC 2976 single-reference comparison, while preserving the single-image
and synthetic/observational evidence boundaries.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import set_cell, set_font
from build_manuscript_v42_closed_book_scope import three_line_table
from build_manuscript_v44_independent_psf import after, find, replace_text, rewrite

ROOT = Path(__file__).resolve().parents[1]
MAIN_SOURCE = ROOT / "supplementary/AstroCFR_Crowded_Field_Manuscript_v44_independent_psf.docx"
MAIN_DEST = ROOT / "supplementary/AstroCFR_Crowded_Field_Manuscript_v45_scene_scope.docx"
SUP_SOURCE = ROOT / "supplementary/AstroCFR_Supplementary_Materials_v44.docx"
SUP_DEST = ROOT / "supplementary/AstroCFR_Supplementary_Materials_v45.docx"
SCENE_DIR = ROOT / "results/astrophysical_scene_benchmark"
ANGST_DIR = ROOT / "results/non_globular_runs/angst_single_reference"


def fmt_pct(value):
    return f"{100 * float(value):.1f}%"


def fmt_pp(value):
    return f"{100 * float(value):+.2f}"


def add_table(doc, headers, rows, size=7.0):
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, size=size)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell(cell, value, size=size)
    three_line_table(table)
    return table


def scene_lookup(payload, scene, method):
    return next(r for r in payload["results"] if r.get("scene") == scene and r.get("method") == method)


def angst_lookup(payload, field, method):
    return next(r for r in payload["results"] if r.get("field") == field and r.get("method") == method)


def main():
    scenes = json.loads((SCENE_DIR / "summary.json").read_text(encoding="utf-8"))
    angst = json.loads((ANGST_DIR / "summary.json").read_text(encoding="utf-8"))

    main_doc = Document(MAIN_SOURCE)
    abstract = find(main_doc, "We present AstroCFR")
    replace_text(
        abstract,
        "AstroCFR therefore provides a reproducible recovery-precision-cost frontier",
        "Additional scope checks cover PHAT M31, the ANGST M81-DEEP and NGC 2976 fields, and fixed-truth Galactic-centre-like, thin-disk-like, and dwarf-galaxy-like scenes; they preserve the recovery-precision-cost trade-off rather than identify a universal winner. AstroCFR therefore provides a reproducible recovery-precision-cost frontier",
    )

    anchor = find(main_doc, "We therefore ran an independent-PSF stress test")
    heading = after(anchor, "5.3.1 Cross-morphology and non-globular scope checks",
                    style="Heading 2", indent=0)

    gc_a = scene_lookup(scenes, "galactic_center_like", "astrocfr_epsf")
    gc_p = scene_lookup(scenes, "galactic_center_like", "photutils_psf")
    td_a = scene_lookup(scenes, "thin_disk_like", "astrocfr_epsf")
    td_p = scene_lookup(scenes, "thin_disk_like", "photutils_psf")
    dw_a = scene_lookup(scenes, "dwarf_galaxy_like", "astrocfr_epsf")
    dw_p = scene_lookup(scenes, "dwarf_galaxy_like", "photutils_psf")
    tests = {r["scene"]: r for r in scenes["paired_recovery_tests"]}
    synthetic_text = (
        "To broaden morphology coverage under exhaustive truth, we added three deterministic 384 x 384 single-image stress tests. "
        "They represent a Galactic-centre-like nuclear cusp with patchy extinction, a warped Milky-Way-thin-disk-like field with a dust lane, and an elliptical dwarf-galaxy-like resolved population. "
        "All methods receive the same Poisson/read-noise realization and are matched one-to-one within 2 pixels; stars are rendered with a spatially varying elliptical Moffat PSF that is not reused by any recovery branch. "
        f"AstroCFR ePSF versus Photutils recall is {fmt_pct(gc_a['recall'])} versus {fmt_pct(gc_p['recall'])} in the nuclear scene, "
        f"{fmt_pct(td_a['recall'])} versus {fmt_pct(td_p['recall'])} in the thin-disk scene, and {fmt_pct(dw_a['recall'])} versus {fmt_pct(dw_p['recall'])} in the dwarf scene. "
        f"The paired gains are {fmt_pp(tests['galactic_center_like']['paired_recall_difference'])}, {fmt_pp(tests['thin_disk_like']['paired_recall_difference'])}, and {fmt_pp(tests['dwarf_galaxy_like']['paired_recall_difference'])} percentage points, respectively. "
        "These gains are modest and require more CPU time; AstroCFR does not dominate precision or astrometric RMS in every scene. Full scene parameters, catalogues, and method-level results are in Supplementary Section S11."
    )
    synthetic = after(heading, synthetic_text)

    m81_a = angst_lookup(angst, "m81_deep", "astrocfr_epsf")
    m81_p = angst_lookup(angst, "m81_deep", "photutils_psf")
    n29_a = angst_lookup(angst, "ngc2976_deep", "astrocfr_epsf")
    n29_p = angst_lookup(angst, "ngc2976_deep", "photutils_psf")
    m81_h = angst_lookup(angst, "m81_deep", "astrocfr_photutils_hybrid")
    n29_h = angst_lookup(angst, "ngc2976_deep", "astrocfr_photutils_hybrid")
    angst_text = (
        "We also ran an observational non-globular comparison on the official ANGST F814W reference images and DOLPHOT-derived GST catalogues for M81-DEEP and the dwarf galaxy NGC 2976. "
        "The catalogue is finite, so the reported quantity is held-out GST recovery rather than blind completeness or purity. "
        f"AstroCFR ePSF recovers {fmt_pct(m81_a['catalogue_recovery'])} versus Photutils {fmt_pct(m81_p['catalogue_recovery'])} in M81 and "
        f"{fmt_pct(n29_a['catalogue_recovery'])} versus {fmt_pct(n29_p['catalogue_recovery'])} in NGC 2976. "
        f"Passing AstroCFR proposals to Photutils retains {fmt_pct(m81_h['catalogue_recovery'])} and {fmt_pct(n29_h['catalogue_recovery'])}, while improving AstroCFR's conditional measurement agreement, but it is slower and does not retain every proposal. "
        "Thus the non-globular evidence supports a two-layer candidate-plus-measurement product rather than a universal all-in-one branch."
    )
    after(synthetic, angst_text)

    limitation = find(main_doc, "The primary development experiments use CSST-like simulations")
    rewrite(
        limitation,
        "The primary development experiments use CSST-like simulations because real CSST survey images are not yet available. Frozen zero-shot transfer to Pan-STARRS1 and Legacy Survey data is unreliable, and target-adapted results are supervised few-shot adaptation rather than sensor-independent generalization. Observational single-image evidence now includes three ACSGGCT globular clusters, PHAT M31, ANGST M81-DEEP, and the dwarf galaxy NGC 2976. Galactic-centre and Milky-Way thin-disk coverage remains a controlled morphology simulation rather than a real-field validation. All these experiments use one stacked or reference image per field rather than a multi-exposure end-to-end catalogue-production comparison.",
    )
    after(
        find(main_doc, "The primary NGC 6752 fixed-scene artificial-star curves"),
        "The three morphology stress tests use exhaustive simulated truth and therefore permit genuine precision estimates, unlike finite observational comparison catalogues. They use one deterministic realization per morphology and an analytic spatially varying Moffat injector; consequently they test failure modes and branch trade-offs, not the population statistics of the Galactic centre, the Milky Way thin disk, or dwarf galaxies as a class.",
    )
    conclusion = find(main_doc, "The Pareto frontier exposes three practically useful operating points")
    replace_text(
        conclusion,
        "an independent Anderson-PSF injection on PHAT M31 preserves the faint-star proposal advantage over DAO/Photutils;",
        "an independent Anderson-PSF injection on PHAT M31 preserves the faint-star proposal advantage over DAO/Photutils; ANGST M81 and NGC 2976 retain higher GST recovery for AstroCFR ePSF than Photutils, while the three fixed-truth morphology tests show smaller recall gains with explicit precision and CPU costs;",
    )
    availability = find(main_doc, "Data availability:")
    rewrite(
        availability,
        "Data availability: The CSST challenge data are distributed through the National Astronomical Data Center CSST Data Challenge 2026 page and are not redistributed in this repository. The HST/ACS evaluation uses public MAST HLSP ACSGGCT v2 F606W images and official catalogues for NGC 6397, NGC 6752, and NGC 1851; the PHAT M31 and ANGST M81-DEEP/NGC 2976 inputs are also public HST products. The three morphology stress tests, their exhaustive truth catalogues, and all method outputs are generated reproducibly from the disclosed repository configuration.",
        indent=0,
    )
    main_doc.save(MAIN_DEST)

    sup = Document(SUP_SOURCE)
    rewrite(sup.add_paragraph(), "S11 Cross-morphology and non-globular scope expansion",
            style="Heading 1", indent=0)
    append = lambda text, style="Normal", indent=0.35: rewrite(
        sup.add_paragraph(), text, style=style, indent=indent)
    append(
        "The morphology benchmark uses three fixed-truth 384 x 384 images: a Galactic-centre-like nuclear cusp, a Milky-Way-thin-disk-like field, and a dwarf-galaxy-like resolved population. Source positions, extinction, structured backgrounds, and spatially varying elliptical Moffat PSFs are generated from fixed seeds. Poisson and Gaussian read noise are added once per scene, and the identical image is supplied to DAOStarFinder, SEP/SExtractor-style extraction, Photutils PSFPhotometry, and AstroCFR ePSF plus residual deblending. Evaluation uses exhaustive truth and greedy one-to-one matching within 2 pixels. The labels '-like' are mandatory: these scenes are morphology stress tests rather than observations of a named field."
    )
    pic = sup.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(SCENE_DIR / "astrophysical_scene_comparison.png"), width=Inches(6.4))
    append(
        "Fig. S18. Fixed-truth cross-morphology comparison. Left: one noisy realization for each disclosed scene. Right: recall and precision under the common 2-pixel one-to-one protocol. The three scenes are synthetic stress tests, not observed-field images.",
        style="Caption", indent=0,
    )

    append(
        "Table S22. Method-level fixed-truth results. Position RMS is one-coordinate-equivalent scatter from the matched radial offsets at 50 mas/pixel. Magnitude RMS is conditional scatter after removing one median method offset. Runtime is CPU wall time per megapixel on the same workstation.",
        style="Caption", indent=0,
    )
    scene_names = {
        "galactic_center_like": "Galactic-centre-like",
        "thin_disk_like": "Thin-disk-like",
        "dwarf_galaxy_like": "Dwarf-galaxy-like",
    }
    rows = []
    for r in scenes["results"]:
        if "error" in r:
            continue
        rows.append([
            scene_names[r["scene"]], r["label"], fmt_pct(r["recall"]), fmt_pct(r["precision"]),
            f"{r['f1']:.3f}", fmt_pct(r["high_density_ge_5_recall"]),
            f"{r.get('astrometric_rms_mas', float('nan')):.2f}",
            f"{r.get('photometric_rms_mag', float('nan')):.3f}",
            f"{r['runtime_s_per_mpix']:.2f}",
        ])
    add_table(sup, ["Scene", "Method", "Recall", "Precision", "F1", "High-density recall",
                    "Pos. RMS / mas", "Mag. RMS", "s/MPix"], rows, size=6.2)

    append(
        "Table S23. Paired recovery comparison between AstroCFR ePSF and Photutils. Confidence intervals resample the per-truth recovery difference; exact McNemar tests use discordant recovery outcomes.",
        style="Caption", indent=0,
    )
    pair_rows = []
    for r in scenes["paired_recovery_tests"]:
        lo, hi = r["paired_bootstrap_ci95"]
        pair_rows.append([
            scene_names[r["scene"]], f"{fmt_pp(r['paired_recall_difference'])} pp",
            f"[{100*lo:.2f}, {100*hi:.2f}] pp", str(r["a_only"]), str(r["b_only"]),
            f"{r['mcnemar_exact_p']:.3g}",
        ])
    add_table(sup, ["Scene", "Recall difference", "95% paired CI", "AstroCFR only",
                    "Photutils only", "McNemar p"], pair_rows, size=7.0)

    append(
        "Table S24. Independent ANGST non-globular single-reference-image comparison. GST recovery is catalogue-conditioned and is not blind completeness. The finite GST catalogue does not support a false-positive or blind-purity label for unmatched detections.",
        style="Caption", indent=0,
    )
    angst_rows = []
    for r in angst["results"]:
        if "error" in r:
            continue
        angst_rows.append([
            "M81-DEEP" if r["field"] == "m81_deep" else "NGC 2976",
            r["label"], fmt_pct(r["catalogue_recovery"]),
            fmt_pct(r["dense_f814w_le_27_recovery"]),
            f"{r.get('astrometric_rms_mas', float('nan')):.2f}",
            f"{r.get('photometric_rms_mag', float('nan')):.3f}",
            f"{r['runtime_s_per_mpix']:.2f}",
        ])
    add_table(sup, ["Field", "Method", "GST recovery", "Dense F814W<=27 recovery",
                    "Pos. RMS / mas", "F814W RMS / mag", "s/MPix"], angst_rows, size=6.5)
    append(
        "Interpretation boundary. The synthetic scenes close a morphology-coverage gap only at the controlled fixed-truth level. The ANGST fields add real non-globular single-reference-image evidence, including the dwarf galaxy NGC 2976, but remain catalogue-conditioned. A real Galactic-centre or Milky-Way thin-disk claim requires an independently deeper or artificial-star truth set. None of these additions changes the manuscript's single-image scope or establishes a DOLPHOT/ALLFRAME-class multi-exposure result."
    )
    sup.save(SUP_DEST)
    print(MAIN_DEST)
    print(SUP_DEST)


if __name__ == "__main__":
    main()
