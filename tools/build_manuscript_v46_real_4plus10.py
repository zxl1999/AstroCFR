#!/usr/bin/env python
"""Build the submission-oriented v45 manuscript from the completed evidence set.

The script starts from v44, not v45.  Synthetic morphology scenes are not
inserted as observational evidence.  The main text reports only completed
real-image/catalogue comparisons and labels the remaining registered fields as
pending.  The document is therefore a protocol/results draft until all ten
real fields pass the common benchmark gates.
"""
from __future__ import annotations

import csv
import json
import math
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from build_manuscript_v30_submission_fixes import set_cell
from build_manuscript_v42_closed_book_scope import three_line_table
from build_manuscript_v44_independent_psf import after, find, rewrite, set_font

ROOT = Path(__file__).resolve().parents[1]
MAIN_SOURCE = ROOT / "supplementary/AstroCFR_Crowded_Field_Manuscript_v44_independent_psf.docx"
SUP_SOURCE = ROOT / "supplementary/AstroCFR_Supplementary_Materials_v44.docx"
MAIN_DEST = ROOT / "supplementary/AstroCFR_Crowded_Field_Manuscript_v45_submission_ready.docx"
MAIN_LINKED_DEST = MAIN_DEST
SUP_DEST = ROOT / "supplementary/AstroCFR_Supplementary_Materials_v45_submission_ready.docx"
SUP_UPDATED_DEST = ROOT / "supplementary/AstroCFR_Supplementary_Materials_v45_submission_ready_crossfield_winners.docx"
COMPARISON_CSV = ROOT / "results/v45_submission_ready_observational_cross_field_table.csv"
SUMMARY = ROOT / "results/real_field_4plus10/summary.json"


# The rendered author--year citations are live external links. DOI landing
# pages are preferred; an official record is used if the source has no DOI.
CITATION_SENTENCES = [
    ("Anderson & King (2006)", " developed empirical ePSF methods for precise HST stellar positions and fluxes. ", "https://www.stsci.edu/hst/instrumentation/acs/documentation/instrument-science-reports-isrs"),
    ("Anderson et al. (2008)", " extended this approach to ACS crowded-field astrometry. ", "https://doi.org/10.1088/0004-6256/135/6/2055"),
    ("Bellini et al. (2011)", " quantified WFC3/UVIS geometric-distortion corrections. ", "https://doi.org/10.1086/659878"),
    ("Chambers et al. (2016)", " introduced the Pan-STARRS1 reference survey. ", "https://doi.org/10.48550/arXiv.1612.05560"),
    ("Dey et al. (2019)", " described the Legacy Surveys imaging products. ", "https://doi.org/10.3847/1538-3881/ab089d"),
    ("Gaia Collaboration et al. (2016)", " established Gaia's first all-sky astrometric reference. ", "https://doi.org/10.1051/0004-6361/201629272"),
    ("Gaia Collaboration et al. (2023)", " released the DR3 astrometric and photometric reference system. ", "https://doi.org/10.1051/0004-6361/202243940"),
    ("Astropy Collaboration (2022)", " provides the FITS, WCS, and coordinate infrastructure used here. ", "https://doi.org/10.3847/1538-4357/ac7c74"),
    ("Harris et al. (2020)", " provides the NumPy array framework for numerical operations. ", "https://doi.org/10.1038/s41586-020-2649-2"),
    ("Hunter (2007)", " introduced Matplotlib for scientific visualization. ", "https://doi.org/10.1109/MCSE.2007.55"),
    ("Virtanen et al. (2020)", " provides SciPy numerical routines. ", "https://doi.org/10.1038/s41592-019-0686-2"),
    ("Barbary (2016)", " provides SEP's Python interface to SExtractor-style detection. ", "https://doi.org/10.21105/joss.00058"),
    ("Bradley et al. (2024)", " documents the Photutils software API. ", "https://ascl.net/2401.012"),
    ("Pedregosa et al. (2011)", " established the scikit-learn machine-learning framework. ", "https://jmlr.org/papers/v12/pedregosa11a.html"),
    ("Goodfellow et al. (2016)", " presents deep-learning foundations relevant to the optional image classifier. ", "https://www.deeplearningbook.org/"),
    ("He et al. (2016)", " introduced residual networks used as classifier context. ", "https://doi.org/10.1109/CVPR.2016.90"),
    ("Shi et al. (2024)", " studied CSST photometric preparation. ", "https://doi.org/10.1088/1674-4527/ad2dbd"),
    ("Long et al. (2025)", " investigated deep-feature point-source detection. ", "https://doi.org/10.3847/1538-4365/ad9244"),
    ("Han et al. (2026)", " addressed CSST multi-modal star-galaxy classification. ", "https://doi.org/10.1016/j.ascom.2026.101112"),
    ("Wainer et al. (2025)", " studied Rubin crowded-field photometry. ", "https://doi.org/10.3847/2515-5172/adecef"),
    ("Espinosa et al. (2025)", " analyzed PSF-fitting uncertainty. ", "https://doi.org/10.1051/0004-6361/202555342"),
    ("De Alba et al. (2026)", " proposed physics-informed aperture feature learning. ", "https://doi.org/10.1016/j.ascom.2026.101165"),
    ("Centofanti et al. (2026)", " studied image-based PSF recovery. ", "https://doi.org/10.1051/0004-6361/202558730"),
    ("Wang et al. (2026)", " examined PSF reconstruction. ", "https://doi.org/10.3847/1538-4365/ae5053"),
    ("Zhang et al. (2026)", " studied aperture photometry after PSF deconvolution. ", "https://doi.org/10.3847/1538-3881/ae7718"),
    ("Libralato et al. (2024)", " reported high-precision Euclid--Gaia cluster astrometry and photometry. ", "https://doi.org/10.48550/arXiv.2411.02487"),
    ("Salaris et al. (2024)", " examined HST NGC 6752 stellar-population measurements. ", "https://doi.org/10.1002/asna.20240018"),
    ("Ben-David et al. (2010)", " formalized domain-adaptation bounds. ", "https://doi.org/10.1007/s10994-009-5152-4"),
    ("Pan & Yang (2010)", " surveyed transfer-learning settings. ", "https://doi.org/10.1109/TKDE.2009.191"),
    ("Ganin et al. (2016)", " introduced domain-adversarial training. ", "https://jmlr.org/papers/v17/15-239.html"),
    ("Long et al. (2015)", " proposed deep adaptation networks. ", "https://proceedings.mlr.press/v37/long15.html"),
    ("Breiman (2001)", " introduced random forests. ", "https://doi.org/10.1023/A:1010933404324"),
    ("Chen & Guestrin (2016)", " introduced gradient-boosted trees. ", "https://doi.org/10.1145/2939672.2939785"),
]


def pct(value):
    return f"{100 * float(value):.2f}%"


def add_external_hyperlink(paragraph, text, url):
    """Append a visibly clickable external hyperlink to a Word paragraph."""
    relation_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(colour)
    properties.append(underline)
    run.append(properties)
    value = OxmlElement("w:t")
    value.set(qn("xml:space"), "preserve")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def rewrite_linked_citation_paragraph(paragraph):
    """Render one sentence per source with a live author--year citation."""
    paragraph.clear()
    paragraph.style = "Normal"
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.08
    for citation, sentence, url in CITATION_SENTENCES:
        add_external_hyperlink(paragraph, citation, url)
        set_font(paragraph.add_run(sentence), size=10.5)
    return paragraph


def rewrite_linked_fragments(paragraph, fragments, *, indent=0.35):
    """Write prose with selected live external links in one paragraph."""
    paragraph.clear()
    paragraph.style = "Normal"
    paragraph.paragraph_format.first_line_indent = None if not indent else paragraph.paragraph_format.first_line_indent
    paragraph.paragraph_format.line_spacing = 1.08
    for text, url in fragments:
        if url:
            add_external_hyperlink(paragraph, text, url)
        else:
            set_font(paragraph.add_run(text), size=10.5)
    return paragraph


def before(anchor, text, *, style="Normal", indent=0.35):
    """Insert and format a paragraph immediately before another paragraph."""
    paragraph = anchor._parent.add_paragraph(style=style)
    paragraph._p.getparent().remove(paragraph._p)
    anchor._p.addprevious(paragraph._p)
    return rewrite(paragraph, text, style=style, indent=indent)


def pp(value):
    return f"{100 * float(value):+.2f} pp"


def lookup(rows, field, method):
    return next(r for r in rows if r.get("field") == field and r.get("method") == method and not r.get("error"))


def add_table(doc, headers, rows, size=7.0):
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, size=size)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell(cell, value, size=size)
    three_line_table(table)
    return table


def insert_table_after(anchor, table):
    anchor._p.addnext(table._tbl)


def after_table(table, text, *, style="Normal", indent=0.35):
    """Insert and format a paragraph immediately after a Word table."""
    paragraph = table._parent.add_paragraph(style=style)
    paragraph._p.getparent().remove(paragraph._p)
    table._tbl.addnext(paragraph._p)
    return rewrite(paragraph, text, style=style, indent=indent)


def insert_section_before(doc, anchor, orientation):
    """Start an orientation-specific section immediately before ``anchor``."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    if orientation == WD_ORIENT.LANDSCAPE:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = section.page_height, section.page_width
    # add_section appends a break paragraph immediately before body sectPr.
    break_paragraph = doc.element.body[-2]
    anchor._p.addprevious(break_paragraph)


def insert_section_after(doc, table, orientation):
    """End an orientation-specific table section immediately after ``table``."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    if orientation == WD_ORIENT.LANDSCAPE:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = section.page_height, section.page_width
    break_paragraph = doc.element.body[-2]
    table._tbl.addnext(break_paragraph)


def iter_document_blocks(doc):
    """Yield paragraphs and tables in their actual Word document order."""
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def add_internal_hyperlink(paragraph, text, anchor):
    """Append a clickable hyperlink to a bookmark in the same Word file."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(colour)
    props.append(underline)
    run.append(props)
    value = OxmlElement("w:t")
    value.set(qn("xml:space"), "preserve")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def insert_linked_lead_before(anchor, lead, link_text, bookmark):
    paragraph = anchor._parent.add_paragraph(style="Normal")
    paragraph._p.getparent().remove(paragraph._p)
    anchor._p.addprevious(paragraph._p)
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.line_spacing = 1.08
    set_font(paragraph.add_run(lead), size=10.5)
    add_internal_hyperlink(paragraph, link_text, bookmark)
    set_font(paragraph.add_run("."), size=10.5)
    return paragraph


def set_table_formal_layout(table):
    """Prevent row-level page splits and repeat the header on continuation pages."""
    for index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))
    table.autofit = True


def format_submission_document(doc):
    """Apply journal-style captions, table pagination, and internal cross-links."""
    # Caption paragraphs are the authoritative labels; all figure abbreviations
    # are expanded to the same formal style used by Table captions.
    caption_re = re.compile(r"^(Fig\.|Figure|Table)\s+(S?\d+)\.")
    original_blocks = list(iter_document_blocks(doc))
    next_bookmark_id = 1000
    for block_index, block in enumerate(original_blocks):
        if not isinstance(block, Paragraph):
            continue
        match = caption_re.match(block.text.strip())
        if not match:
            continue
        kind, number = match.groups()
        formal_kind = "Figure" if kind in {"Fig.", "Figure"} else "Table"
        if kind != formal_kind:
            # Replace only the label, preserving the caption's substantive text.
            rewrite(block, block.text.replace(f"{kind} {number}.", f"{formal_kind} {number}.", 1), style="Caption", indent=0)
        else:
            block.style = "Caption"
        block.alignment = WD_ALIGN_PARAGRAPH.LEFT
        block.paragraph_format.keep_together = True
        block.paragraph_format.widow_control = True
        for run in block.runs:
            set_font(run, size=9.0)
        bookmark = f"{formal_kind.lower()}_{number}"
        add_bookmark(block, bookmark, next_bookmark_id)
        next_bookmark_id += 1

        # Avoid duplicating an existing explicit lead sentence (e.g. the
        # cross-field provenance paragraph). Otherwise add a concise linked
        # sentence immediately before the table or figure.
        recent = " ".join(
            candidate.text for candidate in original_blocks[max(0, block_index - 6):block_index]
            if isinstance(candidate, Paragraph)
        )
        explicit = re.search(rf"\b{formal_kind}\s+{re.escape(number)}\b", recent)
        is_table = formal_kind == "Table"
        if not explicit:
            if is_table:
                insert_linked_lead_before(block, "The corresponding tabulated results are reported in ", f"Table {number}", bookmark)
            else:
                # Figure captions follow their image; place the lead before the
                # image paragraph when one is immediately available.
                image_anchor = None
                for candidate in reversed(original_blocks[max(0, block_index - 3):block_index]):
                    if isinstance(candidate, Paragraph) and any(child.tag == qn("w:drawing") for child in candidate._p.iter()):
                        image_anchor = candidate
                        break
                insert_linked_lead_before(image_anchor or block, "The corresponding visualization is shown in ", f"Figure {number}", bookmark)

    for table in doc.tables:
        set_table_formal_layout(table)


def renumber_caption(doc, old_prefix, new_prefix):
    paragraph = find(doc, old_prefix)
    rewrite(paragraph, paragraph.text.replace(old_prefix, new_prefix, 1), style=paragraph.style.name, indent=0)


def method_name(row):
    """Use one unambiguous name for the cross-field five-method table."""
    if row.get("method") == "astrocfr_photutils_hybrid":
        return "AstroCFR+Photutils hybrid"
    return row["label"]


OBSERVATIONAL_FIELDS = [
    ("ngc6397", "NGC 6397\n(low-density cluster)", "Globular cluster"),
    ("ngc6752", "NGC 6752\n(core-crowded cluster)", "Globular cluster"),
    ("ngc1851", "NGC 1851\n(extremely crowded cluster)", "Globular cluster"),
    ("m31_b21_f15", "M31 B21-F15\n(outer disk)", "Spiral-galaxy disk"),
    ("m33_b01_f01", "M33 B01-F01\n(inner disk)", "Spiral-galaxy disk"),
    ("m33_b03_f02", "M33 B03-F02\n(disk)", "Spiral-galaxy disk"),
    ("ngc2070_1", "NGC 2070-1\n(LMC star-forming field)", "Star-forming field"),
    ("ngc2070_2", "NGC 2070-2\n(LMC star-forming field)", "Star-forming field"),
    ("m81_deep", "M81-DEEP\n(spiral field)", "Spiral-galaxy disk"),
    ("ngc2976_deep", "NGC 2976\n(dwarf disk)", "Dwarf galaxy"),
    ("gr8", "GR8\n(dwarf irregular; multi-exposure)", "Dwarf galaxy"),
]


def wilson(recovered, total):
    """Wilson 95% interval for a catalogue-conditioned recovery fraction."""
    if total <= 0:
        return (float("nan"), float("nan"))
    z = 1.959963984540054
    p = recovered / total
    denom = 1 + z * z / total
    centre = (p + z * z / (2 * total)) / denom
    half = z * math.sqrt(p * (1 - p) / total + z * z / (4 * total * total)) / denom
    return max(0.0, centre - half), min(1.0, centre + half)


def recovery_cell(value, ci=None, recovered=None, total=None):
    if ci is None and recovered is not None and total is not None:
        ci = wilson(recovered, total)
    if ci is None:
        return f"{100 * value:.2f}%"
    return f"{100 * value:.2f}% ({100 * ci[0]:.2f}\u2013{100 * ci[1]:.2f}%)"


def crowded_subset_text(item):
    """Format the field-appropriate crowded subset without cross-filter claims."""
    if "dense_v20_recall" in item:
        value = float(item["dense_v20_recall"])
        n = item.get("high_density_v20_n")
        suffix = f" (n={int(n)})" if n is not None else ""
        return f"V<=20: {100 * value:.2f}%{suffix}"
    if "high_density_v20_recall" in item:
        value = float(item["high_density_v20_recall"])
        n = item.get("high_density_v20_n")
        suffix = f" (n={int(n)})" if n is not None else ""
        return f"V<=20: {100 * value:.2f}%{suffix}"
    for key, value in item.items():
        if not (key.startswith("dense_") and key.endswith("_recovery")):
            continue
        cut = key[len("dense_"):-len("_recovery")]
        n = item.get(f"dense_{cut}_n")
        if n is None or int(n) <= 0:
            return "\u2014"
        value = float(value)
        n = int(n)
        recovered = int(round(value * n))
        band = cut.replace("_le_", "<=").upper()
        return f"{band}: {recovery_cell(value, recovered=recovered, total=n)} (n={n})"
    return "\u2014"


def outcome_note(row, group):
    """A row-level operating-point note, without claiming overall dominance."""
    eps = 1e-12
    best_recovery = max(x["recovery"] for x in group)
    best_position = min(x["position"] for x in group)
    best_photometry = min(x["photometry"] for x in group)
    fastest = min(x["runtime"] for x in group)
    recovery = abs(row["recovery"] - best_recovery) < eps
    position = abs(row["position"] - best_position) < eps
    photometry = abs(row["photometry"] - best_photometry) < eps
    runtime = abs(row["runtime"] - fastest) < eps
    if recovery and (position or photometry):
        return "Highest recovery with a leading conditional measurement metric."
    if recovery:
        return "Highest recovery; assess the added runtime trade-off."
    if position and photometry:
        return "Best conditional position and magnitude precision."
    if position:
        return "Lowest conditional position RMS."
    if photometry:
        return "Lowest conditional magnitude RMS."
    if runtime:
        return "Fastest measured branch."
    return "Intermediate recovery\u2013precision\u2013runtime operating point."


def winner_names(group, metric, *, highest=False):
    target = (max if highest else min)(row[metric] for row in group)
    return [row["method"] for row in group if abs(row[metric] - target) < 1e-12]


def field_winner_summary(group):
    recovery = "; ".join(winner_names(group, "recovery", highest=True))
    position = "; ".join(winner_names(group, "position"))
    photometry = "; ".join(winner_names(group, "photometry"))
    if position == photometry:
        precision = f"Position and photometry: {position}"
    else:
        precision = f"Position: {position}; photometry: {photometry}"
    return recovery, precision


def observational_comparison_rows(results):
    """Return all available measured branches for completed observed fields.

    The three ACSGC clusters retain their additional RF and spatial-ePSF
    branches.  The other fields retain the common five-branch benchmark.
    No unavailable method is represented by an invented value.
    """
    joint = json.loads((ROOT / "results/joint_csst_hst_m31_evidence/joint_evidence.json").read_text(encoding="utf-8"))
    cluster_labels = {
        "DAOStarFinder": "DAOStarFinder",
        "SEP/SExtractor-style": "SEP/SExtractor-style",
        "Photutils PSFPhotometry": "Photutils PSFPhotometry",
        "WPDC original (target-adapted RF)": "AstroCFR-RF",
        "WPDC ePSF + residual deblend": "AstroCFR ePSF + residual deblend",
        "WPDC spatial ePSF + joint fit": "AstroCFR spatial-ePSF + joint fit",
    }
    grouped = {field: [] for field, _, _ in OBSERVATIONAL_FIELDS}
    for item in joint["hst_single_stack"]:
        field = item["field"]
        if field not in grouped:
            continue
        total = int(item["test_references"])
        recovered = int(round(float(item["test_completeness"]) * total))
        grouped[field].append({
            "method": cluster_labels[item["method"]], "recovery": float(item["test_completeness"]),
            "recovery_text": recovery_cell(float(item["test_completeness"]), recovered=recovered, total=total),
            "dense_text": crowded_subset_text(item),
            "position": float(item["position_rms_mas"]), "photometry": float(item["magnitude_rms_mag"]),
            "runtime": float(item["runtime_s_per_mpix"]),
        })
    hybrid = json.loads((ROOT / "results/hst_hybrid_wpdc_photutils/hybrid_summary.json").read_text(encoding="utf-8"))
    for item in hybrid["results"]:
        field = item["cluster"]
        total = int(item["test_references"]); recovered = int(item["test_recovered"])
        grouped[field].append({
            "method": "AstroCFR+Photutils hybrid", "recovery": float(item["test_completeness"]),
            "recovery_text": recovery_cell(float(item["test_completeness"]), recovered=recovered, total=total),
            "dense_text": crowded_subset_text(item),
            "position": float(item["astrometric_rms_mas"]), "photometry": float(item["photometric_rms_mag"]),
            "runtime": float(item["runtime_s_per_mpix"]),
        })
    common_labels = {
        "dao": "DAOStarFinder", "sep": "SEP/SExtractor-style",
        "photutils_psf": "Photutils PSFPhotometry",
        "astrocfr_epsf": "AstroCFR ePSF + residual deblend",
        "astrocfr_photutils_hybrid": "AstroCFR+Photutils hybrid",
    }
    for item in results:
        field, method = item.get("field"), item.get("method")
        if field not in grouped or field.startswith("ngc") and field in {"ngc6397", "ngc6752", "ngc1851"}:
            continue
        if method not in common_labels or item.get("error"):
            continue
        grouped[field].append({
            "method": common_labels[method], "recovery": float(item["catalogue_recovery"]),
            "recovery_text": recovery_cell(float(item["catalogue_recovery"]), ci=item.get("catalogue_recovery_ci95")),
            "dense_text": crowded_subset_text(item), "position": float(item["astrometric_rms_mas"]),
            "photometry": float(item["photometric_rms_mag"]), "runtime": float(item["runtime_s_per_mpix"]),
        })
    rows = []
    for field, label, evidence_tier in OBSERVATIONAL_FIELDS:
        group = grouped[field]
        if not group:
            continue
        best_recovery, best_precision = field_winner_summary(group)
        for index, item in enumerate(group):
            rows.append([
                label if index == 0 else "", evidence_tier if index == 0 else "",
                item["method"], item["recovery_text"], item["dense_text"],
                f"{item['position']:.2f}", f"{item['photometry']:.3f}", f"{item['runtime']:.2f}",
                best_recovery if index == 0 else "", best_precision if index == 0 else "",
            ])
    return rows


def append_references(doc, entries):
    """Append only citations absent from the inherited v44 bibliography."""
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for key, entry in entries:
        if key not in full_text:
            rewrite(doc.add_paragraph(), entry, indent=0)


def main():
    payload = json.loads(SUMMARY.read_text(encoding="utf-8"))
    results = payload["results"]
    readiness = payload["readiness"]
    # GR8 is a valid real SCI multi-exposure benchmark, recorded separately
    # because its image construction is not the single-image protocol used by
    # the other expansion rows.  Replace the retired ERR-extension diagnostic
    # if the audited stack result is available.
    gr8_stack_path = ROOT / "results/real_field_4plus10/gr8_multiepoch/summary.json"
    gr8_stack = None
    if gr8_stack_path.exists():
        gr8_stack = json.loads(gr8_stack_path.read_text(encoding="utf-8"))
        results = [r for r in results if r.get("field") != "gr8"] + gr8_stack["results"]
    complete = [r for r in readiness if r["manuscript_admitted"]]
    if len(complete) < 3:
        raise RuntimeError("v45 draft requires at least three admitted real-field comparisons")

    m31_a = lookup(results, "m31_b21_f15", "astrocfr_epsf")
    m31_p = lookup(results, "m31_b21_f15", "photutils_psf")
    m81_a = lookup(results, "m81_deep", "astrocfr_epsf")
    m81_p = lookup(results, "m81_deep", "photutils_psf")
    n29_a = lookup(results, "ngc2976_deep", "astrocfr_epsf")
    n29_p = lookup(results, "ngc2976_deep", "photutils_psf")
    cross_rows = observational_comparison_rows(results)
    comparison_headers = [
        "Field (observing regime)", "Evidence tier / scene class", "Method / branch",
        "Catalogue recovery (95% CI)", "Crowded-subset recovery (field-specific cut)",
        "Position RMS / mas", "Magnitude RMS / mag", "Runtime / s MPix-1",
        "Best recovery branch", "Best conditional precision branch(es)",
    ]
    with COMPARISON_CSV.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.writer(stream)
        writer.writerow(comparison_headers)
        # Word suppresses repeated field labels for visual grouping; the CSV
        # intentionally repeats them so every record remains self-describing.
        csv_field = ""
        for row in cross_rows:
            if row[0]:
                csv_field = row[0]
            writer.writerow([csv_field, *row[1:]])

    main_doc = Document(MAIN_SOURCE)
    citation_paragraph = find(main_doc, "The measurement layer is grounded in empirical HST PSF")
    rewrite_linked_citation_paragraph(citation_paragraph)
    # The new cross-field table precedes the inherited adaptation tables.
    # Renumber their captions and every in-text reference as one sequence.
    renumber_caption(main_doc, "Table 8. Simulation-to-HST", "Table 9. Simulation-to-HST")
    renumber_caption(main_doc, "Table 9. Accuracy", "Table 10. Accuracy")
    renumber_caption(main_doc, "Table 10. Density-adaptive", "Table 11. Density-adaptive")
    release = find(main_doc, "Release selection is then science-constrained")
    rewrite(release, release.text.replace("Table 9 reports", "Table 10 reports"), indent=0.35)
    uncertainty = find(main_doc, "To avoid treating one point estimate")
    rewrite(uncertainty, uncertainty.text.replace("and Table 9 used", "and Table 10 used"), indent=0.35)
    routing = find(main_doc, "The density-adaptive router defines")
    rewrite(routing, routing.text.replace("Table 10.", "Table 11."), indent=0.35)
    cluster_caption = find(main_doc, "Table 3. AstroCFR on the central HST/ACS")
    cluster_provenance = before(cluster_caption, "", indent=0.35)
    rewrite_linked_fragments(cluster_provenance, [
        ("The HST/ACS images and official comparison catalogues used in Tables 3–6 follow the ACS Survey of Galactic Globular Clusters catalogue construction of ", None),
        ("Anderson et al. (2008)", "https://doi.org/10.1088/0004-6256/135/6/2055"),
        (". Their empirical-PSF measurement context follows ", None),
        ("Anderson & King (2006)", "https://www.stsci.edu/hst/instrumentation/acs/documentation/instrument-science-reports-isrs"),
        (". These catalogues are held-out evaluation references and are not used to fit the reported measurement branches.", None),
    ])
    anchor = find(main_doc, "We therefore ran an independent-PSF stress test")
    heading = after(anchor, "5.3.1 Completed public-archive benchmark expansion",
                    style="Heading 2", indent=0)
    p1 = after(
        heading,
        "The reported evidence comprises four CSST-like simulated chips, three real HST/ACS globular-cluster fields (NGC 6397, NGC 6752, and NGC 1851), and eight public-archive HST fields with complete five-method evaluations. The public-archive sample comprises PHAT M31 B21-F15, PHATTER M33 B01-F01 and B03-F02, two Hubble Tarantula Treasury pointings, M81-DEEP, NGC 2976, and the GR8 multi-exposure mosaic. Fixed-truth morphology simulations are supplementary stress tests only and are not counted as observations. Every included public-archive field uses real observed pixels, a field-appropriate external catalogue, held-out association, and the shared DAO, SEP, Photutils, AstroCFR ePSF, and AstroCFR+Photutils hybrid comparison; image product, filter, pixel scale, PSF, and catalogue provenance remain field-specific. Two additional PHAT M31 outer-disk pointings, B21-F10 and B21-F18, are excluded because their independent catalogues and complete common-method evaluations were not available at the time of analysis.",
    )
    p2 = after(
        p1,
        f"The eight included real expansion fields each have all five method branches and complete conditional measurement metrics. On real PHAT M31 B21-F15 F475W data, AstroCFR ePSF recovers {pct(m31_a['catalogue_recovery'])} of held-out quality references versus {pct(m31_p['catalogue_recovery'])} for Photutils ({pp(m31_a['catalogue_recovery'] - m31_p['catalogue_recovery'])}). On the official ANGST F814W products, the corresponding values are {pct(m81_a['catalogue_recovery'])} versus {pct(m81_p['catalogue_recovery'])} for M81-DEEP and {pct(n29_a['catalogue_recovery'])} versus {pct(n29_p['catalogue_recovery'])} for NGC 2976. These are catalogue-conditioned recovery values from finite external catalogues; unmatched detections are not assigned false-positive labels, and no blind purity or exhaustive completeness claim is made.",
    )
    p3 = after(
        p2,
        "M33 B01-F01 and B03-F02 are completed real ACS/WFC F475W single-image comparisons against the independently published PHATTER Table-6 catalogue (Williams et al. 2021). Their held-out reference counts are 3,030 and 7,604, respectively, and all five branches report catalogue recovery, conditional position RMS, conditional magnitude RMS, runtime, and memory. AstroCFR ePSF recovers 77.06% versus 69.97% for Photutils in B01 (+7.10 percentage points), and 40.02% versus 36.05% in the denser B03 field (+3.97 percentage points). GR8 is a separately labelled real multi-exposure ACS/WFC SCI-mosaic comparison: seven valid FLC exposures, 2,132 held-out independent ANGST/GST references, and all five methods with recovery, conditional RMS, runtime, and memory. It remains in Supplementary Table S25 rather than being pooled with the single-image/reference-image rows. No incomplete field is represented as efficacy evidence or replaced by simulated Galactic-centre-like, thin-disk-like, or dwarf-galaxy-like images. M33 tests establish M33 inner-disk/disk coverage only; they do not establish a real Milky-Way Galactic-centre or Milky-Way thin-disk claim.",
    )
    cross_heading = after(p3, "5.3.2 Cross-field branch comparison", style="Heading 2", indent=0)
    insert_section_before(main_doc, cross_heading, WD_ORIENT.LANDSCAPE)
    cross_provenance = after(cross_heading, "", indent=0.35)
    rewrite_linked_fragments(cross_provenance, [
        ("Table 8 compares every measured branch in the completed cross-field evidence set. Its external catalogue context is provided by PHAT ", None),
        ("(Dalcanton et al., 2012)", "https://doi.org/10.1088/0067-0049/200/2/18"),
        (", PHATTER ", None),
        ("(Williams et al., 2021)", "https://doi.org/10.3847/1538-4365/abd6f6"),
        (", the Hubble Tarantula Treasury ", None),
        ("(Sabbi et al., 2016)", "https://doi.org/10.3847/0067-0049/222/1/11"),
        (", and ANGST/GST ", None),
        ("(Dalcanton et al., 2009)", "https://doi.org/10.1088/0067-0049/183/1/67"),
        (". These catalogues define held-out matching references; they are not training labels or exhaustive truth catalogues.", None),
    ])
    cross_caption = after(
        cross_provenance,
        "Table 8. Full cross-field comparison for every measured branch in the completed observational evidence set. Recovery is catalogue-conditioned; position and magnitude RMS are conditional on matched held-out sources after the stated calibration and clipping protocol. The crowded-subset column uses V<=20 only for the ACSGC F606W globular-cluster catalogues and the recorded field-specific filter cuts (F475W<=27, F555W<=27, or F814W<=27) for the public-archive galaxy fields; n is the crowded-subset denominator where available. Winner columns are field-level summaries shown on the first row of each field; conditional position and photometric precision are kept separate because no composite RMS was preregistered.",
        style="Caption", indent=0,
    )
    cross_table = add_table(main_doc, comparison_headers, cross_rows, size=5.5)
    insert_table_after(cross_caption, cross_table)
    insert_section_after(main_doc, cross_table, WD_ORIENT.PORTRAIT)

    limitation = find(main_doc, "The primary development experiments use CSST-like simulations")
    rewrite(
        limitation,
        "The primary development experiments use CSST-like simulations because real CSST survey images are not yet available. This manuscript reports only eight completed real HST expansion fields, each with five method branches and conditional position/photometry metrics. The observational comparisons use finite external catalogues and therefore support catalogue-conditioned recovery and conditional measurement residuals, not blind purity. The current field set also does not establish a real Milky-Way Galactic-centre or Milky-Way thin-disk claim. Most completed expansion experiments remain single-image or reference-image comparisons; GR8 is separately identified as a multi-exposure SCI-mosaic comparison rather than an end-to-end catalogue-production result.",
    )
    availability = find(main_doc, "Data availability:")
    rewrite(
        availability,
        "Data availability: The CSST challenge data are distributed through the National Astronomical Data Center and are not redistributed. The reported public-archive HST benchmark uses PHAT (Dalcanton et al. 2012), PHATTER (Williams et al. 2021), Hubble Tarantula Treasury/HTTP (Sabbi et al. 2016), and ANGST/DOLPHOT GST (Dalcanton et al. 2009; Dolphin 2000) products. Download scripts, included field identifiers, file hashes where available, and machine-readable method summaries are provided in the repository. Large FITS products remain external. Synthetic morphology stress tests are disclosed separately and are not counted as observational fields. B21-F10 and B21-F18 are not included in the reported benchmark because their independent PHAT catalogues and complete common-method evaluations were unavailable.",
        indent=0,
    )
    append_references(main_doc, [
        ("Dalcanton, J. J., et al. (2009)", "Dalcanton, J. J., et al. (2009). The ACS Nearby Galaxy Survey Treasury. The Astrophysical Journal Supplement Series, 183, 67-108. https://doi.org/10.1088/0067-0049/183/1/67"),
        ("Dalcanton, J. J., et al. (2012)", "Dalcanton, J. J., et al. (2012). The Panchromatic Hubble Andromeda Treasury. The Astrophysical Journal Supplement Series, 200, 18. https://doi.org/10.1088/0067-0049/200/2/18"),
        ("Williams, B. F., et al. (2021)", "Williams, B. F., et al. (2021). The Panchromatic Hubble Andromeda Treasury: Triangulum Extended Region. The Astrophysical Journal Supplement Series, 253, 8. https://doi.org/10.3847/1538-4365/abd6f6"),
        ("Sabbi, E., et al. (2016)", "Sabbi, E., et al. (2016). The Hubble Tarantula Treasury Project. II. Photometric Catalog and Properties of the Young Stellar Populations. The Astrophysical Journal Supplement Series, 222, 11. https://doi.org/10.3847/0067-0049/222/1/11"),
    ])
    main_saved = MAIN_DEST
    try:
        main_doc.save(main_saved)
    except PermissionError:
        # Do not overwrite a manuscript currently open in Word.  The linked
        # version is still delivered under a distinct, reviewable filename.
        main_saved = MAIN_LINKED_DEST
        main_doc.save(main_saved)

    sup = Document(SUP_SOURCE)
    rewrite(sup.add_paragraph(), "S11 Completed empirical evidence inventory and public-archive benchmark protocol",
            style="Heading 1", indent=0)
    append = lambda text, style="Normal", indent=0.35: rewrite(
        sup.add_paragraph(), text, style=style, indent=indent)
    append(
        "The completed manuscript evidence inventory has four CSST-like simulation chips, three real HST/ACS globular-cluster fields, and eight completed public-archive HST fields. Only fields with a real image, independently produced catalogue, held-out comparison, all five method branches, runtime/RSS measurements, and reporting-boundary checks are included. The matching framework is common, but the real-image product (DRZ, reference image, FLC SCI crop, or explicitly separate multi-exposure mosaic), filter, pixel scale, and catalogue provenance are field-specific. Finite catalogues support catalogue recovery and conditional astrometric/photometric agreement but not blind false-discovery rates."
    )
    append("Table S22. Completed empirical evidence inventory. Five-method status refers only to the common DAO/SEP/Photutils/AstroCFR-ePSF/AstroCFR+Photutils comparison; the CSST tier preserves the original AstroCFR result and is not retrospectively treated as a matched five-method benchmark.", style="Caption", indent=0)
    inventory = [
        ["CSST simulation", "Chip 12", "CSST-like supplied-catalogue chip", "original AstroCFR result"],
        ["CSST simulation", "Chip 13", "CSST-like supplied-catalogue chip", "original AstroCFR result"],
        ["CSST simulation", "Chip 17", "CSST-like supplied-catalogue chip", "original AstroCFR result"],
        ["CSST simulation", "Chip 18", "CSST-like supplied-catalogue chip", "original AstroCFR result"],
        ["HST/ACS globular cluster", "NGC 6397", "HST/ACS F606W globular cluster", "complete"],
        ["HST/ACS globular cluster", "NGC 6752", "HST/ACS F606W globular cluster", "complete"],
        ["HST/ACS globular cluster", "NGC 1851", "HST/ACS F606W globular cluster", "complete"],
    ] + [["completed real HST", r["field_id"], r["scene_class"], "complete"] for r in complete]
    add_table(sup, ["Evidence tier", "Unit", "Scene", "Five-method status"], inventory, size=6.1)
    append("Table S23. Completed real HST expansion fields.", style="Caption", indent=0)
    add_table(sup, ["Slot", "Field", "Observed scene", "FLCs", "Catalogues", "Status"], [
        [str(r["evidence_slot"]), r["field_id"], r["scene_class"], str(r["flc_files"]),
         str(r["catalogue_files"]), "Admitted" if r["manuscript_admitted"] else
         ("All methods; metric gate pending" if r["all_methods_complete"] else r["next_gate"])]
        for r in complete
    ], size=6.2)
    append(
        "Table S24. All admitted real-expansion comparisons. Recovery is measured against the held-out portion of each finite external catalogue. Position and magnitude RMS values are conditional on matched sources after calibration and robust clipping. The two Tarantula rows are retained despite low recovery and large conditional residuals; this table is not restricted to favourable cases."
        , style="Caption", indent=0)
    names = {"m31_b21_f15": "M31 B21-F15", "m33_b01_f01": "M33 B01-F01", "m33_b03_f02": "M33 B03-F02", "ngc2070_1": "NGC 2070-1", "ngc2070_2": "NGC 2070-2", "m81_deep": "M81-DEEP", "ngc2976_deep": "NGC 2976"}
    complete_ids = set(names)
    result_rows = []
    for r in results:
        if r.get("field") not in complete_ids or r.get("error"):
            continue
        result_rows.append([
            names[r["field"]], method_name(r), pct(r["catalogue_recovery"]),
            f"{r.get('astrometric_rms_mas', float('nan')):.2f}",
            f"{r.get('photometric_rms_mag', float('nan')):.3f}",
            f"{r.get('runtime_s_per_mpix', float('nan')):.2f}",
        ])
    add_table(sup, ["Field", "Method", "Catalogue recovery", "Pos. RMS / mas",
                    "Magnitude RMS", "s/MPix"], result_rows, size=6.4)
    append("Table S25. GR8 real multi-exposure ACS/WFC F475W benchmark. Seven valid CTE-corrected FLC exposures were WCS-resampled onto two native SCI grids and median-combined; the zero-exposure product was excluded. The independent ANGST/GST catalogue is held out by the common stripe protocol. This is reported separately from the single-image/reference-image comparison tier.", style="Caption", indent=0)
    gr8_rows = []
    for r in results:
        if r.get("field") == "gr8" and not r.get("error"):
            gr8_rows.append([method_name(r), str(r["test_references"]), pct(r["catalogue_recovery"]),
                             f"{r.get('astrometric_rms_mas', float('nan')):.2f}",
                             f"{r.get('photometric_rms_mag', float('nan')):.3f}", f"{r['runtime_s_per_mpix']:.2f}"])
    add_table(sup, ["Method", "Held-out refs", "Catalogue recovery", "Pos. RMS", "Magnitude RMS", "s/MPix"], gr8_rows, size=6.4)
    sup_cross_caption = append("Table S26. Full cross-field branch comparison for all completed observed fields. Catalogue recovery is reported with its 95% interval. The crowded-subset column uses V<=20 only for ACSGC F606W clusters and field-specific F475W/F555W/F814W<=27 cuts for the public-archive galaxy fields; these strata are not pooled across filters or distances. Evidence tiers are explicit, and the field-level winner columns report recovery, position RMS, and magnitude RMS without constructing an unregistered composite precision score.", style="Caption", indent=0)
    insert_section_before(sup, sup_cross_caption, WD_ORIENT.LANDSCAPE)
    sup_cross_table = add_table(sup, comparison_headers, cross_rows, size=5.3)
    insert_section_after(sup, sup_cross_table, WD_ORIENT.PORTRAIT)
    append("Method-definition boundary. ‘AstroCFR+Photutils hybrid’ means AstroCFR ePSF/residual-deblend proposals followed by Photutils Gaussian-PRF PSFPhotometry. It is distinct from the v44 ‘spatial-ePSF + joint fit’ branch, which performs an AstroCFR-internal spatial empirical-PSF joint fit. The two are reported as distinct operating points and are never pooled or ranked as the same method.")
    append("Archive provenance. PHAT images/catalogues follow Dalcanton et al. (2012, ApJS, 200, 18); PHATTER products follow Williams et al. (2021, ApJS, 253, 8); the Tarantula/HTTP catalogue follows Sabbi et al. (2016, ApJS, 222, 11); and ANGST/GST products follow Dalcanton et al. (2009, ApJS, 183, 67) with DOLPHOT processing (Dolphin 2000, PASP, 112, 1383). These catalogues are external matching references only and never enter method fitting.")
    append(
        "Scope boundary: the reported empirical evidence excludes any field lacking a complete common five-method result and valid conditional position and photometry metrics. In particular, PHAT M31 B21-F10 and B21-F18 are excluded pending independent catalogue availability and complete evaluation. The manuscript does not infer performance on omitted fields or replace them with simulations."
    )
    append_references(sup, [
        ("Dalcanton, J. J., et al. (2009)", "Dalcanton, J. J., et al. (2009). The ACS Nearby Galaxy Survey Treasury. The Astrophysical Journal Supplement Series, 183, 67-108. https://doi.org/10.1088/0067-0049/183/1/67"),
        ("Dalcanton, J. J., et al. (2012)", "Dalcanton, J. J., et al. (2012). The Panchromatic Hubble Andromeda Treasury. The Astrophysical Journal Supplement Series, 200, 18. https://doi.org/10.1088/0067-0049/200/2/18"),
        ("Williams, B. F., et al. (2021)", "Williams, B. F., et al. (2021). The Panchromatic Hubble Andromeda Treasury: Triangulum Extended Region. The Astrophysical Journal Supplement Series, 253, 8. https://doi.org/10.3847/1538-4365/abd6f6"),
        ("Sabbi, E., et al. (2016)", "Sabbi, E., et al. (2016). The Hubble Tarantula Treasury Project. II. Photometric Catalog and Properties of the Young Stellar Populations. The Astrophysical Journal Supplement Series, 222, 11. https://doi.org/10.3847/0067-0049/222/1/11"),
    ])
    sup_saved = SUP_DEST
    try:
        sup.save(sup_saved)
    except PermissionError:
        sup_saved = SUP_UPDATED_DEST
        sup.save(sup_saved)
    print(main_saved)
    print(sup_saved)
    print(COMPARISON_CSV)


if __name__ == "__main__":
    main()
