#!/usr/bin/env python
"""Build v31: rename the manuscript system to AstroCFR and add context.

The repository is now published as AstroCFR. The reusable Python package
retains the historical ``src/wpdc`` import path for compatibility. This
revision also adds DOI-verified adjacent work and a density-adaptive routing
operating point derived from the registered fixed-scene artificial-star
summary. Version v30 is not modified.
"""
from __future__ import annotations

import json
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_manuscript_v30_submission_fixes import (
    append_text,
    bookmark,
    internal_link,
    rewrite_reference,
    set_cell,
    set_font,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v30_submission_fixes.docx"
DEST = ROOT / "supplementary" / "AstroCFR_Multimedia_Systems_SCI_manuscript_v32.docx"
ROUTING = ROOT / "results" / "hst_density_adaptive_routing" / "density_adaptive_branch_routing.json"
FIGURES = ROOT / "results" / "astrocfr_manuscript_figures"

MEDIA_REPLACEMENTS = {
    "word/media/image1.png": FIGURES / "fig1_astrocfr_architecture.png",
    "word/media/image9.png": FIGURES / "fig9_astrocfr_sextractor_comparison.png",
    "word/media/image14.png": FIGURES / "fig14_astrocfr_zero_shot.png",
    "word/media/image16.png": FIGURES / "fig16_astrocfr_hst_benchmark.png",
    "word/media/image18.png": FIGURES / "fig18_astrocfr_controlled_comparison.png",
    "word/media/image19.png": FIGURES / "fig19_astrocfr_artificial_recovery.png",
    "word/media/image20.png": FIGURES / "fig20_astrocfr_six_branch_comparison.png",
    "word/media/image21.png": FIGURES / "fig21_astrocfr_density_recovery.png",
    "word/media/image23.png": FIGURES / "fig23_astrocfr_failure_cases.png",
    "word/media/image24.png": FIGURES / "fig24_astrocfr_density_magnitude_recovery.png",
}


NEW_REFERENCES = [
    ("Yan, Y., Liu, C., Li, J., & Wang, F. (2026). Detection and classification of astronomical sources with "
     "Astro-RetinaNet in crowded stellar fields. The Astrophysical Journal Supplement Series, 283(2), 72. "
     "https://doi.org/10.3847/1538-4365/ae4a24"),
    ("Lai, J., Lam, M. I., Chen, J., Zhang, X., Tian, H., Chen, X., et al. (2026). Stellar density classification "
     "and regression for CSST multi-color imaging using deep learning. Publications of the Astronomical Society "
     "of the Pacific, 138(6), 064501. https://doi.org/10.1088/1538-3873/ae6d75"),
    ("Yan, Y., Wu, Y., Nie, J., Zhang, T., Liu, C., Ban, Z., et al. (2026). A robust geometric distortion solution "
     "for the Main Survey Camera of CSST. The Astronomical Journal, 171(5), 264. "
     "https://doi.org/10.3847/1538-3881/ae505f"),
    ("Nie, J., Wei, P., Cao, Z., Yan, Y., Liu, C., Tian, H., et al. (2025). Toward high-precision astrometry with "
     "CSST using multi-Gaussian fitting of PSF. Research in Astronomy and Astrophysics, 25(11), 115006. "
     "https://doi.org/10.1088/1674-4527/adfd22"),
    ("Shaw, R. A., Fotopoulou, S., Birkinshaw, M., Maddox, N., & Stewart, H. (2025). DRUID: Source detection and "
     "deblending in astronomical images with persistent homology. RAS Techniques and Instruments, 4, rzaf006. "
     "https://doi.org/10.1093/rasti/rzaf006"),
    ("Burke, C. J., Aleo, P. D., Chen, Y.-C., Liu, X., Peterson, J. R., Sembroski, G. H., et al. (2019). "
     "Deblending and classifying astronomical sources with Mask R-CNN deep learning. Monthly Notices of the "
     "Royal Astronomical Society, 490(3), 3952-3965. https://doi.org/10.1093/mnras/stz2845"),
]


def replace_text_nodes(paragraph, old: str, new: str) -> None:
    for node in paragraph._p.iter(qn("w:t")):
        if node.text:
            node.text = node.text.replace(old, new)


def all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def normal_paragraph(doc, indent: float = 0.35):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Inches(indent) if indent else None
    p.paragraph_format.line_spacing = 1.08
    return p


def heading(doc, text: str, style: str):
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        set_font(run, size=10.5, bold=True)
    return p


def move_before(nodes, target):
    cursor = target._p.getprevious()
    for node in nodes:
        node.getparent().remove(node)
        cursor.addnext(node)
        cursor = node


def border(cell, edge, val="single", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    element = borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), val)
    element.set(qn("w:sz"), size)
    element.set(qn("w:color"), "000000")


def three_line_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            border(cell, "top", "single" if row_i == 0 else "nil", "14" if row_i == 0 else "0")
            border(cell, "bottom", "single" if row_i in (0, len(table.rows) - 1) else "nil",
                   "8" if row_i == 0 else ("14" if row_i == len(table.rows) - 1 else "0"))
            border(cell, "left", "nil", "0")
            border(cell, "right", "nil", "0")


def append_related_work(doc):
    target = next(p for p in doc.paragraphs if p.text.startswith("We present AstroCFR as"))
    p = normal_paragraph(doc)
    append_text(p, "Recent adjacent work includes blended-source detection with Astro-RetinaNet (")
    internal_link(p, "Yan et al., 2026", "ref_42")
    append_text(p, "), density-regime classification for CSST processing (")
    internal_link(p, "Lai et al., 2026", "ref_43")
    append_text(p, "), the WPDC-2P geometric-distortion correction method (")
    internal_link(p, "Yan et al., 2026", "ref_44")
    append_text(p, "), multi-Gaussian CSST PSF astrometry (")
    internal_link(p, "Nie et al., 2025", "ref_45")
    append_text(p, "), persistent-homology deblending (")
    internal_link(p, "Shaw et al., 2025", "ref_46")
    append_text(p, "), and Mask R-CNN source deblending (")
    internal_link(p, "Burke et al., 2019", "ref_47")
    append_text(p, "). These studies address detector-level recovery, density gating, geometric correction, PSF measurement, or segmentation rather than the complete AstroCFR deployment protocol. Their reported metrics therefore provide context rather than a pooled SOTA leaderboard. AstroCFR is unrelated to WPDC-2P; the latter denotes Weighted Polynomial Distortion Correction in 2-Phase for CSST astrometry.")
    move_before([p._p], target)


def extend_literature_table(doc):
    table = doc.tables[11]
    rows = [
        ("Yan et al. (2026), Astro-RetinaNet", "CSST-simulated NGC 2298 and HST M31",
         "At i=25, 82% crowded and 95% noncrowded recovery; M31 F814W=27 candidate count exceeds Photutils/SExtractor",
         "Closest detector-level study, but images, limiting magnitudes, association rule, and outputs differ"),
        ("Lai et al. (2026)", "CSST multi-color density simulations",
         "98.83% density-category accuracy; 0.0824 dex bright-star-count MAE",
         "Density gate rather than source-recovery or measurement baseline"),
        ("Yan et al. (2026), WPDC-2P", "18 CSST detectors and BASS",
         "0.013-0.107 pixel simulated scatter; 5.494/9.981 mas BASS positional scatter",
         "Astrometric geometric-distortion module only; acronym explicitly distinguished from AstroCFR"),
        ("Nie et al. (2025)", "CSST simulated sparse and crowded fields",
         "Sub-mas sparse-field centering; improved faint-star astrometry versus SExtractor/DOLPHOT",
         "PSF/astrometry context; no common AstroCFR candidate-recovery protocol"),
        ("Shaw et al. (2025), DRUID", "LoTSS and KiDS extended-source images",
         "Persistent-homology detection and nested-source segmentation",
         "Different source morphologies and no crowded stellar photometry benchmark"),
        ("Burke et al. (2019)", "PhoSim simulated astronomical scenes",
         "Mask R-CNN detection, classification, and deblending",
         "Learning-based deblending context; different simulator and evaluation task"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell(cell, value, size=7.6)


def append_routing_protocol(doc):
    results = next(p for p in doc.paragraphs if p.text.strip() == "5 Results")
    h = heading(doc, "4.4 Density-adaptive branch-routing protocol", "Heading 2")
    p = normal_paragraph(doc)
    append_text(p, "Motivated by the heterogeneous-density processing problem, we evaluate a deterministic deployment router on the registered NGC 6752 fixed-scene artificial-star aggregates. High-density strata use AstroCFR spatial-ePSF plus joint fitting, whereas low-density strata use Photutils PSFPhotometry. The router uses the predefined injection-density stratum, fits no parameters, and receives the same scenes as the fixed-branch baselines. Wilson 95% intervals are calculated from pooled recovery counts. Runtime is the injection-weighted combination of the canonical five-repeat CPU medians. Because the aggregate contains proposal-recovery counts rather than per-candidate residuals, purity and measurement RMS are not inferred for this auxiliary experiment.")
    move_before([h._p, p._p], results)


def append_routing_result(doc):
    data = json.loads(ROUTING.read_text(encoding="utf-8"))
    by_policy = {x["policy"]: x for x in data["policies"]}
    discussion = next(p for p in doc.paragraphs if p.text.strip() == "6 Discussion")
    h = heading(doc, "5.22 Density-adaptive deployment routing", "Heading 1")
    p = normal_paragraph(doc)
    append_text(p, "The density-adaptive router defines an intermediate recovery-cost operating point. On 800 identical artificial-star injections it recovers 400 sources (50.00%, 95% CI: 46.54-53.46%) at an injection-weighted 20.44 s/MPix. This is 10.75 percentage points above Photutils-only recovery, while the always-on spatial-ePSF branch recovers 60.38% at 32.23 s/MPix. Thus routing saves 36.6% of the spatial-ePSF runtime at a 10.38-point recovery cost; it does not dominate either endpoint. The result is summarized in ")
    internal_link(p, "Table 25", "table_25")
    append_text(p, ".")
    caption = doc.add_paragraph("Table 25. Density-adaptive branch routing on the identical NGC 6752 expanded artificial-star scenes. Recovery intervals are Wilson 95% intervals over 800 injections. Runtime is an injection-weighted CPU operating-point estimate; purity and measurement RMS are not available from this proposal-recovery aggregate.", style="Caption")
    bookmark(caption, "table_25", 3025)
    for run in caption.runs:
        set_font(run, size=9.0)
    table = doc.add_table(rows=1, cols=5)
    headers = ["Policy", "High-density branch", "Low-density branch", "Recovery (95% CI)", "s/MPix"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, size=8.0)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    labels = {
        "photutils_only": ("Photutils-only", "Photutils", "Photutils"),
        "epsf_only": ("AstroCFR ePSF-only", "ePSF+deblend", "ePSF+deblend"),
        "spatial_epsf_only": ("AstroCFR spatial-only", "spatial-ePSF+joint", "spatial-ePSF+joint"),
        "density_adaptive": ("Density-adaptive", "spatial-ePSF+joint", "Photutils"),
    }
    for key in ("photutils_only", "epsf_only", "spatial_epsf_only", "density_adaptive"):
        r = by_policy[key]
        label, high, low = labels[key]
        lo, hi = (100 * x for x in r["recovery_ci95"])
        values = [label, high, low, f"{100*r['recovery']:.2f}% [{lo:.2f}-{hi:.2f}%]", f"{r['runtime_s_per_mpix_weighted']:.2f}"]
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell(cell, value, size=7.8)
    three_line_table(table)
    move_before([h._p, p._p, caption._p, table._element], discussion)


def append_references(doc):
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "References")
    existing = [p for p in paragraphs[start + 1:] if p.text.strip()]
    if len(existing) != 41:
        raise RuntimeError(f"Expected 41 existing references, found {len(existing)}")
    for number, text in enumerate(NEW_REFERENCES, 42):
        paragraph = doc.add_paragraph(style="Normal")
        rewrite_reference(paragraph, text, number)


def replace_media(source: Path, destination: Path) -> None:
    """Replace selected raster assets while preserving the DOCX package."""
    for path in MEDIA_REPLACEMENTS.values():
        if not path.exists():
            raise FileNotFoundError(path)
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=destination.parent) as handle:
        temporary = Path(handle.name)
    try:
        with ZipFile(source) as inp, ZipFile(temporary, "w", ZIP_DEFLATED) as out:
            for info in inp.infolist():
                payload = (MEDIA_REPLACEMENTS[info.filename].read_bytes()
                           if info.filename in MEDIA_REPLACEMENTS else inp.read(info.filename))
                if info.filename.endswith(".xml") or info.filename.endswith(".rels"):
                    payload = payload.replace(
                        b"github.com/zxl1999/WPDC", b"github.com/zxl1999/AstroCFR"
                    )
                out.writestr(info, payload)
        temporary.replace(destination)
    finally:
        temporary.unlink(missing_ok=True)


def main():
    doc = Document(SOURCE)

    # Rename visible manuscript-facing identifiers. Lowercase package paths
    # such as src/wpdc remain historical compatibility paths.
    for paragraph in all_paragraphs(doc):
        replace_text_nodes(paragraph, "WPDC", "AstroCFR")

    title = doc.paragraphs[0]
    title.clear()
    set_font(title.add_run("AstroCFR: An Astronomical Crowded-Field Recovery System for Large-Scale Survey Images"), size=14, bold=True)
    abstract = next(p for p in doc.paragraphs if p.text.startswith("We present AstroCFR,"))
    replace_text_nodes(abstract, "We present AstroCFR, a", "We present AstroCFR (Astronomical Crowded-Field Recovery System), a")
    replace_text_nodes(
        abstract,
        "These results support a modular system with explicit operating-point trade-offs, not a universally superior method.",
        "A fixed density-adaptive router provides an intermediate Pareto point: 50.00% artificial-star recovery at 20.44 s/MPix, versus 39.25% at 8.64 s/MPix for Photutils-only and 60.38% at 32.23 s/MPix for the always-on spatial-ePSF branch. These results support a modular system with explicit operating-point trade-offs, not a universally superior method.",
    )

    code = next(p for p in doc.paragraphs if p.text.startswith("Code availability:"))
    replace_text_nodes(code, "tag v1.0.0", "tag v1.1.1")
    append_text(code, " The Python package retains the historical src/wpdc import path for backward compatibility; the repository and manuscript-facing system name is AstroCFR.")

    append_related_work(doc)
    extend_literature_table(doc)
    append_routing_protocol(doc)
    append_routing_result(doc)
    append_references(doc)

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=DEST.parent) as handle:
        intermediate = Path(handle.name)
    try:
        doc.save(intermediate)
        replace_media(intermediate, DEST)
    finally:
        intermediate.unlink(missing_ok=True)
    print(DEST)


if __name__ == "__main__":
    main()
