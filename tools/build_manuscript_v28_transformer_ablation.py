#!/usr/bin/env python
"""Add the controlled Transformer classifier ablation to manuscript v27."""
from __future__ import annotations

import json
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v27_diagnostics.docx"
DEST = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v28_transformer_ablation.docx"
RESULT = ROOT / "results" / "transformer_candidate_ablation" / "transformer_candidate_ablation.json"


def style_run(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size); run.font.bold = bold; run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text, style="Normal", indent=0.35):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Inches(indent) if indent else None
    p.paragraph_format.line_spacing = 1.08
    style_run(p.add_run(text))
    return p


def set_three_line(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders"); tbl_pr.append(borders)
    for edge in ("top", "bottom"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None: node = OxmlElement(f"w:{edge}"); borders.append(node)
        node.set(qn("w:val"), "single"); node.set(qn("w:sz"), "12"); node.set(qn("w:space"), "0"); node.set(qn("w:color"), "000000")
    inside = borders.find(qn("w:insideH"))
    if inside is None: inside = OxmlElement("w:insideH"); borders.append(inside)
    inside.set(qn("w:val"), "single"); inside.set(qn("w:sz"), "6"); inside.set(qn("w:space"), "0"); inside.set(qn("w:color"), "000000")
    for edge in ("left", "right", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None: node = OxmlElement(f"w:{edge}"); borders.append(node)
        node.set(qn("w:val"), "nil")


def main():
    payload = json.loads(RESULT.read_text(encoding="utf-8"))
    doc = Document(SOURCE)
    anchor = next(p for p in doc.paragraphs if p.text.strip() == "6 Discussion")
    heading = add_para(doc, "5.21 Classifier architecture ablation", style="Heading 1", indent=0)
    text = ("To test whether a newer attention-based classifier is necessary, we performed a controlled simulation-domain ablation. "
            "Candidate generation, the association rule, the four CSST-like chips, and the train/validation/test partition were fixed. "
            "Only the candidate-quality classifier was changed: the established 17-feature RandomForest, the original WPDC StarBogusNet, "
            "and a lightweight patch Transformer using the same 25 × 25 cutout and 17 handcrafted features. "
            "The Transformer uses 25 non-overlapping 5 × 5 patches, 64-dimensional tokens, two four-head encoder layers, and a small feature-fusion head.")
    p1 = add_para(doc, text)
    p2 = add_para(doc, "The results are reported in Table 24. The Transformer improves recall and F1 over the original CNN in this split, but the RF branch remains the strongest operating point under the pre-registered threshold rule. The attention model also requires approximately twice the training time of the CNN and is not a replacement for the WPDC system. This negative result is useful: it shows that the paper's contribution is the deployable multi-stage pipeline and its explicit operating points, rather than an unsupported claim that a Transformer is universally superior.")
    cap = add_para(doc, "Table 24. Controlled simulation-domain candidate-classifier ablation. Values are measured on the untouched 20% test partition; thresholds are selected on the validation partition to retain at least 90% validation recall. GPU figures apply to the CNN and Transformer classifier stages only.", style="Caption", indent=0)
    table = doc.add_table(rows=1, cols=9)
    table.autofit = True
    headers = ["Model", "Params", "Recall", "Precision", "F1", "AUROC", "AUPRC", "Train/s", "Inference ms/candidate"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        for run in cell.paragraphs[0].runs: style_run(run, bold=True)
    for name, label in (("RandomForest", "RandomForest"), ("Original_CNN", "Original WPDC CNN"), ("Patch_Transformer", "Patch Transformer")):
        e = payload["models"][name]; m = e["test"]
        vals = [label, "--" if e["parameters"] is None else f"{e['parameters']:,}", f"{m['recall']:.4f}", f"{m['precision']:.4f}", f"{m['f1']:.4f}", f"{m['auroc']:.4f}", f"{m['auprc']:.4f}", f"{e['training_s']:.2f}", f"{e['inference_ms_per_candidate']:.4f}"]
        cells = table.add_row().cells
        for cell, value in zip(cells, vals):
            cell.text = value
            for run in cell.paragraphs[0].runs: style_run(run)
    set_three_line(table)
    nodes = [heading._p, p1._p, p2._p, cap._p, table._tbl]
    cursor = anchor._p.getprevious()
    for node in nodes: node.getparent().remove(node)
    for node in nodes: cursor.addnext(node); cursor = node
    doc.save(DEST)
    print(DEST)


if __name__ == "__main__": main()
