#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Order the bibliography by first in-text appearance and apply hanging indent."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
SOURCE=ROOT/'WPDC_Multimedia_Systems_SCI_manuscript_v18_crossrefs.docx'
DEST=ROOT/'WPDC_Multimedia_Systems_SCI_manuscript_v19_reference_order.docx'
DROP=('Foreman-Mackey,','McKinney,','Shorten,','Turlach,')
EARLY=('Anderson, J., et al. (2008). The ACS Survey','Libralato, M., et al. (2024). Euclid:','Salaris, M., et al. (2024). The HST Large','Bradley, L., Sipocz, B., Robitaille, T., et al.')

def set_font(p):
    for r in p.runs:
        r.font.name='Times New Roman';r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体');r.font.size=Pt(10.5);r.font.color.rgb=RGBColor(0,0,0)

def pattern(ref):
    pairs=[
      ('Anderson, J., & King','Anderson & King, 2006'),('Anderson, J., et al.','Anderson et al., 2008'),
      ('Bellini,','Bellini et al., 2011'),('Chambers,','Chambers et al., 2016'),('Dey,','Dey et al., 2019'),
      ('Gaia Collaboration, et al. (2016)','Gaia Collaboration et al., 2016'),('Gaia Collaboration, et al. (2023)','Gaia Collaboration et al., 2016, 2023'),
      ('Astropy Collaboration,','Astropy Collaboration, 2022'),('Harris,','Harris et al., 2020'),('Hunter,','Hunter, 2007'),('Virtanen,','Virtanen et al., 2020'),('Barbary,','Barbary, 2016'),
      ('Bradley,','Bradley et al., 2024'),('Pedregosa,','Pedregosa et al., 2011'),('Goodfellow,','Goodfellow et al., 2016'),('He,','He et al., 2016'),
      ('Shi,','Shi et al., 2024'),('Long, M., Xin','Long et al., 2025'),('Wainer,','Wainer et al., 2025'),('Espinosa,','Espinosa et al., 2025'),('Centofanti,','Centofanti et al., 2026'),('Libralato,','Libralato et al., 2024'),('Salaris,','Salaris et al., 2024'),
      ('Ben-David,','Ben-David et al., 2010'),('Pan,','Pan & Yang, 2010'),('Ganin,','Ganin et al., 2016'),('Long, M., Cao','Long et al., 2015'),('Breiman,','Breiman, 2001'),('Chen,','Chen & Guestrin, 2016'),
      ('Stetson,','Stetson, 1987'),('Bertin,','Bertin & Arnouts, 1996'),('Dolphin,','Dolphin, 2000'),('Schlafly,','Schlafly et al., 2018'),('Melchior,','Melchior et al., 2018'),('Wang,','Wang et al., 2024'),
      ('Zhang, Y.,','Zhang et al., 2024a'),('Zhang, S.,','Zhang et al., 2024b')]
    for prefix,needle in pairs:
        if ref.startswith(prefix):return needle
    raise KeyError(ref[:80])

def main():
    doc=Document(SOURCE)
    # Cite the retained implementation references in the introduction.
    for p in doc.paragraphs:
        if p.text.startswith('The positioning also follows established measurement'):
            p.text=('The positioning also follows established measurement and deployment practice: empirical ACS ePSFs and crowded-field catalogues (Anderson & King, 2006; Anderson et al., 2008; Bellini et al., 2011), survey-scale reference systems (Chambers et al., 2016; Dey et al., 2019; Gaia Collaboration et al., 2016, 2023), and reproducible scientific-computing components (Astropy Collaboration, 2022; Harris et al., 2020; Hunter, 2007; Virtanen et al., 2020; Barbary, 2016; Bradley et al., 2024; Pedregosa et al., 2011). Deep-learning foundations and residual networks provide the image-model context (Goodfellow et al., 2016; He et al., 2016). Recent context includes CSST photometric preparation (Shi et al., 2024), deep-feature point-source detection (Long et al., 2025), Rubin crowded-field photometry (Wainer et al., 2025), PSF-fitting uncertainty analysis (Espinosa et al., 2025), image-based PSF recovery (Centofanti et al., 2026), and recent HST/Euclid cluster studies (Libralato et al., 2024; Salaris et al., 2024). Transfer-learning claims are framed using domain-shift theory and adaptation methods (Ben-David et al., 2010; Pan & Yang, 2010; Ganin et al., 2016; Long et al., 2015), while the classifier baselines are standard Random Forest and gradient-boosting references (Breiman, 2001; Chen & Guestrin, 2016).')
            set_font(p)
    idx=next(i for i,p in enumerate(doc.paragraphs) if p.text.strip()=='References')
    heading=doc.paragraphs[idx]
    refs=[p.text.strip() for p in doc.paragraphs[idx+1:] if p.text.strip()]
    # Bradley is unique but was previously misplaced before the heading.
    early=[]
    for p in list(doc.paragraphs[:idx]):
        if p.text.strip().startswith(EARLY):
            if p.text.strip().startswith('Bradley,'):early.append(p.text.strip())
            p._element.getparent().remove(p._element)
    refs=[r for r in refs+early if not r.startswith(DROP)]
    # Deduplicate exact entries while preserving their content.
    unique=[]
    for r in refs:
        if r not in unique:unique.append(r)
    # Removing misplaced pre-heading entries shifts paragraph indices.
    idx=next(i for i,p in enumerate(doc.paragraphs) if p.text.strip()=='References')
    heading=doc.paragraphs[idx]
    body='\n'.join(p.text for p in doc.paragraphs[:idx])
    ranked=[]
    for r in unique:
        needle=pattern(r);pos=body.find(needle)
        if pos<0:raise RuntimeError(f'Uncited reference: {needle}')
        ranked.append((pos,r))
    ranked.sort(key=lambda x:x[0]);ordered=[r for _,r in ranked]
    # Remove old bibliography paragraphs after the heading.
    for p in list(doc.paragraphs[idx+1:]):p._element.getparent().remove(p._element)
    cursor=heading
    for r in ordered:
        p=doc.add_paragraph(r,style='Normal');p.paragraph_format.left_indent=Inches(0.20);p.paragraph_format.first_line_indent=Inches(-0.20);p.paragraph_format.line_spacing=1.0;set_font(p)
        cursor._p.addnext(p._p);cursor=p
    doc.save(DEST);print(DEST,'references',len(ordered))
if __name__=='__main__':main()
