#!/usr/bin/env python
"""Strengthen the defensible methodological contribution without inventing new results."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import internal_link, set_font

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v39_final.docx"
DEST = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v40_methodology.docx"
SUP_SOURCE = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v39.docx"
SUP_DEST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v40.docx"


def find(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def replace(p, value, indent=.35):
    p.clear(); p.style = "Normal"
    p.paragraph_format.first_line_indent = Inches(indent) if indent else None
    p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(value), size=10.5)


def insert_after(p, value, bold=False, indent=.35):
    q = p._parent.add_paragraph(style="Normal")
    q._p.getparent().remove(q._p); p._p.addnext(q._p)
    replace(q, value, indent)
    for r in q.runs:
        r.font.bold = bold
    return q


def bookmark(p, name, ident):
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(ident)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(ident))
    p._p.insert(0, start); p._p.append(end)


def add_reference(doc, value, anchor, ident):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(.28)
    p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(value), size=10.5)
    bookmark(p, anchor, ident)
    return p


def main():
    doc = Document(SOURCE)
    sup = Document(SUP_SOURCE)

    # State a method contribution that is actually supported by the experiments: common-protocol,
    # confidence-aware multi-objective branch selection—not a claim of a new primitive detector.
    p = find(doc, "Recent detector-level work on blends")
    replace(p,
        "Recent detector-level work on blends, density classification, distortion, PSF fitting, and segmentation "
        "motivates individual modules. AstroCFR's methodological contribution is a calibration-aware decision "
        "protocol: candidate-recovery branches are evaluated under one image, association rule, spatial split, "
        "and injected scene set, then compared as uncertainty-qualified recovery–precision–cost operating points. "
        "This contribution is therefore a reproducible selection procedure, not a claim of a new primitive detector.")
    q = insert_after(p,
        "For ACS/WFC calibration context, the manuscript also follows the effective-PSF and detector-photometry "
        "foundations of ", indent=.35)
    # Rebuild this paragraph so citations remain genuine Word hyperlinks.
    q.clear(); q.paragraph_format.first_line_indent = Inches(.35); q.paragraph_format.line_spacing = 1.08
    set_font(q.add_run("For ACS/WFC calibration context, the manuscript also follows the effective-PSF and detector-photometry foundations of "), size=10.5)
    internal_link(q, "Anderson & King (2000; Ref. 48)", "ref_48")
    set_font(q.add_run(" and "), size=10.5)
    internal_link(q, "Sirianni et al. (2005; Ref. 49)", "ref_49")
    set_font(q.add_run("."), size=10.5)

    # Make the selection rule explicit in the problem-setting section.
    scope = find(doc, "AstroCFR is treated as a modular astronomical processing chain")
    h = insert_after(scope, "2.3 Calibration-aware operating-point protocol", bold=True, indent=0)
    protocol = insert_after(h,
        "For a branch b, we report the decision vector z_b = (1−C_b, 1−R_b, Epos_b, Emag_b, T_b, M_b), where "
        "C is held-out completeness, R is dense-subset recovery, Epos and Emag are held-out positional and "
        "photometric RMS, and T and M are runtime and peak memory. A branch is Pareto-admissible only when no "
        "other evaluated branch is no worse in every reported component and better in at least one, under the "
        "same image, matching rule, spatial test split, and injected scenes. This definition prevents a detector "
        "from appearing preferable merely because it was evaluated under a different threshold or truth protocol.")
    insert_after(protocol,
        "Release selection is then science-constrained rather than globally ranked: low-latency screening selects "
        "DAOStarFinder or AstroCFR-RF, precision-constrained measurement selects Photutils, and recovery-constrained "
        "dense-field analysis selects an ePSF branch only when its Wilson interval and computational budget meet "
        "the declared science requirement. Table 9 reports these common-protocol operating points.")

    # Turn the observed learning curve into a concrete, falsifiable calibration-sufficiency rule.
    budget = find(doc, "Thus the defensible deployment statement is not that a fixed number of real images")
    insert_after(budget,
        "The learning curves support a practical calibration-sufficiency rule. Begin with one spatially isolated "
        "tile and increase the budget only until the calibration set contains at least 30 positive and 8 conservative "
        "negative candidates and the five-selection recall half-width is at most 0.05. NGC 6397 (38/8 labels; "
        "half-width 0.036) and NGC 6752 (101/15; 0.021) satisfy this rule with one tile. NGC 1851 fails the "
        "stability condition at one and three tiles and first satisfies it at six tiles (364/19; 0.002). This is a "
        "data-backed stopping rule for the reported protocol, not a universal label-budget law.")

    # The CNN is optional; prevent an untested cutout-size limitation from being mistaken for a core-method gap.
    cnn_lim = find(doc, "Finally, the present CNN cutout size is fixed at 25 x 25 pixels")
    replace(cnn_lim,
        "The 25 × 25-pixel CNN cutout is retained only as an optional simulation-domain baseline. AstroCFR's "
        "primary HST claims do not depend on it, and the manuscript does not claim its scale is optimal. A "
        "multi-scale CNN comparison would be a separate learning-model study rather than evidence required for "
        "the candidate-recovery and measurement operating points reported here.")

    # Concrete multi-exposure extension route, without presenting it as completed work.
    future = find(doc, "The next improvement direction is full-field joint PSF fitting")
    replace(future,
        "The next evaluation should use multiple HST/ACS FLC exposures rather than one stacked image. The extension "
        "will share source coordinates across exposures, use exposure-specific spatial ePSFs, backgrounds, and noise "
        "weights, and optimize their joint pixel likelihood over local neighbour groups. It should then compare the "
        "same sources against DOLPHOT or HST1PASS under matched exposure lists and artificial-star scenes. This is "
        "the appropriate route to a multi-exposure crowded-field claim; it is deliberately outside the present result.")

    # Append two directly relevant, verified core references.  Their first citations are immediately above.
    add_reference(doc,
        "Anderson, J., & King, I. R. (2000). Toward high-precision astrometry with WFPC2. I. Deriving an accurate point-spread function. Publications of the Astronomical Society of the Pacific, 112, 1360–1382. https://doi.org/10.1086/316632",
        "ref_48", 948)
    add_reference(doc,
        "Sirianni, M., Jee, M. J., Benítez, N., Blakeslee, J. P., Martel, A. R., et al. (2005). The photometric performance and calibration of the Hubble Space Telescope Advanced Camera for Surveys. Publications of the Astronomical Society of the Pacific, 117, 1049–1112. https://doi.org/10.1086/444553",
        "ref_49", 949)

    doc.save(DEST)
    sup.save(SUP_DEST)
    print(DEST)
    print(SUP_DEST)


if __name__ == "__main__":
    main()
