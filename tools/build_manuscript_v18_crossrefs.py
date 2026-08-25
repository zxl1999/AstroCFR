#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Add Word internal cross-references for every table and figure."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[1]
SOURCE=ROOT/'WPDC_Multimedia_Systems_SCI_manuscript_v17_sci_format.docx'
DEST=ROOT/'WPDC_Multimedia_Systems_SCI_manuscript_v18_crossrefs.docx'

def bookmark(paragraph,name,ident):
    start=OxmlElement('w:bookmarkStart');start.set(qn('w:id'),str(ident));start.set(qn('w:name'),name)
    end=OxmlElement('w:bookmarkEnd');end.set(qn('w:id'),str(ident))
    paragraph._p.insert(0,start);paragraph._p.append(end)

def hyperlink(paragraph,label,anchor):
    h=OxmlElement('w:hyperlink');h.set(qn('w:anchor'),anchor);h.set(qn('w:history'),'1')
    r=OxmlElement('w:r');rp=OxmlElement('w:rPr');color=OxmlElement('w:color');color.set(qn('w:val'),'0563C1');u=OxmlElement('w:u');u.set(qn('w:val'),'single');rp.append(color);rp.append(u);r.append(rp)
    t=OxmlElement('w:t');t.text=label;r.append(t);h.append(r);paragraph._p.append(h)

def reference_paragraph(doc,caption,kind,num,ident):
    p=doc.add_paragraph();p.style='Normal';p.paragraph_format.first_line_indent=Inches(0.35);p.paragraph_format.line_spacing=1.08
    p.add_run('The corresponding ' + ('quantitative results are summarized in ' if kind=='Table' else 'visualization is shown in '))
    hyperlink(p,f'{kind} {num}',ident);p.add_run('.')
    for r in p.runs:r.font.name='Times New Roman';r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体');r.font.size=Pt(10.5);r.font.color.rgb=RGBColor(0,0,0)
    # Put the citation before an immediately preceding picture paragraph.
    target=caption._p.getprevious()
    if kind=='Fig.' and target is not None and target.tag.endswith('}p') and target.xpath('.//w:drawing'):
        target.addprevious(p._p)
    else:caption._p.addprevious(p._p)

def main():
    doc=Document(SOURCE);ident=100
    for p in doc.paragraphs:
        m=re.match(r'^Fig\.\s*(\d+)\.?\s+(.*)$',p.text.strip(),re.I)
        if m:
            p.text=f'Fig. {m.group(1)}. {m.group(2)}';p.style='Caption'
    # Work on a stable snapshot because new citation paragraphs are inserted.
    captions=[p for p in doc.paragraphs if p.style.name=='Caption' and re.match(r'^(Table|Fig\.)\s+\d+\.',p.text.strip())]
    for p in captions:
        m=re.match(r'^(Table|Fig\.)\s+(\d+)\.',p.text.strip());kind,num=m.group(1),int(m.group(2));name=('table_' if kind=='Table' else 'figure_')+str(num)
        bookmark(p,name,ident);reference_paragraph(doc,p,kind,num,name);ident+=1
    doc.save(DEST);print(DEST,'crossrefs',len(captions))
if __name__=='__main__':main()
