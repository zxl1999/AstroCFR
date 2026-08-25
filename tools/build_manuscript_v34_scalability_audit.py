#!/usr/bin/env python
"""Add the registered scalability and density-gate audit to the v33 pair."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import set_font
from build_manuscript_v33_slim import format_three_line


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "AstroCFR_Multimedia_Systems_SCI_manuscript_v33_slim.docx"
DEST = ROOT / "supplementary" / "AstroCFR_Multimedia_Systems_SCI_manuscript_v34_scalability.docx"
SUP_SOURCE = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v33.docx"
SUP_DEST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v34.docx"
SCALING = ROOT / "results" / "hst_tile_parallel_scaling" / "tile_parallel_scaling.json"
GATE_DIR = ROOT / "results" / "hst_automatic_density_gate"


def insert_after(paragraph, text: str, *, style: str = "Normal", size: float = 10.5):
    new = paragraph._parent.add_paragraph(style=style)
    new._p.getparent().remove(new._p)
    paragraph._p.addnext(new._p)
    run = new.add_run(text)
    set_font(run, size=size)
    new.paragraph_format.line_spacing = 1.08
    if style == "Normal":
        new.paragraph_format.first_line_indent = Inches(.35)
    return new


def replace_text(paragraph, old: str, new: str) -> None:
    for node in paragraph._p.iter(qn("w:t")):
        if node.text:
            node.text = node.text.replace(old, new)


def paragraph_start(doc, prefix: str):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def gate_results() -> dict:
    values = {}
    for cluster in ("ngc6397", "ngc6752", "ngc1851"):
        payload = json.loads((GATE_DIR / f"{cluster}_automatic_density_gate_diagnostic.json").read_text(encoding="utf-8"))
        values[cluster] = payload["lightweight_image_only_gate"]["held_out_artificial_protocol_positions"]
    return values


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=8, bold=(row_index == 0))
    format_three_line(table)
    return table


def add_main() -> None:
    doc = Document(SOURCE)
    gates = gate_results()
    gate_text = ", ".join(f"{cluster.upper()} {value['sensitivity']:.3f}" for cluster, value in gates.items())
    abstract = next(p for p in doc.paragraphs if p.text.startswith("We present AstroCFR"))
    replace_text(abstract,
                 "a fixed density-adaptive router achieves 50.00% recovery",
                 "a fixed density-stratified deployment router (using registered artificial-star strata, not an image-trained gate) achieves 50.00% recovery")
    router = paragraph_start(doc, "The density-adaptive router defines an intermediate")
    insert_after(router,
                 "To test whether this policy could be automated without reading the evaluation catalogue, we ran a separate image-only density-gate audit. A shallow RandomForest using multi-scale candidate counts and local image texture was calibrated only in the left x < 400-pixel strip and evaluated in the right x >= 800-pixel strip. On held-out artificial-star query positions, high-density sensitivity was " + gate_text + ". Because this cross-field behavior is unstable, the headline router remains the registered density-stratified deployment policy rather than an automatically gated end-to-end claim.")
    resource = paragraph_start(doc, "The controlled HST evaluation now reports wall-clock")
    insert_after(resource,
                 "Tile-level parallelism was also measured for the AstroCFR ePSF plus residual-deblend branch on the NGC 6752 1200 x 1200 crop, split into four identical 600 x 600 tiles. Including worker initialization, median wall time was 64.61 s with one process, 31.69 s with two (2.04x speed-up), and 21.42 s with four (3.02x speed-up); aggregate peak RSS increased from 786 MB to 1,530 MB and 3,025 MB, while all configurations returned the same 7,096 candidates. This establishes tile-level CPU parallelism on one workstation, not TB/PB-scale throughput.")
    # Make the limitation explicit even if a journal extracts only Discussion.
    limitations = paragraph_start(doc, "The primary development experiments use CSST-like simulations")
    insert_after(limitations,
                 "The image-only density-gate audit is likewise a negative control rather than a claimed contribution: the proxy was useful on some held-out catalogue positions but did not provide stable high-density sensitivity across all three fields. We therefore retain the simpler, pre-registered density strata for the reported Pareto operating point.")
    for table in doc.tables:
        format_three_line(table)
    doc.save(DEST)


def add_supplement() -> None:
    doc = Document(SUP_SOURCE)
    gates = gate_results()
    h = doc.add_paragraph("S3 Image-only density-gate audit", style="Heading 1")
    for run in h.runs: set_font(run, size=12, bold=True)
    p = doc.add_paragraph("The registered density-adaptive result in the main text uses artificial-star density strata. To assess whether those strata can be inferred from the image alone, we calibrated a shallow RandomForest on multi-scale candidate counts and local pixel texture in x < 400 pixels and evaluated x >= 800 pixels. No catalogue labels enter inference. The cross-field instability is why the main text does not call the router an automatic density classifier.")
    for run in p.runs: set_font(run, size=10.5)
    cap = doc.add_paragraph("Table S3. Image-only density-gate audit on spatially held-out artificial-protocol positions. These are gate-classification diagnostics, not source-recovery or photometric metrics.", style="Caption")
    for run in cap.runs: set_font(run, size=9)
    add_table(doc, ["Field", "n", "High-density sensitivity", "Specificity", "Balanced accuracy"],
              [[cluster.upper(), str(value["n"]), f"{value['sensitivity']:.3f}",
                f"{value['specificity']:.3f}", f"{value['balanced_accuracy']:.3f}"]
               for cluster, value in gates.items()])

    scaling = json.loads(SCALING.read_text(encoding="utf-8"))["summary"]
    h2 = doc.add_paragraph("S4 Tile-level CPU parallel scaling", style="Heading 1")
    for run in h2.runs: set_font(run, size=12, bold=True)
    p = doc.add_paragraph("The ePSF plus residual-deblend branch was applied independently to four 600 x 600 tiles from the NGC 6752 central crop. Timings include worker start-up and per-worker image/PSF initialization. Candidate totals were identical across configurations. These values establish tile-level workstation scaling only.")
    for run in p.runs: set_font(run, size=10.5)
    cap = doc.add_paragraph("Table S4. Tile-level CPU scaling and aggregate process RSS for AstroCFR ePSF+deblend.", style="Caption")
    for run in cap.runs: set_font(run, size=9)
    add_table(doc, ["Processes", "Wall time (s)", "Speed-up", "Efficiency", "Throughput (MPix/s)", "Peak RSS (MB)"],
              [[str(row["workers"]), f"{row['wall_s_median']:.2f}", f"{row['speedup_vs_1']:.2f}x",
                f"{row['parallel_efficiency']:.3f}", f"{row['throughput_mpix_s_median']:.4f}",
                f"{row['aggregate_peak_rss_mb_median']:.0f}"] for row in scaling])
    doc.save(SUP_DEST)


if __name__ == "__main__":
    add_main(); add_supplement(); print(DEST); print(SUP_DEST)
