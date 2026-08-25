#!/usr/bin/env python
"""Produce a focused submission manuscript from the fully audited v35 report.

The v35 document remains the complete technical record.  This builder removes
diagnostic and superseded results from the journal-facing narrative, archives
them in the supplement, and renumbers all surviving figures/tables and links.
"""
from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

from build_manuscript_v30_submission_fixes import set_font
from build_manuscript_v33_slim import format_three_line


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v35_reframed.docx"
DEST = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v36_focused.docx"
SUP_SOURCE = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v35.docx"
SUP_DEST = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v36.docx"
MEDIA_DIR = ROOT / "results" / "astrocfr_supplementary_figures" / "v36_moved"


TABLES_TO_SUPPLEMENT = {2, 3, 4, 5, 7, 8, 9, 11, 12, 13, 15, 17, 22}
FIGURES_TO_SUPPLEMENT = {2, 3, 4, 5, 6, 7, 9, 12, 14}


def caption_map(doc, kind: str):
    rx = re.compile(rf"^{kind}\.\s*(\d+)\.", re.I) if kind == "Fig" else re.compile(r"^Table\s+(\d+)\.", re.I)
    out = {}
    for p in doc.paragraphs:
        m = rx.match(p.text.strip())
        if m:
            out[int(m.group(1))] = p
    return out


def next_table_element(caption):
    node = caption._p.getnext()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return node
        if node.tag == qn("w:p") and "Table " in "".join(node.itertext()):
            break
        node = node.getnext()
    raise RuntimeError(f"No table after {caption.text}")


def preceding_image_paragraph(caption):
    node = caption._p.getprevious()
    while node is not None:
        if node.tag == qn("w:p") and any(True for _ in node.iter(qn("w:drawing"))):
            return node
        if node.tag == qn("w:p") and "Fig." in "".join(node.itertext()):
            break
        node = node.getprevious()
    raise RuntimeError(f"No image before {caption.text}")


def image_blob(doc, image_paragraph):
    blips = list(image_paragraph.iter(qn("a:blip")))
    if not blips:
        raise RuntimeError("Drawing has no embedded image")
    rid = blips[0].get(qn("r:embed"))
    part = doc.part.related_parts[rid]
    suffix = ".png" if "png" in part.content_type else ".jpg"
    return part.blob, suffix


def remove_node(node):
    parent = node.getparent()
    if parent is not None:
        parent.remove(node)


def remove_paragraphs_between(doc, start_text: str, end_text: str):
    start = next(p for p in doc.paragraphs if p.text.startswith(start_text))._p
    end = next(p for p in doc.paragraphs if p.text.startswith(end_text))._p
    node = start.getnext()
    while node is not None and node is not end:
        following = node.getnext()
        remove_node(node)
        node = following


def insert_after(paragraph, values):
    anchor = paragraph._p
    for value in values:
        new = paragraph._parent.add_paragraph(style="Normal")
        new._p.getparent().remove(new._p)
        anchor.addnext(new._p)
        anchor = new._p
        new.paragraph_format.first_line_indent = Inches(0.35)
        new.paragraph_format.line_spacing = 1.08
        run = new.add_run(value)
        set_font(run, size=10.5)


def replace_paragraph(p, value):
    p.clear(); p.style = "Normal"
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(value), size=10.5)


def archive_removed_items(main, supplement):
    """Copy superseded tables/figures into a clearly labelled audit section."""
    table_caps = caption_map(main, "Table")
    fig_caps = caption_map(main, "Fig")
    h = supplement.add_paragraph("S7 Archived diagnostics and superseded pilots", style="Heading 1")
    for run in h.runs:
        set_font(run, size=12, bold=True)
    p = supplement.add_paragraph(
        "These items are retained for auditability but are not part of the primary evidential chain. "
        "Classifier diagnostics document implementation choices; pilot artificial-star results are "
        "superseded by the n=200 fixed-scene experiment in the main text."
    )
    for run in p.runs:
        set_font(run, size=10.5)

    next_table = 7
    for old in sorted(TABLES_TO_SUPPLEMENT):
        cap = table_caps[old]
        new_cap = re.sub(r"^Table\s+\d+\.", f"Table S{next_table}.", cap.text)
        cp = supplement.add_paragraph(new_cap, style="Caption")
        for run in cp.runs:
            set_font(run, size=9)
        supplement._element.body.insert(-1, copy.deepcopy(next_table_element(cap)))
        format_three_line(supplement.tables[-1])
        next_table += 1

    MEDIA_DIR.mkdir(parents=True, exist_ok=True)
    next_fig = 6
    for old in sorted(FIGURES_TO_SUPPLEMENT):
        cap = fig_caps[old]
        blob, suffix = image_blob(main, preceding_image_paragraph(cap))
        path = MEDIA_DIR / f"figS{next_fig}_from_main_fig{old}{suffix}"
        path.write_bytes(blob)
        supplement.add_picture(str(path), width=Inches(6.15))
        supplement.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_cap = re.sub(r"^Fig\.\s*\d+\.", f"Fig. S{next_fig}.", cap.text)
        cp = supplement.add_paragraph(new_cap, style="Caption")
        for run in cp.runs:
            set_font(run, size=9)
        next_fig += 1


def remove_captioned_items(doc):
    tables = caption_map(doc, "Table")
    figures = caption_map(doc, "Fig")
    for old in TABLES_TO_SUPPLEMENT:
        remove_node(next_table_element(tables[old]))
        remove_node(tables[old]._p)
    for old in FIGURES_TO_SUPPLEMENT:
        remove_node(preceding_image_paragraph(figures[old]))
        remove_node(figures[old]._p)


def renumber(doc):
    """Renumber surviving captions, textual references, bookmarks and hyperlinks."""
    table_old = sorted(caption_map(doc, "Table"))
    fig_old = sorted(caption_map(doc, "Fig"))
    tm = {old: new for new, old in enumerate(table_old, 1)}
    fm = {old: new for new, old in enumerate(fig_old, 1)}
    table_rx = re.compile(r"\bTable\s+(\d+)\b")
    fig_rx = re.compile(r"\b(Fig(?:ure)?\.?)\s*(\d+)\b")
    for node in doc._element.body.iter(qn("w:t")):
        if not node.text:
            continue
        node.text = table_rx.sub(lambda m: f"Table {tm.get(int(m.group(1)), int(m.group(1)))}", node.text)
        node.text = fig_rx.sub(lambda m: f"{m.group(1)} {fm.get(int(m.group(2)), int(m.group(2)))}", node.text)
    for node in doc._element.body.iter(qn("w:bookmarkStart")):
        name = node.get(qn("w:name")) or ""
        m = re.fullmatch(r"table_(\d+)", name)
        if m and int(m.group(1)) in tm:
            node.set(qn("w:name"), f"table_{tm[int(m.group(1))]}")
        m = re.fullmatch(r"figure_(\d+)", name)
        if m and int(m.group(1)) in fm:
            node.set(qn("w:name"), f"figure_{fm[int(m.group(1))]}")
    for node in doc._element.body.iter(qn("w:hyperlink")):
        anchor = node.get(qn("w:anchor")) or ""
        m = re.fullmatch(r"table_(\d+)", anchor)
        if m and int(m.group(1)) in tm:
            node.set(qn("w:anchor"), f"table_{tm[int(m.group(1))]}")
        m = re.fullmatch(r"figure_(\d+)", anchor)
        if m and int(m.group(1)) in fm:
            node.set(qn("w:anchor"), f"figure_{fm[int(m.group(1))]}")


def insert_zero_shot_figure(doc, source):
    """Keep one compact illustration for the negative control/adaptation transition."""
    cap = caption_map(source, "Fig")[10]
    blob, suffix = image_blob(source, preceding_image_paragraph(cap))
    MEDIA_DIR.mkdir(parents=True, exist_ok=True)
    path = MEDIA_DIR / f"zero_shot_to_few_shot{suffix}"
    path.write_bytes(blob)
    doc.add_picture(str(path), width=Inches(5.6))
    image_p = doc.paragraphs[-1]
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_cap = doc.add_paragraph(
        "Fig. 10. Frozen simulation-trained screening and supervised few-shot target adaptation "
        "on the two external survey cutouts. The illustration is qualitative; all quoted recall is "
        "measured on a spatially untouched target partition.", style="Caption"
    )
    for run in new_cap.runs:
        set_font(run, size=9)
    target = next(p for p in doc.paragraphs if p.text.startswith("5.3 HST/ACS crowded-field recovery"))._p
    target.addprevious(image_p._p)
    target.addprevious(new_cap._p)


def focus_narrative(doc):
    # A single methodological rule replaces apologetic explanations of each scale.
    p = next(p for p in doc.paragraphs if p.text.startswith("The dominant dense-field failure mode"))
    replace_paragraph(p,
        "The dominant dense-field failure mode is structural blending. Candidate groups are processed "
        "by paired-PSF fitting, joint multi-source fitting, or residual deblending according to group "
        "complexity. All spatial scales and structural settings—including grouping radius, polynomial "
        "order, and calibration-tile size—were treated as hyperparameters selected by validation-set "
        "performance on the CSST-like simulations and then kept fixed for HST transfer, except for "
        "explicit target measurements such as FWHM. They are deployment settings rather than physical constants."
    )
    p = next(p for p in doc.paragraphs if p.text.startswith("Astrometric refinement uses a three-stage"))
    replace_paragraph(p,
        "Astrometric refinement uses a validation-gated cascade: a polynomial removes large-scale WCS "
        "distortion, a look-up table models residual spatial structure, and Gaussian-process correction "
        "is retained only when it improves held-out residuals. The polynomial order follows the common "
        "hyperparameter-selection rule above and is not asserted to be instrument-independent."
    )
    p = next(p for p in doc.paragraphs if p.text.startswith("To prevent target-test leakage"))
    replace_paragraph(p,
        "To prevent target-test leakage, images are split into spatially disjoint fit, threshold-selection, "
        "and test regions. The 200 × 200 pixel tile is a validation-selected sampling unit, not a physical "
        "scale; changing instruments requires revalidation in PSF-normalized units. Final metrics are "
        "reported only on regions excluded from model fitting and threshold selection."
    )

    # Collapse classifier development to one result paragraph before the retained comparison.
    remove_paragraphs_between(doc, "5.1 CSST-like simulation validation", "On the CSST-like simulations, calibrated SExtractor")
    h = next(p for p in doc.paragraphs if p.text.startswith("5.1 CSST-like simulation validation"))
    insert_after(h, [
        "Validation-selected per-chip calibration was necessary for heterogeneous simulated chips: on "
        "chip 12 it raised RandomForest recall from 72.06% to 95.01% while reducing precision from "
        "97.13% to 81.11%. Across the complete pipeline, RandomForest, XGBoost, and hybrid stacking "
        "converged to 94.35–94.60% mean recall; the RandomForest was retained as the conservative branch "
        "because it provided the simplest accuracy–cost operating point. Full classifier and fusion "
        "diagnostics are archived in Supplementary Section S7."
    ])

    # Remove detailed per-chip funnel discussion after its table has moved.
    start = next(p for p in doc.paragraphs if p.text.startswith("The final AstroCFR competition configuration"))
    replace_paragraph(start,
        "Across the four CSST-like chips, the registered AstroCFR configuration reaches 94.6% mean recall, "
        "10.9 mas astrometric RMS, and 0.066 mag photometric RMS. Per-chip catalogues and candidate-funnel "
        "diagnostics are provided in Supplementary Section S7; these simulation-catalogue values are not "
        "interpreted as blind real-field purity."
    )
    for prefix in ("The candidate-count trajectory", "The funnel also explains", "Astrometric and photometric correction chains"):
        q = next((p for p in doc.paragraphs if p.text.startswith(prefix)), None)
        if q is not None:
            remove_node(q._p)

    # Zero-shot is a two-paragraph negative control plus one retained illustration.
    remove_paragraphs_between(doc, "5.2 Simulation-to-real target adaptation", "5.3 HST/ACS crowded-field recovery")
    h = next(p for p in doc.paragraphs if p.text.startswith("5.2 Simulation-to-real target adaptation"))
    insert_after(h, [
        "A frozen RandomForest trained on the four CSST-like chips was evaluated on public Pan-STARRS1 "
        "M31 and Legacy Survey M13 cutouts. Without target calibration, the screen collapses: retained "
        "recall is 2.8% and 0.0%, respectively. This is a negative control demonstrating domain shift, "
        "not evidence of zero-shot generalization; full counts are archived in Supplementary Section S7.",
        "We therefore use supervised few-shot target adaptation. FWHM and candidate normalization are "
        "estimated from the target image, while a small labelled spatial region refits the RandomForest "
        "and a disjoint region selects its threshold. Held-out recall rises to 43.3% on M31 and 7.1% on "
        "M13. The accompanying illustration shows the failure and recovery; the more relevant three-field "
        "HST calibration-budget experiment is reported below."
    ])

    # Remove the superseded ePSF precursor/pilot narrative; controlled comparison follows directly.
    begin = next(p for p in doc.paragraphs if p.text.startswith("The HST benchmark in Section"))
    end = next(p for p in doc.paragraphs if p.text.startswith("To make the deployment claim falsifiable"))
    node = begin._p
    while node is not end._p:
        following = node.getnext(); remove_node(node); node = following

    # Pilot n=40 is superseded; keep only the reference-star claim here.
    p = next(p for p in doc.paragraphs if p.text.startswith("The clearest supported claim is local"))
    replace_paragraph(p,
        "The clearest supported claim is local and operational. On dense NGC 6752 V≤20 references "
        "(n=402), AstroCFR ePSF-deblend recovers 87.6% (95% CI 84.0–90.4%), versus 57.0% for DAO, "
        "28.6% for SEP, and 52.7% for AstroCFR-RF. The definitive fixed-scene artificial-star test uses "
        "n=200 per magnitude–density stratum and is reported below; its smaller pilot is archived in S7."
    )

    # Resource detail and literature matrix belong in the supplement.
    p = next((p for p in doc.paragraphs if p.text.startswith("The CNN classifier is a separate")), None)
    if p:
        replace_paragraph(p,
            "CNN CPU/GPU resource measurements and the complete RF/CNN/Transformer ablation are archived "
            "in Supplementary Section S7. They document the optional simulation-domain classifier and do "
            "not alter the HST CPU operating-point comparison."
        )
    p = next((p for p in doc.paragraphs if p.text.startswith("A controlled RF/CNN/lightweight")), None)
    if p:
        remove_node(p._p)


def main():
    main_doc = Document(SOURCE)
    source_doc = Document(SOURCE)
    supplement = Document(SUP_SOURCE)
    archive_removed_items(main_doc, supplement)
    remove_captioned_items(main_doc)
    focus_narrative(main_doc)
    insert_zero_shot_figure(main_doc, source_doc)
    renumber(main_doc)
    for table in main_doc.tables:
        format_three_line(table)
    main_doc.save(DEST)
    supplement.save(SUP_DEST)
    print(DEST)
    print(SUP_DEST)


if __name__ == "__main__":
    main()
