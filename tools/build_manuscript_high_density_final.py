#!/usr/bin/env python
"""Revise the v45 formal manuscript around the completed high-density audit.

The source document is never overwritten.  Paragraph properties and table
styles are retained; only content is replaced and the existing cross-reference
hyperlink convention is reused for newly introduced citations.
"""
from __future__ import annotations

import json
import re
import statistics
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

from build_manuscript_v30_submission_fixes import bookmark, external_link, internal_link, set_cell, set_font

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v45_formal_journal_format.docx"
DST = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v45_high_density_final.docx"
BASE = ROOT / "results" / "acsggct_all11_baselines" / "hst_unified_baseline_summary.json"
LIT = ROOT / "results" / "hst_literature_method_benchmark_all11" / "summary.json"

FIELDS = ("ngc2808", "ngc5286", "ngc6388", "ngc6441", "ngc0104", "ngc0362", "ngc6093", "ngc6624", "ngc6397", "ngc6752", "ngc1851")
FIELD_LABEL = {x: x.upper().replace("NGC", "NGC ").replace(" 0", " 0") for x in FIELDS}
FIELD_LABEL.update({"ngc0104": "NGC 0104", "ngc0362": "NGC 0362"})
BASE_METHOD = {
    "dao": "DAOStarFinder", "sep": "SEP/SExtractor-style", "photutils_psf": "Photutils PSFPhotometry",
    "wpdc_epsf_deblend": "AstroCFR ePSF + residual deblend", "wpdc_spatial_epsf_joint": "AstroCFR spatial-ePSF + joint fit",
}
LIT_METHOD = {"global_epsf_joint": "Global empirical ePSF + neighbour joint", "three_gaussian_dpsf_joint": "Three-Gaussian dPSF + neighbour joint"}

LINKS = {
    "Anderson & King (2000)": "https://doi.org/10.1086/316632",
    "Sirianni et al. (2005)": "https://doi.org/10.1086/444553",
    "Anderson & King (2006)": "https://www.stsci.edu/hst/instrumentation/acs/documentation/instrument-science-reports-isrs",
    "Anderson et al. (2008)": "https://doi.org/10.1088/0004-6256/135/6/2055",
    "Bellini et al. (2011)": "https://doi.org/10.1086/659878",
    "Chambers et al. (2016)": "https://doi.org/10.48550/arXiv.1612.05560",
    "Dey et al. (2019)": "https://doi.org/10.3847/1538-3881/ab089d",
    "Gaia Collaboration et al. (2016)": "https://doi.org/10.1051/0004-6361/201629272",
    "Gaia Collaboration et al. (2023)": "https://doi.org/10.1051/0004-6361/202243940",
    "Astropy Collaboration (2022)": "https://doi.org/10.3847/1538-4357/ac7c74",
    "Harris et al. (2020)": "https://doi.org/10.1038/s41586-020-2649-2",
    "Hunter (2007)": "https://doi.org/10.1109/MCSE.2007.55",
    "Virtanen et al. (2020)": "https://doi.org/10.1038/s41592-019-0686-2",
    "Barbary (2016)": "https://doi.org/10.21105/joss.00058",
    "Bradley et al. (2024)": "https://ascl.net/2401.012",
    "Pedregosa et al. (2011)": "https://jmlr.org/papers/v12/pedregosa11a.html",
    "Goodfellow et al. (2016)": "https://www.deeplearningbook.org/",
    "He et al. (2016)": "https://doi.org/10.1109/CVPR.2016.90",
    "Shi et al. (2024)": "https://doi.org/10.1088/1674-4527",
    "Long et al. (2025)": "https://doi.org/10.3847/1538-4365/ad9244",
    "Han et al. (2026)": "https://doi.org/10.1016/j.ascom.2026.101112",
    "Wainer et al. (2025)": "https://doi.org/10.3847/2515-5172/adecef",
    "Shaw et al. (2025)": "https://doi.org/10.1093/rasti/rzaf006",
    "Nie et al. (2025)": "https://doi.org/10.1088/1674-4527/adfd22",
    "Libralato et al. (2024)": "https://doi.org/10.48550/arXiv.2411.02487",
    "Wang et al. (2026)": "https://doi.org/10.3847/1538-4365/ae5053",
    "Zhang et al. (2026)": "https://doi.org/10.3847/1538-3881/ae7718",
    "Espinosa et al. (2025)": "https://doi.org/10.1051/0004-6361/202555342",
    "De Alba et al. (2026)": "https://doi.org/10.1016/j.ascom.2026.101165",
    "Centofanti et al. (2026)": "https://doi.org/10.1051/0004-6361/202558730",
    "Salaris et al. (2024)": "https://doi.org/10.1002/asna.20240018",
    "Ben-David et al. (2010)": "https://doi.org/10.1007/s10994-009-5152-4",
    "Pan & Yang (2010)": "https://doi.org/10.1109/TKDE.2009.191",
    "Ganin et al. (2016)": "https://jmlr.org/papers/v17/15-239.html",
    "Breiman (2001)": "https://doi.org/10.1023/A:1010933404324",
    "Chen & Guestrin (2016)": "https://doi.org/10.1145/2939672.2939785",
    "Dolphin (2000)": "https://doi.org/10.1086/316630",
    "Stetson (1987)": "https://doi.org/10.1086/131977",
    "Bertin & Arnouts (1996)": "https://doi.org/10.1051/aas:1996164",
    "Melchior et al. (2018)": "https://doi.org/10.1016/j.ascom.2018.07.001",
    "Long et al. (2015)": "https://proceedings.mlr.press/v37/long15.html",
    "Schlafly et al. (2018)": "https://doi.org/10.3847/1538-4365/aaa3e2",
    "Wang et al. (2024)": "https://doi.org/10.1088/1674-4527/ad4df5",
    "Zhang et al. (2023)": "https://doi.org/10.1093/mnras/stad3815",
    "Zhang et al. (2024)": "https://doi.org/10.48550/arXiv.2409.13296",
    "Yan et al. (2026a)": "https://doi.org/10.3847/1538-4365/ae4a24",
    "Yan et al. (2026b)": "https://doi.org/10.3847/1538-3881/ae505f",
    "Lai et al. (2026)": "https://doi.org/10.1088/1538-3873/ae6d75",
    "Burke et al. (2019)": "https://doi.org/10.1093/mnras/stz2845",
    "Dalcanton et al. (2009)": "https://doi.org/10.1088/0067-0049/183/1/67",
    "Dalcanton et al. (2012)": "https://doi.org/10.1088/0067-0049/200/2/18",
    "Williams et al. (2021)": "https://doi.org/10.3847/1538-4365/abd6f6",
    "Sabbi et al. (2016)": "https://doi.org/10.3847/0067-0049/222/1/11",
}

# The journal requires the reference list to follow the order in which sources
# are first cited.  This exact cumulative v41+v45 list keeps the literature
# depth while preventing uncited legacy entries from remaining in the paper.
REFERENCE_ORDER = (
    "Anderson, J., & King, I. R. (2000)",
    "Anderson, J., & King, I. R. (2006)",
    "Anderson, J., et al. (2008)",
    "Sirianni, M.,",
    "Bellini, A., Anderson, J., & Bedin, L. R. (2011)",
    "Chambers, K. C., et al. (2016)",
    "Dey, A., et al. (2019)",
    "Gaia Collaboration, et al. (2016)",
    "Gaia Collaboration, et al. (2023)",
    "Astropy Collaboration, et al. (2022)",
    "Harris, C. R., et al. (2020)",
    "Hunter, J. D. (2007)",
    "Virtanen, P., et al. (2020)",
    "Stetson, P. B. (1987)",
    "Bertin, E., & Arnouts, S. (1996)",
    "Barbary, K. (2016)",
    "Bradley, L.,",
    "Dolphin, A. E. (2000)",
    "Melchior, P.,",
    "Schlafly, E. F.,",
    "Wang, Y., Sun, R., Deng, T., Zhao, C., Zhao, P., Yang, J., Jia, P., Liu, H., & Zhou, J. (2024)",
    "Shi, R.-F., Huang, Y., Li, X.-Y., & Zhang, H.-W. (2024)",
    "Zhang, Y., Cao, Z., Wang, F., Lam, M. I., Deng, H., Mei, Y., & Tan, L. (2023)",
    "Zhang, S., Fang, G., Song, J., Li, R., Gu, Y., Lin, Z., Zhou, C., Dai, Y., & Kong, X. (2024)",
    "Long, M., Xin, J., Du, J., Zhao, J., Wang, X., et al. (2025)",
    "Han, Z., Zhang, T., Liu, C., & Ling, C. (2026)",
    "Yan, Y., Liu, C., Li, J., & Wang, F. (2026)",
    "Lai, J., Lam, M. I., Chen, J., Zhang, X., Tian, H., Chen, X., et al. (2026)",
    "Burke, C. J.,",
    "Shaw, R. A., Fotopoulou, S., Birkinshaw, M., Maddox, N., & Stewart, H. (2025)",
    "Libralato, M., et al. (2024)",
    "Nie, J., Wei, P., Cao, Z., Yan, Y., Liu, C., Tian, H., et al. (2025)",
    "Espinosa, S.,",
    "Centofanti, E.,",
    "Wang, P., Wei, P., Liu, C., Wang, R., Wang, F., & Zhang, X. (2026)",
    "Zhang, S., Wang, L., Diao, Y., Yan, Z., Peng, X., Liu, Y., Shan, H., Wang, G., Liu, F., Wei, C., Nie, L., Chen, X., Ding, H., & Zheng, Z. (2026)",
    "De Alba, K.,",
    "Wainer, T. M.,",
    "Salaris, M., et al. (2024)",
    "Dalcanton, J. J., et al. (2009)",
    "Dalcanton, J. J., et al. (2012)",
    "Williams, B. F., et al. (2021)",
    "Sabbi, E., et al. (2016)",
    "Ben-David, S., Blitzer, J., Crammer, K., & Pereira, F. (2010)",
    "Pan, S. J., & Yang, Q. (2010)",
    "Ganin, Y., et al. (2016)",
    "Long, M., Cao, Y., Wang, J., & Jordan, M. I. (2015)",
    "Breiman, L. (2001)",
    "Chen, T., & Guestrin, C. (2016)",
    "Pedregosa, F., et al. (2011)",
    "Goodfellow, I., Bengio, Y., & Courville, A. (2016)",
    "He, K., Zhang, X., Ren, S., & Sun, J. (2016)",
    "Yan, Y., Wu, Y., Nie, J., Zhang, T., Liu, C., Ban, Z., et al. (2026)",
)

# Citation spellings used in the Introduction and methods map.  They are
# paired with REFERENCE_ORDER when reference bookmarks are created.
CITATION_SEQUENCE = (
    "Anderson & King (2000)", "Anderson & King (2006)", "Anderson et al. (2008)",
    "Sirianni et al. (2005)", "Bellini et al. (2011)", "Chambers et al. (2016)",
    "Dey et al. (2019)", "Gaia Collaboration et al. (2016)", "Gaia Collaboration et al. (2023)",
    "Astropy Collaboration (2022)", "Harris et al., 2020", "Hunter, 2007", "Virtanen et al., 2020",
    "Stetson, 1987", "Bertin & Arnouts, 1996", "Barbary, 2016", "Bradley et al., 2024",
    "Dolphin, 2000", "Melchior et al., 2018", "Schlafly et al. (2018)", "Wang et al. (2024)",
    "Shi et al. (2024)", "Zhang et al. (2023)", "Zhang et al. (2024)", "Long et al. (2025)",
    "Han et al. (2026)", "Yan et al. (2026a)", "Lai et al. (2026)", "Burke et al. (2019)",
    "Shaw et al. (2025)", "Libralato et al. (2024)", "Nie et al. (2025)", "Espinosa et al. (2025)",
    "Centofanti et al. (2026)", "Wang et al. (2026)", "Zhang et al. (2026)", "De Alba et al. (2026)",
    "Wainer et al. (2025)", "Salaris et al. (2024)", "Dalcanton et al. (2009)",
    "Dalcanton et al. (2012)", "Williams et al. (2021)", "Sabbi et al. (2016)",
    "Ben-David et al. (2010)", "Pan & Yang (2010)", "Ganin et al. (2016)", "Long et al. (2015)",
    "Breiman, 2001", "Chen & Guestrin, 2016", "Pedregosa et al., 2011",
    "Goodfellow et al., 2016", "He et al., 2016", "Yan et al. (2026b)",
)
CURRENT_REF_ANCHORS = {}


def pfind(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix) or prefix in p.text)


def refresh_embedded_figure(doc, partname: str, source: Path) -> None:
    """Replace one existing PNG payload without altering its Word layout anchor."""
    for part in doc.part.related_parts.values():
        if str(part.partname) == partname:
            part._blob = source.read_bytes()
            return
    raise KeyError(f"Embedded figure not found: {partname}")


def label_registered_csst_branch(doc):
    """Disambiguate the registered CSST AstroCFR branch from spatial-ePSF."""
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip() for c in table.rows[0].cells]
        if "Tool" not in headers:
            continue
        tool_col = headers.index("Tool")
        for row in table.rows[1:]:
            if row.cells[tool_col].text.strip() == "AstroCFR":
                set_cell(row.cells[tool_col], "AstroCFR ePSF + residual deblend\n(registered full-frame branch)", size=7.2)
    for p in doc.paragraphs:
        if p.text.startswith("Table 2. Registered CSST-like measurement audit"):
            new_text = p.text.replace(
                "Table 2. Registered CSST-like measurement audit.",
                "Table 2. Registered CSST-like measurement audit using AstroCFR ePSF + residual deblend (registered full-frame branch).",
                1,
            )
            p.clear()
            run = p.add_run(new_text)
            set_font(run, size=9.0)


def enrich_chip_characteristics_table(doc):
    """Make Table 1 descriptive labels reproducible rather than impressionistic."""
    table = doc.tables[0]
    if len(table.columns) == 5:
        table.add_column(Inches(1.35))
    headers = ["Chip", "Background RMS (DN)", "PSF FWHM (px)",
               "Supplied-reference density (arcmin^-2)", "Magnitude span", "Dominant difficulty"]
    rows = (
        ("12", "3.55", "1.27", "32.7", "18.03-23.03", "highest supplied-reference density; faint sources"),
        ("13", "3.97", "1.33", "19.8", "17.36-21.90", "intermediate supplied-reference density"),
        ("17", "4.39", "1.67", "7.7", "15.51-20.89", "bright sources and higher background"),
        ("18", "5.36", "1.93", "7.7", "15.30-20.83", "broader PSF and bright-star artifacts/blends"),
    )
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, size=7.0)
    for row, values in zip(table.rows[1:], rows):
        for cell, value in zip(row.cells, values):
            set_cell(cell, value, size=7.0)


def set_main_title(doc):
    """Use a science-facing title that identifies the primary method and task."""
    if not doc.paragraphs:
        return
    p = doc.paragraphs[0]
    p.clear()
    run = p.add_run("AstroCFR: A Single-Image Crowded-Field Measurement Pipeline with Spatial-ePSF Joint Fitting")
    set_font(run, size=16.0, bold=True)


def clear_preserve_format(paragraph):
    # Paragraph.clear() removes runs but leaves pPr (including indentation,
    # spacing, outline level and keep-with-next settings) untouched.
    paragraph.clear()


def remove_paragraph(paragraph):
    """Remove an obsolete source paragraph while retaining surrounding layout."""
    element = paragraph._element
    element.getparent().remove(element)


def linked_rewrite(paragraph, text, links=None):
    links = dict(links or LINKS)
    # Author-year prose and parenthetical citations use both "A et al.
    # (2025)" and "(A et al., 2025)" forms in the formal source.  Give
    # both forms the same hyperlink target.
    for label, url in list(links.items()):
        m = re.fullmatch(r"(.+) \(((?:19|20)\d{2}[a-z]?)\)", label)
        if m:
            links.setdefault(f"{m.group(1)}, {m.group(2)}", url)
    clear_preserve_format(paragraph)
    # Keep the document's existing body font and do not touch paragraph
    # indentation.  Hyperlink runs use the same Times New Roman convention as
    # the source manuscript.
    cursor = 0
    choices = sorted(((pos, label, url) for label, url in links.items() if (pos := text.find(label)) >= 0), key=lambda x: x[0])
    for pos, label, url in choices:
        if pos < cursor:
            continue
        if pos > cursor:
            run = paragraph.add_run(text[cursor:pos]); set_font(run, size=10.5)
        # In-text citations navigate to the corresponding References entry;
        # DOI/official URLs remain in the bibliography itself.
        anchor = CURRENT_REF_ANCHORS.get(label)
        if anchor:
            internal_link(paragraph, label, anchor, size=10.5)
        else:
            external_link(paragraph, label, url, size=10.5)
        cursor = pos + len(label)
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:]); set_font(run, size=10.5)


def plain_rewrite(paragraph, text):
    clear_preserve_format(paragraph)
    run = paragraph.add_run(text)
    set_font(run, size=10.5)


def insert_after(paragraph, style="Normal"):
    new = paragraph._parent.add_paragraph(style=style)
    new._p.getparent().remove(new._p)
    paragraph._p.addnext(new._p)
    return new


def normalize_a4_layout(doc):
    """Remove obsolete landscape section breaks and restore A4 portrait pages."""
    for paragraph in doc.paragraphs:
        p_pr = paragraph._p.pPr
        if p_pr is None:
            continue
        sect_pr = p_pr.find(qn("w:sectPr"))
        if sect_pr is not None:
            p_pr.remove(sect_pr)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)


def equation_after(doc, anchor, number):
    """Insert a placeholder that Word subsequently converts to native OMML.

    Word's equation engine (rather than python-docx) is the authoritative
    renderer for UnicodeMath alignment markers such as ``#(1)``.
    """
    paragraph = doc.add_paragraph(style="Normal")
    paragraph._p.getparent().remove(paragraph._p)
    anchor.addnext(paragraph._p)
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = None
    paragraph.paragraph_format.space_after = None
    paragraph.add_run(f"[[WORD_NATIVE_EQUATION_{number}]]")
    return paragraph


def _border(parent, edge, value="nil", size=None):
    element = parent.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        parent.append(element)
    element.set(qn("w:val"), value)
    if size is not None:
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), "000000")


def apply_three_line_table(table):
    """Apply the journal three-line table convention to a data table."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    _border(borders, "top", "single", 8)
    _border(borders, "bottom", "single", 8)
    for edge in ("left", "right", "insideH", "insideV"):
        _border(borders, edge, "nil")
    # The header's lower rule is the middle rule; all data-cell internal
    # borders are deliberately absent.
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        _border(tc_borders, "bottom", "single", 6)
        for edge in ("top", "left", "right"):
            _border(tc_borders, edge, "nil")
    for row in table.rows[1:]:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "bottom", "left", "right"):
                _border(tc_borders, edge, "nil")


def table_caption_rewrite(paragraph, body):
    """Match the existing Caption style: bold label, regular title text."""
    clear_preserve_format(paragraph)
    paragraph.style = "Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.keep_with_next = True
    first = paragraph.add_run("Table 8.")
    set_font(first, size=10.0, bold=True)
    rest = paragraph.add_run(" " + body)
    set_font(rest, size=10.0)
    bookmark(paragraph, "table_8", 5108)


def table_reference_rewrite(paragraph, text):
    """Make the narrative Table 8 mention a true internal cross-reference."""
    clear_preserve_format(paragraph)
    label = "Table 8"
    before, after = text.split(label, 1)
    if before:
        run = paragraph.add_run(before); set_font(run, size=10.5)
    internal_link(paragraph, label, "table_8", size=10.5)
    run = paragraph.add_run(after); set_font(run, size=10.5)


def pct(value, digits=1):
    return f"{100 * float(value):.{digits}f}%"


def ci(value, digits=1):
    if not value:
        return ""
    return f"{pct(value[0], digits)}–{pct(value[1], digits)}"


def load_rows():
    base = json.loads(BASE.read_text(encoding="utf-8"))["results"]
    lit = json.loads(LIT.read_text(encoding="utf-8"))["results"]
    by = {(r["cluster"], r["method"]): r for r in base}
    lb = {(r["cluster"], r["method"]): r for r in lit}
    return by, lb


def runtime_density_summary(base):
    """Return robust field-median runtime and magnitude-stratified summaries."""
    methods = ("dao", "sep", "photutils_psf", "wpdc_epsf_deblend", "wpdc_spatial_epsf_joint")
    out = {}
    for method in methods:
        rows = [base[(f, method)] for f in FIELDS]
        out[method] = {
            "runtime_s_per_mpix": statistics.median(r["runtime_s_per_mpix"] for r in rows),
            "runtime_s": statistics.median(r["runtime_s"] for r in rows),
            "recall_v_le_18": statistics.median(r["recall_v_le_18"] for r in rows),
            "recall_v_le_20": statistics.median(r["recall_v_le_20"] for r in rows),
            "recall_v_le_22": statistics.median(r["recall_v_le_22"] for r in rows),
        }
    return out


def rebuild_cross_field_table(doc, base, lit, table=None, table_index=7):
    table = table if table is not None else doc.tables[table_index]
    headers = ["Field", "Evidence tier / scene class", "Method / branch", "Catalogue recovery (95% CI)", "Crowded-subset recovery", "Position RMS / mas", "Magnitude RMS / mag", "Runtime / s MPix⁻¹", "Best recovery branch", "Best conditional precision branch(es)"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, size=7.0)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    for field in FIELDS:
        field_rows = []
        for method in ("dao", "sep", "photutils_psf", "wpdc_epsf_deblend", "wpdc_spatial_epsf_joint"):
            field_rows.append((method, base[(field, method)]))
        for method in ("global_epsf_joint", "three_gaussian_dpsf_joint"):
            field_rows.append((method, lit[(field, method)]))
        baseline = [x for method, x in field_rows if method in BASE_METHOD]
        best = max(baseline, key=lambda x: x["high_density_v20_recall"])
        best_name = BASE_METHOD[best["method"]]
        best_pos = min(baseline, key=lambda x: x["astrometric_rms_mas"])
        best_mag = min(baseline, key=lambda x: x["photometric_rms_mag"])
        precision = f"Position: {BASE_METHOD[best_pos['method']]}; photometry: {BASE_METHOD[best_mag['method']]}"
        for idx, (method, x) in enumerate(field_rows):
            if method in BASE_METHOD:
                rec = f"{pct(x['test_completeness'])} [{ci(x.get('test_completeness_ci95'))}]"
                dense = f"V<=20, >=3 neighbours: {pct(x['high_density_v20_recall'])} [{ci(x.get('high_density_v20_ci95'))}]; n={x['high_density_v20_n']}"
                runtime = f"{x['runtime_s_per_mpix']:.2f}"
            else:
                rec_value = x.get("matched_test_dense_n", 0) / max(x.get("dense_reference_test_n", 1), 1)
                rec = f"{pct(rec_value)} (dense held-out denominator)"
                dense = f"dense test: {pct(rec_value)}; n={x['dense_reference_test_n']}"
                runtime = "—"
            values = [
                (FIELD_LABEL[field] if idx == 0 else ""),
                ("ACSGGCT HST F606W; 11-field controlled tier" if idx == 0 else ""),
                BASE_METHOD.get(method, LIT_METHOD.get(method, method)), rec, dense,
                (f"{x['astrometric_rms_mas']:.2f}" if x.get('astrometric_rms_mas') is not None else "—"),
                (f"{x['photometric_rms_mag']:.3f}" if x.get('photometric_rms_mag') is not None else "—"), runtime,
                (best_name if idx == 0 else ""), (precision if idx == 0 else ""),
            ]
            cells = table.add_row().cells
            for cell, value in zip(cells, values):
                set_cell(cell, value, size=6.8)
    # Preserve formal-table pagination behaviour for the newly added rows:
    # rows are never split internally, and the header repeats on continuation
    # pages.  The table itself may continue across pages when it is longer than
    # one page.
    for ridx, row in enumerate(table.rows):
        trpr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trpr.append(cant)
        if ridx == 0:
            repeat = OxmlElement("w:tblHeader")
            trpr.append(repeat)
    apply_three_line_table(table)


def rebuild_compact_main_table(doc, old_table, base):
    """Replace the long main-text matrix with a compact paired field table."""
    table = doc.add_table(rows=1, cols=7)
    table.style = old_table.style
    table._tbl.getparent().remove(table._tbl)
    old_table._tbl.addprevious(table._tbl)
    old_table._tbl.getparent().remove(old_table._tbl)
    headers = ["Field", "N dense", "Spatial-ePSF recovery / %", "Photutils recovery / %", "Delta / pp", "Pos. RMS S / P / mas", "Mag. RMS S / P / mag"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, size=7.5)
    spatial_rows, phot_rows = [], []
    for field in FIELDS:
        spatial = base[(field, "wpdc_spatial_epsf_joint")]
        phot = base[(field, "photutils_psf")]
        spatial_rows.append(spatial); phot_rows.append(phot)
        values = [
            FIELD_LABEL[field], str(spatial["high_density_v20_n"]),
            f"{pct(spatial['high_density_v20_recall'])} [{ci(spatial['high_density_v20_ci95'])}]",
            f"{pct(phot['high_density_v20_recall'])} [{ci(phot['high_density_v20_ci95'])}]",
            f"{100*(spatial['high_density_v20_recall']-phot['high_density_v20_recall']):+.1f}",
            f"{spatial['astrometric_rms_mas']:.2f} / {phot['astrometric_rms_mas']:.2f}",
            f"{spatial['photometric_rms_mag']:.3f} / {phot['photometric_rms_mag']:.3f}",
        ]
        cells = table.add_row().cells
        for cell, value in zip(cells, values): set_cell(cell, value, size=7.2)
    values = [
        "Median", "—",
        f"{pct(statistics.median(x['high_density_v20_recall'] for x in spatial_rows))}",
        f"{pct(statistics.median(x['high_density_v20_recall'] for x in phot_rows))}",
        f"{statistics.median(100*(s['high_density_v20_recall']-p['high_density_v20_recall']) for s,p in zip(spatial_rows, phot_rows)):+.1f}",
        f"{statistics.median(x['astrometric_rms_mas'] for x in spatial_rows):.2f} / {statistics.median(x['astrometric_rms_mas'] for x in phot_rows):.2f}",
        f"{statistics.median(x['photometric_rms_mag'] for x in spatial_rows):.3f} / {statistics.median(x['photometric_rms_mag'] for x in phot_rows):.3f}",
    ]
    cells = table.add_row().cells
    for cell, value in zip(cells, values): set_cell(cell, value, size=7.2)
    for ridx, row in enumerate(table.rows):
        trpr = row._tr.get_or_add_trPr(); trpr.append(OxmlElement("w:cantSplit"))
        if ridx == 0: trpr.append(OxmlElement("w:tblHeader"))
    apply_three_line_table(table)
    return table


def reorder_references(doc):
    global CURRENT_REF_ANCHORS
    refs_heading = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "References")
    paragraphs = doc.paragraphs
    ref_paras = paragraphs[refs_heading + 1:]
    ordered = []
    for prefix in REFERENCE_ORDER:
        matches = [p for p in ref_paras if p.text.startswith(prefix)]
        if len(matches) != 1:
            raise RuntimeError(f"Expected exactly one reference beginning {prefix!r}; found {len(matches)}")
        ordered.append(matches[0])
    # APA requires year suffixes for the two distinct Yan first-author papers.
    # The source manuscript predates their joint use and therefore has no a/b
    # suffixes yet.
    for p, suffix in ((ordered[26], "a"), (ordered[52], "b")):
        plain_rewrite(p, p.text.replace("(2026).", f"(2026{suffix}).", 1))
    for p in ref_paras:
        p._p.getparent().remove(p._p)
    anchor = paragraphs[refs_heading]._p
    for p in ordered:
        anchor.addnext(p._p)
        anchor = p._p
    # Normalize every reference paragraph.  The source mixed justified and
    # left-aligned paragraphs, which stretched spaces between Latin words;
    # APA references use a consistent hanging indent and left alignment.
    for p in ordered:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.20)
        p.paragraph_format.first_line_indent = Inches(-0.20)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        # Explicit English language tagging keeps Western punctuation and DOI
        # strings breakable instead of applying East-Asian distributed spacing.
        lang = pPr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            pPr.append(lang)
        lang.set(qn("w:val"), "en-US")
    settings = doc.settings._element
    if settings.find(qn("w:wordWrap")) is None:
        word_wrap = OxmlElement("w:wordWrap")
        word_wrap.set(qn("w:val"), "on")
        settings.append(word_wrap)
    # Add stable internal targets to every bibliography entry and map the
    # rendered citation spellings to those targets.
    CURRENT_REF_ANCHORS = {}
    for idx, (prefix, citation) in enumerate(zip(REFERENCE_ORDER, CITATION_SEQUENCE), 1):
        actual_prefix = prefix
        if idx == 27:
            actual_prefix = prefix.replace("(2026)", "(2026a)")
        elif idx == 53:
            actual_prefix = prefix.replace("(2026)", "(2026b)")
        ref_p = next(p for p in doc.paragraphs if p.text.startswith(actual_prefix))
        bookmark(ref_p, f"ref_{idx}", 4000 + idx)
        CURRENT_REF_ANCHORS[citation] = f"ref_{idx}"
    # linked_rewrite() adds comma-year aliases for parenthetical citations.
    for label, target in list(CURRENT_REF_ANCHORS.items()):
        m = re.fullmatch(r"(.+) \(((?:19|20)\d{2}[a-z]?)\)", label)
        if m:
            CURRENT_REF_ANCHORS.setdefault(f"{m.group(1)}, {m.group(2)}", target)


def main():
    doc = Document(SRC)
    set_main_title(doc)
    label_registered_csst_branch(doc)
    normalize_a4_layout(doc)
    reorder_references(doc)
    base, lit = load_rows()
    cross_field_table = doc.tables[7]
    enrich_chip_characteristics_table(doc)

    linked_rewrite(pfind(doc, "We present AstroCFR"),
        "We present AstroCFR, a single-image crowded-field measurement pipeline with modular candidate-recovery and measurement branches. The primary evidence is a controlled comparison across 11 public HST/ACS F606W globular-cluster fields; four CSST-like simulations provide a separate simulation-domain reference for assessing the gap between known-truth synthetic data and real observations. Across the 11 HST fields, the spatial-ePSF joint branch has a median high-density (V<=20, at least three neighbours within 10 pixels) recovery of 86.3%, versus 45.9% for Photutils, and is higher in all 11 fields. Its position RMS is lower than Photutils in 7/11 fields and its magnitude RMS is lower in 8/11 fields. The registered CSST full-frame audit is reported as an integration check, while the method-complete CSST crop audit is reported as feasibility evidence rather than a second primary leaderboard. The evidence therefore supports a conditional high-density recovery and measurement advantage in real crowded fields, not a universal SOTA claim over Photutils or multi-exposure DOLPHOT/ALLFRAME pipelines.", LINKS)

    plain_rewrite(pfind(doc, "Crowded stellar-field photometry"),
        "High-density stellar-field photometry is a difficult source-extraction and measurement problem. When the separation of neighbouring stars approaches the PSF width, source detection, deblending, centroiding, and flux estimation become one coupled inference task. The single-image constraint is operationally necessary in three common cases: a historical archive may contain only one calibrated exposure for a sky position; a time-domain or transient pipeline must issue a first-frame crowded-field measurement before later visits arrive; and a CSST-like quick-look system must assess each exposure for quality and candidate triggering before visit-level registration and multi-exposure stacking are complete. AstroCFR supports two explicitly different operating roles: a high-recall front-end uses the ePSF/deblend branch with permissive candidate retention for triage before later stacking, whereas a single-frame fallback catalogue uses the conservative Photutils or residual-gated ePSF release threshold when precision is required and no later exposure is available. The threshold and branch must therefore be selected for the role; AstroCFR is not a claim that one exposure should replace DOLPHOT/ALLFRAME when a homogeneous exposure sequence is available. The scientific requirement is reliable recovery of stars whose PSF cores overlap under the disclosed single-frame conditions, rather than the lowest RMS among already isolated detections.")
    linked_rewrite(pfind(doc, "The China Space Station Telescope"),
        "The China Space Station Telescope (CSST) motivates a controlled simulation tier because its wide-field imaging will contain spatially heterogeneous dense stellar fields. We use four registered CSST-like chips with supplied reference catalogues, but treat them as a simulation-to-reality reference rather than as the primary performance sample. The primary observational evidence is provided by 11 public ACSGGCT HST/ACS F606W globular-cluster fields (NGC 2808, 5286, 6388, 6441, 0104, 0362, 6093, 6624, 6397, 6752, and 1851), which provide a deliberately difficult single-image test of real crowding rather than a claim about a complete multi-exposure reduction pipeline.", LINKS)
    single_image_note = insert_after(pfind(doc, "The China Space Station Telescope"))
    plain_rewrite(single_image_note, "When a homogeneous multi-exposure set is available, DOLPHOT/ALLFRAME-class methods remain the appropriate reference backend. The single-image mode is intended for archival frames, rapid first-frame triage, and CSST-like per-exposure quality control before visit-level registration and stacking are complete.")
    plain_rewrite(pfind(doc, "AstroCFR is therefore a modular"),
        "AstroCFR is a modular crowded-field processing framework designed to expose candidate recovery, measurement precision, target-specific calibration, and computational cost as distinct operating choices. The contribution is not a claim that its individual ingredients are new; it is the reproducible integration of spatially varying empirical PSFs, neighbour-aware joint fitting, residual-gated companion acceptance, and held-out cross-field evaluation in a single-image setting. Its primary 11-field HST evidence concerns image-only detection and ePSF-based measurement branches. AstroCFR-RF is an optional candidate-screen module evaluated in the CSST simulation domain and in separate bounded HST target-adaptation experiments; it is not a numerical branch of the 11-field image-only comparison.")
    psf_intro = pfind(doc, "Anderson & King (2006) developed")
    linked_rewrite(psf_intro,
        "Empirical PSF modelling is central to crowded HST photometry. Anderson & King (2000) established the empirical-PSF basis for high-precision undersampled stellar measurements, and Anderson & King (2006) specified the ACS/WFC ePSF framework. Anderson et al. (2008) applied this methodology to construct the ACS globular-cluster catalogues used here for evaluation; the associated ACS calibration context is described by Sirianni et al. (2005), while Bellini et al. (2011) demonstrates the importance of calibrated geometric-distortion treatment. Wide-field catalogues and astrometric frames from Chambers et al. (2016), Dey et al. (2019), Gaia Collaboration et al. (2016), and Gaia Collaboration et al. (2023) provide broader survey context. The reproducible analysis stack uses Astropy Collaboration (2022), NumPy (Harris et al., 2020), Matplotlib (Hunter, 2007), and SciPy (Virtanen et al., 2020).", LINKS)
    linked_rewrite(pfind(doc, "The relevant methodological map"),
        "The direct methodological lineage is compact. DAOPHOT (Stetson, 1987), SExtractor (Bertin & Arnouts, 1996), SEP (Barbary, 2016), Photutils (Bradley et al., 2024), and multi-exposure DOLPHOT-style photometry (Dolphin, 2000) establish the principal detection and PSF-fitting baselines. SCARLET (Melchior et al., 2018) and the dense Galactic-plane products of Schlafly et al. (2018) provide complementary approaches to blend separation and survey-scale crowding. AstroCFR does not seek to replace multi-exposure DOLPHOT-class measurement; it tests a single-image, high-density operating point.", LINKS)
    recent_context = pfind(doc, "Recent detector-level work")
    linked_rewrite(recent_context,
        "For CSST preparation and survey-domain context, Wang et al. (2024) studied dense-field preparation using survey-calibrated image products, while Shi et al. (2024), Zhang et al. (2023), and Zhang et al. (2024) addressed related stellar-parameter or classification tasks. AstroCFR is not positioned as a replacement for such final multi-exposure preparation: it is a proposed per-frame front end for crowded-source quality flags, provisional measurements, and trigger/quick-look candidate triage before a visit-level stack and final catalogue are available. Recent detection and density-estimation studies include Long et al. (2025), Han et al. (2026), Yan et al. (2026a), and Lai et al. (2026). These studies motivate the problem setting; their data products and targets are not numerically interchangeable with the present single-F606W experiment.", LINKS)
    psf_context = insert_after(recent_context)
    psf_context.paragraph_format.first_line_indent = Inches(0.35)
    linked_rewrite(psf_context,
        "For PSF modelling and difficult-blend controls, Burke et al. (2019) and Shaw et al. (2025) provide detection and deblending perspectives; the effective-PSF comparison of Libralato et al. (2024) and the multi-Gaussian model of Nie et al. (2025) are closest to the PSF representations examined here. Related work considers fitting uncertainty, wavefront recovery, CSST PSF reconstruction, deconvolved-image photometry, and physics-informed aperture learning, as described by Espinosa et al. (2025), Centofanti et al. (2026), Wang et al. (2026), Zhang et al. (2026), and De Alba et al. (2026). The Rubin analysis of Wainer et al. (2025), the NGC 6752 study of Salaris et al. (2024), and broader HST catalogue programmes by Dalcanton et al. (2009), Dalcanton et al. (2012), Williams et al. (2021), and Sabbi et al. (2016) further show why catalogue construction, filter choice, and exposure strategy bound valid comparisons.", LINKS)
    classifier_context = insert_after(psf_context)
    classifier_context.paragraph_format.first_line_indent = Inches(0.35)
    linked_rewrite(classifier_context,
        "For classifier context, domain-shift and representation-learning studies by Ben-David et al. (2010), Pan & Yang (2010), Ganin et al. (2016), and Long et al. (2015), together with Random Forests (Breiman, 2001), XGBoost (Chen & Guestrin, 2016), scikit-learn (Pedregosa et al., 2011), deep-learning foundations (Goodfellow et al., 2016), residual-network architecture context (He et al., 2016), and the recent Astro-RetinaNet crowded-field detector (Yan et al., 2026a), inform the archived auxiliary analyses. The cited residual-network and Astro-RetinaNet papers are context for optional learned classifiers, not claims that either was a baseline in the primary HST comparison. AstroCFR's contribution is narrower: a calibration-aware, single-image high-density protocol in which broad proposals are screened, blend groups are fitted with empirical PSFs, residual companions are recovered, and spatial-ePSF joint fitting is assessed on held-out crowded references.", LINKS)
    csst_protocol = pfind(doc, "Validation uses four simulated CSST detector chips")
    plain_rewrite(csst_protocol, "Validation is anchored by 11 public ACSGGCT HST/ACS F606W fields and complemented by four simulated CSST detector chips, labeled 12, 13, 17, and 18. Each CSST chip contains approximately 1000 reference point sources with known positions and magnitudes and is used as a controlled simulation-domain reference. Each HST field is evaluated against an external quality-selected catalogue only after image-only candidate generation and fitting. The primary HST subset is defined by V<=20 references with at least three quality-selected neighbours within 10 pixels.")
    csst_boundary = insert_after(csst_protocol)
    csst_boundary.paragraph_format.first_line_indent = Inches(0.35)
    plain_rewrite(csst_boundary, "The CSST and HST tiers deliberately answer different questions. The registered CSST audit is a full-frame supplied-catalogue integration test of calibrated SExtractor and AstroCFR ePSF + residual deblend (registered full-frame branch). In addition, we performed a controlled-crop audit in which all five image-only branches were run on identical 1200 x 1200 crops from chips 12, 13, 17, and 18. The crop audit is reported separately in Supplementary Table S27: it is method-complete but has only 8 and 10 test references on chips 17 and 18, so it supports protocol feasibility rather than a full-frame CSST benchmark. The two CSST products are not pooled with the 11-field HST medians, and the crop audit does not replace the registered full-frame benchmark.")
    csst_role = insert_after(csst_boundary, style="Normal")
    plain_rewrite(csst_role, "The CSST tier contributes two complementary checks. The full-frame audit verifies detector-scale coordinate registration, zero-point calibration, catalogue export, and residual-deblend bookkeeping on a 9232 x 9216 image; it is an integration and reproducibility result. The controlled crops remove that integration advantage and place five branches on identical footprints, making recovery and conditional RMS comparable within the simulation domain. CSST is therefore retained to expose the simulation-to-reality and full-frame-to-crop gap, not as a second claim of observational superiority.")
    registered_note = insert_after(csst_boundary, style="Normal")
    plain_rewrite(registered_note, "Here, registered full-frame branch has a narrow technical meaning: the archived AstroCFR ePSF + residual-deblend workflow was run on the complete 9232 x 9216 simulated detector frame and evaluated against the supplied per-chip catalogue after per-chip coordinate registration and photometric zero-point calibration. It is not the spatial-ePSF joint branch, it is not a separate physical reference frame, and it is not an independent pipeline family. The label distinguishes this archived full-frame integration audit from the newer 1200 x 1200 controlled-crop five-method experiment.")
    crop_note = insert_after(csst_boundary, style="Normal")
    plain_rewrite(crop_note, "In the controlled CSST crop audit, median recovery was 70.5% for spatial-ePSF joint fitting versus 63.3% for Photutils, and median position RMS was 6.41 versus 9.35 mas; Photutils retained the lower median magnitude RMS (0.0283 versus 0.0463 mag). This mixed recovery–precision pattern is reported as a simulation-domain operating-point trade-off.")
    f606w_note = insert_after(crop_note, style="Normal")
    plain_rewrite(f606w_note, "F606W is used as a deliberately homogeneous first validation band: all 11 primary fields share the same ACS/WFC filter and a common catalogue-quality protocol, allowing crowding and PSF representation to be compared without conflating filter-dependent throughput or colour terms. This single-band design is not intended to establish multiband generality. Extension to F814W and to jointly fitted multi-exposure data is a defined next step.")
    plain_rewrite(pfind(doc, "The problem is framed as dense-field source detection"), "The problem is framed as high-density source detection and calibration. Given a noisy image and an evaluation catalogue, the system produces source positions and magnitudes. Recovery and RMS are reported on held-out spatial partitions; reference catalogues are never used to generate candidates, construct the empirical PSF, or fit the image. CSST and HST are separate evidence tiers and are not averaged into one leaderboard.")
    decision = pfind(doc, "For a branch b, we report")
    plain_rewrite(decision, "For a branch b, we report the decision vector [[WORD_INLINE_DECISION_VECTOR]], where C is held-out catalogue recovery, R_dense is recovery of V<=20 references with at least three neighbours within 10 pixels, E_pos and E_mag are held-out positional and photometric RMS, and T and M are runtime and memory cost. Equations (1)–(3) formalize the high-density recovery, joint-fit objective, and conditional position error used throughout the HST tier.")
    eq1 = equation_after(doc, decision._p, 1)
    eq2 = equation_after(doc, eq1._p, 2)
    equation_after(doc, eq2._p, 3)
    plain_rewrite(pfind(doc, "Release selection is then science-constrained"), "Release selection is science-constrained rather than globally ranked: DAOStarFinder or SEP are appropriate for rapid screening, Photutils remains a strong conditional-precision branch, AstroCFR ePSF-deblend is the recovery-oriented compromise, and spatial-ePSF joint fitting is preferred when dense-field completeness and blend-aware positions are the priority. Equation (2) is optimized independently within each blend group G, while Equations (1) and (3) are evaluated only on the held-out reference partition. The selected operating point must be reported with its computational cost and evidence tier.")
    plain_rewrite(pfind(doc, "The quality gates have an explicit order"), "The quality gates have an explicit order. Broad image-only proposals are screened first by morphology and SNR. Candidate probabilities define the conservative RF catalogue, while residual-improvement tests decide whether a blended companion is retained. Astrometric and photometric corrections are accepted only after validation improvement. When the classifier and residual test disagree, the residual test governs companion acceptance; the classifier threshold governs catalogue release. Reference catalogues are used only for the disclosed simulation evaluation and target-adaptation calibration partitions, not as an inference-time gate in the HST image-only branches.")
    plain_rewrite(pfind(doc, "The previous ePSF branch used one image-derived PSF"), "The previous ePSF branch used one image-derived PSF for the crop. The high-density extension estimates empirical PSFs separately in four spatial quadrants and performs two neighbour-aware coordinate/flux fitting passes. This is the AstroCFR spatial-ePSF joint branch. For literature mapping, the same image-only harness also fits a global empirical ePSF and a three-Gaussian discrete PSF, allowing the effect of spatial variation and PSF parameterisation to be separated from the candidate front end.")
    plain_rewrite(pfind(doc, "The deblending module also changes how the classifier should be interpreted"),
        "The deblending module also changes how the classifier should be interpreted. A candidate that appears ambiguous in isolation may become acceptable after the neighbouring-source model explains a shared PSF core. Conversely, a visually plausible candidate can be rejected if adding it does not reduce the local residual structure. Grouping radius and residual-acceptance settings are registered hyperparameters, and their sensitivity is reported separately in Supplementary Tables S5–S6; they must be revalidated for a new instrument configuration.")
    rl_note = insert_after(pfind(doc, "The deblending module also changes how the classifier should be interpreted"), style="Normal")
    plain_rewrite(rl_note, "Richardson–Lucy restoration is retained only as an optional diagnostic for severe or saturated groups in the legacy three-level deblending path. It is not the defining operation of the primary 11-field HST spatial-ePSF comparison, and no headline recovery or RMS result is attributed to Richardson–Lucy alone. The reported operating points are driven by empirical-PSF fitting, neighbour-aware joint fitting, and residual-gated acceptance. A dedicated iteration-count and noise-amplification ablation remains necessary before restoration is promoted as a general point-source measurement component.")
    remove_paragraph(pfind(doc, "The dominant dense-field failure mode is structural blending. Candidate groups are processed"))
    plain_rewrite(pfind(doc, "Astrometric refinement uses a validation-gated cascade"),
        "Astrometric refinement uses a validation-gated cascade: a polynomial removes large-scale WCS distortion, a look-up table models residual spatial structure, and Gaussian-process correction is retained only when it improves held-out residuals. Polynomial degree, grid settings, and local centroid refinement are registered configuration choices; none is treated as an instrument-independent optimum.")
    remove_paragraph(pfind(doc, "Astrometric refinement uses a three-stage cascade"))
    plain_rewrite(pfind(doc, "Detected sources are matched to supplied references"), "Detected sources are matched to supplied references using a greedy one-to-one two-pixel rule. In the HST tier, the catalogue is held out by spatial stripes and the high-density subset is defined independently of candidate generation. In the CSST tier, the registered full-frame supplied-catalogue audit is retained as a separate simulation result. These denominators are not interchangeable.")
    protocol_anchor = pfind(doc, "4.1 Classifier benchmark")
    protocol_1 = insert_after(protocol_anchor)
    protocol_1.paragraph_format.first_line_indent = Inches(0.35)
    plain_rewrite(protocol_1, "Training, validation, and test handling is protocol-specific. The primary 11-field HST spatial-ePSF/Photutils comparison is not a supervised-training experiment: all branches generate and fit candidates from the image, and the reference catalogue is used only after fitting for held-out scoring. Each 1200 x 1200 crop is divided into a deterministic 3 x 3 spatial grid of 200-pixel cells, not a random catalogue split. For the archived WPDC classifier configuration, cell partition 0 is the training/calibration partition, partition 1 selects the classifier threshold, and partition 2 is the final test partition. No HST RF score is reported by design, not because RF is presumed to fail: the released RF artefact is a CSST-simulation model whose labels, feature normalisation, and threshold are tied to simulated truth. Applying it unchanged to HST would be an unvalidated domain transfer; retraining it on the ACSGGCT evaluation catalogue would instead create a supervised target-specific adaptation experiment. Such a result requires a separately pre-registered cross-field or held-out-field protocol and is not reconstructed during final archival. DAO, SEP, Photutils, ePSF-deblend, and spatial-ePSF joint fitting use no catalogue labels for candidate generation or PSF fitting. Astrometric registration and photometric zero points are estimated from non-test matches and evaluated on partition 2. The spatial-ePSF is estimated from isolated proposals across the science crop, so it can see unlabeled test pixels; this transductive calibration is disclosed separately from catalogue-held-out scoring.")
    protocol_1_text = protocol_1.text.replace("not because RF is presumed to fail: the released", "not because RF is presumed to fail. The released")
    protocol_1_first, protocol_1_rest = protocol_1_text.split(" Applying it unchanged to HST would be", 1)
    plain_rewrite(protocol_1, protocol_1_first.rstrip(".") + ".")
    protocol_1b = insert_after(protocol_1, style="Normal")
    protocol_1b.paragraph_format.first_line_indent = Inches(0.35)
    plain_rewrite(protocol_1b, "Applying it unchanged to HST would be" + protocol_1_rest)
    transductive_note = insert_after(protocol_1, style="Normal")
    plain_rewrite(transductive_note, "The empirical PSF is estimated from image-only isolated DAO proposals in the science crop, so the primary branch permits unlabeled test-partition pixels to contribute to image calibration. This is transductive calibration, not reference-catalogue leakage: no held-out reference position, magnitude, match assignment, or quality flag is used to choose PSF stamps or fit sources. Isolated proposal stamps can be spatially unrepresentative of the crowded test region if the PSF varies within a quadrant, so we ran a partition-isolated sensitivity audit in which only DAO proposals from cells 0/1 supplied global and quadrant ePSF stamps; full-image candidates were then fitted and partition 2 was scored unchanged. Across all 11 fields, dense-recovery change had median -0.13 percentage points (range -2.33 to +0.45 pp), position-RMS change had median +0.09 mas (range -0.05 to +0.34 mas), and magnitude-RMS change had median +0.0033 mag (range -0.0014 to +0.0087 mag). The largest dense-recovery decrease occurred in NGC 5286 (-2.33 pp); the headline 11/11 recovery ordering was unchanged. These small but non-zero shifts quantify the transductive sensitivity under the present spatial protocol; a stricter multi-fold PSF-stamp audit remains useful for future work.")
    transductive_result = insert_after(transductive_note, style="Normal")
    plain_rewrite(transductive_result, "The partition-isolated audit is a robustness check rather than a new leaderboard: it changes the ePSF construction only, while candidate generation, neighbour fitting, registration, association radius, and held-out scoring remain fixed. The machine-readable per-field deltas are archived in results/partition_isolated_spatial_epsf_sensitivity_all11/summary.json and comparison.csv.")
    protocol_2 = insert_after(protocol_1)
    protocol_2.paragraph_format.first_line_indent = Inches(0.35)
    plain_rewrite(protocol_2, "The reproducible CSST candidate-classifier ablation uses 11,615 detector-generated candidates from chips 12, 13, 17, and 18: 3,615 positives within 2 pixels of a supplied simulated reference and 8,000 negatives farther than 8 pixels; the ambiguous 2-8 pixel annulus is excluded. The 2-pixel positive radius is the same one-to-one association radius used throughout the manuscript and is smaller than the 3-pixel FWHM assumed by the detector, whereas 8 pixels is a conservative separation from a catalogue source. Negatives are not random blank-sky positions: they are image-detector proposals whose nearest supplied reference lies beyond 8 pixels. Per chip, they are sampled without replacement from that distant-proposal pool at up to five negatives per positive and capped at 2,000, yielding the reported 8,000 negatives. Thus this is a deliberately clean simulation-domain architecture ablation, not a claim that its negative distribution reproduces real crowded-field false detections. A stratified 60/20/20 train/validation/test split with seed 20260806 is fixed before training. The feature standardizer is fitted on the training split only. The validation split selects the highest threshold retaining at least 90% recall, and all accuracy metrics use the untouched test split. RandomForest uses 400 trees, maximum depth 15, minimum leaf size 2, sqrt feature sampling, class weights {negative: 1, positive: 6}, and the same seed. The CNN and patch Transformer use Adam (learning rate 1e-3, weight decay 1e-4), cosine annealing over 30 epochs, batch size 64 for training and 256 for inference; the Transformer uses 25 non-overlapping 5 x 5 patches, 64-dimensional tokens, two four-head encoder layers, and a 17-feature fusion branch.")
    protocol_2_full = protocol_2.text.replace("separation from a catalogue source. Negatives", "separation from a catalogue source. Negatives")
    protocol_2_first, protocol_2_rest = protocol_2_full.split(" Negatives are not random blank-sky positions:", 1)
    plain_rewrite(protocol_2, protocol_2_first.rstrip(".") + ".")
    protocol_2b = insert_after(protocol_2)
    protocol_2b.paragraph_format.first_line_indent = Inches(0.35)
    plain_rewrite(protocol_2b, "Negatives are not random blank-sky positions:" + protocol_2_rest)
    protocol_3 = insert_after(protocol_2)
    protocol_3.paragraph_format.first_line_indent = Inches(0.35)
    plain_rewrite(protocol_3, "A separate closed-book leave-one-chip-out RF pilot trains on the other three simulation chips, reserves a stratified 25% of those training candidates for threshold selection, and opens the held-out chip catalogue only after the final candidate file is written. Its RF uses 200 trees, maximum depth 15, minimum leaf size 2, sqrt feature sampling, class weights {negative: 1, positive: 6}, six CPU workers, and seed 20260811 plus the held-out chip identifier. This pilot is reported only as a blind-proposal stress test and is not pooled with the 11-field HST measurement result.")
    classifier_anchor = pfind(doc, "The classifier comparison uses identical exported candidate pools,")
    plain_rewrite(classifier_anchor, "The common-method comparison uses identical image crops, spatially held-out partitions, a two-pixel association radius, and method-specific image-only candidate generation. DAOStarFinder, SEP/SExtractor-style extraction, Photutils PSFPhotometry, AstroCFR ePSF-deblend, and AstroCFR spatial-ePSF joint fitting are rerun on all 11 ACSGGCT fields. RF is explicitly recorded as protocol-excluded: a CSST-trained screen is not transferred as an HST result, and an HST-catalogue-trained RF would be a different supervised adaptation experiment rather than an image-only branch in this comparison.")
    method_config = insert_after(classifier_anchor, style="Normal")
    plain_rewrite(method_config, "The baseline configurations are explicit. DAOStarFinder uses a 3-pixel FWHM and 3-sigma threshold for image-only centroid/flux proposals; it is not followed by a separate PSFEx model. SEP uses SExtractor-style thresholding at 3 background RMS with minimum area 5, 32 deblending thresholds, and deblend contrast 0.005; its reported flux is calibrated extraction flux, not PSFEx photometry. Photutils uses the same DAOStarFinder proposal frontend but fits a circular Gaussian PRF in a 9 x 9 pixel window with 2-pixel grouping separation and 3-pixel aperture initialisation. AstroCFR global and spatial branches build empirical PSFs from isolated DAO proposals, fit flux and position jointly within blend groups, and apply the same residual-acceptance and held-out calibration rules. Thus DAO/SEP are fast detection/extraction baselines, Photutils is a fixed analytic-PRF PSF-fitting baseline, and AstroCFR is spatially varying empirical-PSF joint fitting; no branch is described as a PSFEx or multi-exposure DOLPHOT reduction.")
    plain_rewrite(pfind(doc, "RandomForest, XGBoost, and hybrid stacking classifiers"), "The literature-mapped controls are diagnostic PSF-representation controls, not additional end-to-end published-pipeline leaderboards. They share the HST image, candidate front end, spatial hold-out, and neighbour-fitting harness, but the archived outputs retain a dense held-out reference denominator rather than the baseline branch's full held-out metric set. They are therefore used to test the direction of the spatial-PSF effect and are reported in Supplementary Table S26 with their own denominator; they are not pooled into Table 8 medians or used to rank the main operating points. An AstroCFR+Photutils hybrid remains an archived three-field ablation and is not promoted to an 11-field claim.")
    sextractor_scope = pfind(doc, "The SExtractor comparison is intentionally not framed")
    plain_rewrite(sextractor_scope, "The SExtractor comparison has a narrower purpose. The registered CSST audit is full-frame and supplied-catalogue conditioned, whereas the HST comparison is a single F606W stack evaluated against an external catalogue. DOLPHOT/ALLFRAME, crowdsource, Euclid/VVV workflows, and CSST-PSFNet are retained as scope audits because their required multi-exposure, multiband, instrument-specific, or labelled-PSF inputs are absent.")
    csst_fairness = insert_after(sextractor_scope, style="Normal")
    plain_rewrite(csst_fairness, "The CSST full-frame magnitude-RMS contrast must be read as a conditional integration-audit diagnostic, not as an estimator-identical photometry contest. SExtractor supplies a catalogue-extraction magnitude under its calibrated detection configuration, whereas AstroCFR reports a local empirical-PSF fit after residual deblending. Both recovery values use the supplied-catalogue association rule, but each RMS is computed only on that method's successfully matched references after its own registration and zero-point calibration. Thus a lower AstroCFR RMS can reflect both the fitted flux model and the method-specific recovered subset; it does not establish a same-star, same-estimator advantage. The controlled HST analysis therefore reports recovery and conditional RMS jointly, and the separate CSST crop audit includes Photutils PSFPhotometry under one common protocol.")
    plain_rewrite(pfind(doc, "Validation-selected per-chip calibration was necessary"), "Validation-selected per-chip calibration was necessary for heterogeneous simulated chips. In the registered full-frame CSST audit, calibrated SExtractor reaches 83.7–91.6% recovery across chips, while AstroCFR ePSF + residual deblend reaches 91.8–96.9%. The median AstroCFR advantage is 6.4 percentage points in recovery, 15.3 mas in position RMS, and 0.335 mag in magnitude RMS. These are registered CSST-tier values, not a replacement for the HST high-density experiment.")
    chip18_note = insert_after(pfind(doc, "Validation-selected per-chip calibration was necessary"), style="Normal")
    plain_rewrite(chip18_note, "Chip 18 illustrates why the qualitative difficulty label cannot be read as a monotonic recovery ranking. It has the largest background RMS (5.36 DN) and broadest fitted PSF (1.93 pixels), but its supplied top-1000 catalogue has only 7.7 stars arcmin^-2, tied for the lowest chip-level reference density. Its bright magnitude range also supplies many high-SNR detections. Thus bright-star artifacts and blends stress local modelling, while the lower supplied-reference density and high SNR can still yield high catalogue recovery after empirical-PSF fitting; the table's difficulty labels describe dominant image characteristics, not an ordinal expected-recall scale.")
    chip_table_note = insert_after(chip18_note, style="Normal")
    plain_rewrite(chip_table_note, "The density values in Table 1 are catalogue-surface densities computed from the 1,000 supplied references divided by each catalogue's detector-coordinate bounding area, using the CSST-like 0.074 arcsec pixel scale. They are descriptive densities of the supplied reference layer, not exhaustive scene densities; the magnitude span is the corresponding supplied-catalogue range.")
    plain_rewrite(pfind(doc, "On the CSST-like simulations, calibrated SExtractor"), "On the CSST-like simulations, calibrated SExtractor reaches 83.7–91.6% supplied-catalogue recovery, while the registered AstroCFR ePSF + residual deblend branch reaches 91.8–96.9%. This branch has a median full-frame position RMS of 8.95 mas versus 24.25 mas for SExtractor, and a median magnitude RMS of 0.064 versus 0.399 mag. The comparison remains diagnostic rather than a universal ranking because the CSST reference catalogue and assembly protocol define the evidence tier.")
    evidence_anchor = pfind(doc, "The reported evidence comprises four CSST-like simulated chips")
    plain_rewrite(evidence_anchor, "The primary evidence comprises 11 real ACSGGCT HST/ACS F606W fields, all with the five common image-only branches; the two literature-mapped PSF controls are also complete for all 11 fields. Four CSST-like full-frame chips and the method-complete controlled-crop audit are retained as a separate simulation-to-reality reference tier. The hybrid branch is retained only for its three archived fields. RF is protocol-excluded from the HST matrix because neither unvalidated CSST-to-HST transfer nor a new HST-catalogue-supervised adaptation is treated as an image-only comparison branch.")
    plain_rewrite(pfind(doc, "The eight included real expansion fields"), "The 11-field HST comparison is the primary high-density observational result. Across fields, spatial-ePSF joint fitting has higher V<=20 high-density recovery than Photutils in 11/11 fields, lower position RMS in 7/11, and lower magnitude RMS in 8/11. The field medians are 86.3% versus 45.9% for dense recovery, 2.44 versus 3.90 mas for position RMS, and 0.088 versus 0.179 mag for magnitude RMS. Thus the spatial branch is not globally photometrically inferior: it has lower magnitude RMS in 8/11 fields; its limitation is field- and subset-dependent, with Photutils retaining lower conditional magnitude RMS in the remaining three fields and in some isolated/precision-oriented operating points. These medians summarize 15,893 high-density references in total (field n range 9–3,307; median n=1,334). NGC 6397 contributes only nine high-density references and is therefore descriptive rather than decisive; excluding it leaves 15,884 references. A descriptive reference-weighted recovery is 67.3% for spatial-ePSF versus 34.7% for Photutils, while the field-median analysis remains the prespecified headline summary. The literature controls use a different denominator and are not pooled into these medians.")
    dolphot_results_section = insert_after(evidence_anchor, style="Normal")
    plain_rewrite(dolphot_results_section, "For an explicit multi-exposure reference, native DOLPHOT FakeStars on three registered PHAT M31 F475W FLC exposures recovered 177/200 low-density injections (88.5%, 95% CI 83.3–92.2%) and 181/200 high-density injections (90.5%, 95% CI 85.6–93.8%) at F475W=26.5, with magnitude RMS 0.201 and 0.197 mag. This is a measured physical-backend comparison, not a literature-only citation. It is reported as a separate tier because AstroCFR uses one F606W DRZ image, a different renderer, filter, catalogue, and injection protocol; the result establishes the expected multi-exposure reference strength and defines AstroCFR's complementary single-frame use case rather than an input-mismatched ranking.")
    plain_rewrite(pfind(doc, "M33 B01-F01 and B03-F02"), "The earlier non-ACSGGCT comparisons remain contextual and are not merged into the 11-field high-density median. They use different filters, reference catalogues, and crowding distributions. GR8 is retained in Supplementary Table S25 as a separately labelled multi-exposure SCI-mosaic processing audit; its mosaic construction, F475W bandpass, and ANGST/GST reference are not input-equivalent to the ACSGGCT single-F606W tier. S25 therefore documents this boundary and its field-level metrics without introducing an unpooled additional validation claim. The present primary claim is deliberately restricted to the matched ACSGGCT F606W protocol; CSST results are retained as a separate simulation-to-reality reference.")
    table_reference_rewrite(pfind(doc, "Table 8 compares every measured branch"), "Table 8 is intentionally a compact paired comparison between AstroCFR spatial-ePSF joint fitting and Photutils, the two branches most relevant to the central high-density claim. It gives all 11 field-level high-density recoveries and conditional RMS values. The full 78-row comparison, including DAO, SEP, ePSF-deblend, global empirical-ePSF, and three-Gaussian dPSF controls, is moved to Supplementary Table S26 and to the machine-readable all-method matrix.")
    plain_rewrite(pfind(doc, "The controlled measurements expose distinct deployment operating points"), "The controlled measurements expose distinct deployment operating points rather than a single winner. On the primary 11-field HST tier, spatial-ePSF joint fitting is the recovery-oriented high-density mode, Photutils remains the most reliable precision-oriented baseline in several fields, and ePSF-deblend provides an intermediate cost–recovery point. The CSST simulation tier is used to expose simulation-to-reality differences and to audit integration feasibility, not to establish a second primary ranking.")
    difficult_note = insert_after(pfind(doc, "The controlled measurements expose distinct deployment operating points"), style="Normal")
    plain_rewrite(difficult_note, "The NGC 1851 learning curve is difficult for measurable field reasons rather than an unreported algorithmic exception: its image-derived FWHM is broader, local background and blend residuals are more heterogeneous, and the conservative negative pool is less stable under one-tile target calibration. Multiple spatial tiles are therefore needed before threshold selection stabilises. This is a field-condition diagnostic, not evidence of selective hand tuning.")
    runtime_note = insert_after(difficult_note, style="Normal")
    runtime = runtime_density_summary(base)
    plain_rewrite(runtime_note, ("Runtime was measured on the same 1,200 x 1,200 HST crops using the CPU execution path, with wall-clock time normalised by image area. "
        f"Across fields, median costs were {runtime['dao']['runtime_s_per_mpix']:.2f} s MPix⁻¹ for DAOStarFinder, "
        f"{runtime['sep']['runtime_s_per_mpix']:.2f} s MPix⁻¹ for SEP, {runtime['photutils_psf']['runtime_s_per_mpix']:.2f} s MPix⁻¹ for Photutils, "
        f"{runtime['wpdc_epsf_deblend']['runtime_s_per_mpix']:.2f} s MPix⁻¹ for ePSF-deblend, and "
        f"{runtime['wpdc_spatial_epsf_joint']['runtime_s_per_mpix']:.2f} s MPix⁻¹ for spatial-ePSF joint fitting. "
        f"The spatial branch is therefore {runtime['wpdc_spatial_epsf_joint']['runtime_s_per_mpix']/runtime['photutils_psf']['runtime_s_per_mpix']:.1f}x the Photutils cost in this implementation, while its median dense recovery is 86.3% versus 45.9%. "
        "These measurements support a tiered deployment interpretation: DAO/SEP are genuine low-latency screening choices, Photutils is a relatively fast precision branch, and spatial-ePSF is a recovery-oriented single-frame back end or high-quality front end rather than a guaranteed real-time full-detector service. The reported crop timings must not be extrapolated to CSST full-frame throughput; parallel and GPU implementations were not benchmarked."))
    strat_note = insert_after(runtime_note, style="Normal")
    plain_rewrite(strat_note, (f"The magnitude-stratified field medians also show why the headline dense-subset result is not a single easy-source statistic. "
        f"For spatial-ePSF, recovery is {100*runtime['wpdc_spatial_epsf_joint']['recall_v_le_18']:.1f}% for V<=18, {100*runtime['wpdc_spatial_epsf_joint']['recall_v_le_20']:.1f}% for V<=20, and {100*runtime['wpdc_spatial_epsf_joint']['recall_v_le_22']:.1f}% for V<=22, compared with {100*runtime['photutils_psf']['recall_v_le_18']:.1f}%, {100*runtime['photutils_psf']['recall_v_le_20']:.1f}%, and {100*runtime['photutils_psf']['recall_v_le_22']:.1f}% for Photutils. "
        "The separation widens toward the fainter cumulative limit, consistent with the intended role of residual companion recovery; it is not evidence that the spatial branch is uniformly superior on every isolated or bright-star precision subset."))
    routing_scope = insert_after(pfind(doc, "The density-adaptive router defines an intermediate recovery-cost operating point"), style="Normal")
    plain_rewrite(routing_scope, "Table 11 is deliberately a proposal-recovery and runtime experiment, not a catalogue-release comparison. Its 800 injected trials do not provide an exhaustive scene catalogue or a common matched-star set for false-discovery, precision, positional RMS, or magnitude RMS estimation. A higher routed recovery could therefore be accompanied by a different false-candidate rate, and a fast low-density branch could have different conditional measurement scatter. The router should consequently be read only as evidence of a recovery-cost trade-off; precision and measurement quality require a future common-scene, exhaustive-truth routing evaluation before it can be used as a final catalogue-selection policy.")
    plain_rewrite(pfind(doc, "The threshold experiment shows that evaluation protocol"), "The threshold experiment shows that evaluation protocol can be a hidden confound. The completed 11-field audit reduces this risk by fixing the crop, association radius, held-out spatial partition, high-density definition, and external-catalogue handling before comparing branches. Apparent gains are therefore interpreted as operating-point differences under a disclosed protocol, not as instrument-independent superiority.")
    fairness_audit = insert_after(pfind(doc, "The threshold experiment shows that evaluation protocol"), style="Normal")
    plain_rewrite(fairness_audit, "A held-out NGC 6752 robustness audit further checks that this direction is not created by a single matching or proposal threshold. With the test partition untouched, changing the common image-only DAO proposal threshold from 2.5 to 4.0 sigma changed all-quality reference recovery from 77.13% to 74.99%, while the catalogue-match lower bound changed from 93.68% to 94.62%. Holding every produced catalogue fixed and varying only the one-to-one association radius from 1 to 3 pixels left the spatial-ePSF branch above Photutils throughout: 93.19% to 93.45% versus 76.17% to 76.25% all-quality recovery. This is a configuration-robustness check, not test-set parameter selection; the manuscript headline retains the predeclared two-pixel rule and high-density definition.")
    isolated_anchor = pfind(doc, "The isolated benchmark identifies useful classifier differences")
    plain_rewrite(isolated_anchor, "The isolated benchmark identifies useful branch differences, but the full high-density result is driven by the interaction of candidate generation, empirical-PSF representation, and neighbour-aware fitting. Spatial-ePSF joint fitting recovers more crowded references than Photutils in every ACSGGCT field, while Photutils remains competitive or superior in individual conditional RMS measures. This is why the paper reports both recovery and measurement precision rather than a single score.")
    rf_ablation_note = insert_after(isolated_anchor, style="Normal")
    plain_rewrite(rf_ablation_note, "The RF ablation clarifies the division of labour: on the dense NGC 6752 subset, AstroCFR-RF recovers 52.7%, whereas the ePSF-deblend branch recovers 87.6% under the same association rule. The classifier is therefore an auxiliary catalogue-screening module; the dominant high-density gain comes from empirical-PSF residual deblending and neighbour-aware fitting, not from the RF score alone.")
    plain_rewrite(pfind(doc, "This has a practical engineering implication"), "This has a practical engineering implication. If final catalogue precision for already recovered stars is the primary objective, Photutils is a strong low-complexity choice. If missing blended companions is the dominant scientific risk, AstroCFR ePSF-deblend or spatial-ePSF joint fitting is preferable, with a substantial runtime cost. The hybrid ablation is not a reliable shortcut: replacing the ePSF measurement stage with a fixed Gaussian Photutils fit improves neither recovery nor precision consistently.")
    plain_rewrite(pfind(doc, "The result should not be read as an argument against deep learning"), "The result should not be read as an argument against deep learning or against multi-exposure photometry. It shows that, for this single-image high-density task, explicit spatial PSF structure and neighbour-aware fitting explain more of the measured recovery advantage than an isolated classifier leaderboard.")
    dolphot_note = insert_after(pfind(doc, "The result should not be read as an argument against deep learning"), style="Normal")
    linked_rewrite(dolphot_note, "DOLPHOT/ALLFRAME-class reductions are the appropriate reference when homogeneous, registered multi-exposure data are available: they exploit repeated measurements and exposure-specific PSFs that are absent from this single-image experiment (Dolphin, 2000). AstroCFR is therefore complementary rather than substitutive: it targets a disclosed single-image operating point, while a direct precision ranking against a multi-exposure reduction would require the same exposure list, detector calibration, spatial hold-outs, and artificial-star protocol.", LINKS)
    dolphot_result = insert_after(dolphot_note, style="Normal")
    plain_rewrite(dolphot_result, "A separate physical-backend reference is available in the reproducibility package but is not pooled with the HST single-image matrix: on three registered PHAT F475W FLC exposures, native DOLPHOT FakeStars recovered 177/200 (88.5%) faint low-density and 181/200 (90.5%) faint high-density injections at F475W=26.5, with magnitude RMS of 0.201 and 0.197 mag, respectively. Those values demonstrate the expected strength of a multi-exposure reduction. They are not ranked against AstroCFR because the AstroCFR independent test uses a single DRZ image, separate one-star injections, and a different renderer; the result defines the complementary use case rather than a hidden DOLPHOT comparison.")
    dolphot_public = insert_after(dolphot_result, style="Normal")
    linked_rewrite(dolphot_public, "This ordering is consistent with the public DOLPHOT literature: Dolphin (2000) describes simultaneous multi-image PSF photometry and artificial-star testing as the route to calibrated crowded-field completeness and errors. The present three-FLC FakeStars numbers provide a concrete local reference for that class of reduction, while the 11-field AstroCFR result quantifies what can be recovered before those repeated exposures are available. We therefore interpret AstroCFR as a first-frame or archival complement, not as a replacement or an unstated claim of equal multi-exposure precision.", LINKS)
    dolphot_gate_note = insert_after(dolphot_public, style="Normal")
    plain_rewrite(dolphot_gate_note, "The attempted same-field DOLPHOT run was not a post hoc choice of a favourable failure case. NGC 1851 was one of the prespecified 11 ACSGGCT fields and was selected for the audit before inspecting the retry output because six native F606W FLC exposures were available. The acceptance gate was also defined before scoring: inter-exposure alignment residuals had to remain sub-pixel and stable enough for the common 2-pixel association, train-only affine registration, and milliarcsecond-scale held-out RMS protocol. A 0.698-pixel or 1.805-pixel warning is not evidence that DOLPHOT is unusable in general; DOLPHOT can model exposure-specific transformations internally. It means only that this particular available FLC set did not yield an input-registered output that could be compared to the single-image catalogue under our declared spatial-holdout protocol. Relaxing the gate would change the coordinate-error budget and would not constitute a fair DOLPHOT ranking.")
    dao_allframe_note = insert_after(dolphot_public, style="Normal")
    plain_rewrite(dao_allframe_note, "Compared with DAOPHOT/ALLFRAME, AstroCFR is not claimed to introduce a new fundamental alternative to image-based PSF fitting; both families model overlapping stars iteratively, and DAOStarFinder is not a DAOPHOT benchmark. The differentiator tested here is a pre-registered single-image operating protocol: broad image-only proposals, residual-gated companion acceptance, spatial-quadrant empirical PSFs, two-pass neighbour fitting, and spatially held-out scoring. Its quantitative evidence is limited to that protocol: across 15,893 dense HST references in 11 F606W fields, spatial-ePSF has an 86.3% field-median dense recovery versus 45.9% for the fixed-PRF Photutils branch, with lower position RMS in 7/11 fields and lower magnitude RMS in 8/11. This is not a numerical DAOPHOT or ALLFRAME advantage. Use AstroCFR's high-recall ePSF/deblend mode for a single-frame front-end or incomplete archival exposure, and use DAOPHOT/ALLFRAME-class multi-image reduction when registered repeated exposures exist. A direct DAOPHOT single-image and ALLFRAME multi-image comparison would require a common FLC exposure set, PSF configuration, artificial-star scene, and held-out scoring protocol; it remains future work rather than an implicit claim.")
    primary_limit = pfind(doc, "The primary development experiments use CSST-like simulations")
    plain_rewrite(primary_limit, "The primary evidence is the 11-field real HST/ACS comparison. The four CSST-like full frames and the unified five-branch crop audit are retained as a complementary simulation-to-reality reference tier. The two tiers use different truth/reference constructions and are not averaged. The HST suite is still a single-image catalogue-conditioned experiment rather than a multi-exposure DOLPHOT/ALLFRAME comparison. Its observational scope is specifically F606W globular-cluster fields: it does not establish performance in F814W or other passbands, in Galactic-bulge or LMC-like crowding, in dwarf-galaxy cores, or across a broad stellar-colour distribution. The single-image setting is practically relevant for archival frames with incomplete exposure sequences, rapid candidate triage before a slower multi-exposure reduction, and time-domain or trigger workflows in which a catalogue is needed before all exposures are available. It is not intended to replace DOLPHOT when a homogeneous registered exposure set is available. RF is not reported numerically because a CSST-trained screen is not validated for HST transfer and a new HST-catalogue-trained RF would be a separate supervised-adaptation experiment; the hybrid branch has only three archived HST fields.")
    adaptation_limit = insert_after(primary_limit, style="Normal")
    linked_rewrite(adaptation_limit, "We use the term target adaptation descriptively rather than as a formal transfer-learning theorem: simulation-to-HST shifts in PSF, background, crowding, catalogue depth, and detector response invalidate an assumption of unchanged feature and label distributions (Ben-David et al. (2010); Pan & Yang (2010)). The small labelled-tile experiment is consequently a bounded recalibration study, not evidence of automatic survey-scale domain transfer or a sample-complexity guarantee. The observed zero-shot degradation is an explicit limitation of simulation-trained classification, not a validation of the CSST simulation fidelity.", LINKS)
    plain_rewrite(pfind(doc, "The primary NGC 6752 fixed-scene artificial-star curves"), "The NGC 6752 artificial-star curves remain a controlled, model-matched stress test and are not substituted for the 11-field observational matrix. The independent HST field expansion and the literature-mapped ePSF controls provide the broader high-density evidence; all are still catalogue-conditioned rather than exhaustive blind truth experiments.")
    plain_rewrite(pfind(doc, "The per-chip reference catalogs contain approximately 1000 stars"), "The CSST per-chip reference catalogues contain approximately 1000 stars, while the HST catalogues are finite quality-selected external references. The four CSST chips are consequently a controlled simulation-domain reference, not a generalization proof for the future instrument; the target-adaptation experiments explicitly demonstrate that simulation-trained screening can shift on real images. Larger and more heterogeneous simulation suites, real CSST commissioning data, and homogeneous multi-exposure HST data are required for a definitive cross-domain and multi-exposure benchmark.")
    plain_rewrite(pfind(doc, "Table 5. Ablation of spatially varying ePSFs"), "Table 5. Ablation of spatially varying ePSFs and two-pass joint fitting on NGC 6752. The refinement improves the reported photometric RMS from 0.042 to 0.037 mag and dense-subset recovery from 87.6% to 88.3%; the confidence intervals overlap, so this is an incremental operating-point change.")
    plain_rewrite(pfind(doc, "The intervals reinforce the intended operating-point interpretation"), "The intervals reinforce the intended operating-point interpretation. In this crop, spatial-ePSF/joint fitting trades approximately 3.7 times the Photutils runtime for higher dense-field recovery and improved flux measurement relative to the original AstroCFR ePSF branch; AstroCFR-RF remains the fast catalogue mode. The branches are therefore reported as complementary operating points rather than collapsed into a single score.")
    plain_rewrite(pfind(doc, "We therefore ran an independent-PSF stress test"), "We therefore ran a pilot independent-PSF stress test on the public PHAT M31 B21-F15 F475W DRZ image. The official Anderson ACS/WFC F475W standard PSF, rather than AstroCFR's image-derived ePSF, was evaluated at each source location through the WCS solutions of three registered 370-s FLC exposures, broadened with a fixed 0.55-pixel Gaussian output kernel to match the measured 2.2-pixel DRZ core, and normalized on the DRZ grid. On common denominators of 41 low-density and 42 high-density F475W=26.5 injections, AstroCFR ePSF recovered 39 and 35 stars versus 28 and 24 for Photutils. The sample is too small to serve as a primary validation set; it is retained only as a pilot check that the proposal gain is not solely caused by reusing the recovery ePSF as the injector. Full paired statistics, renderer diagnostics, and limitations are reported in Supplementary Section S10.")
    pareto = pfind(doc, "The Pareto frontier exposes three practically useful operating points")
    plain_rewrite(pareto, "The 15-field evidence supports three practically useful operating points. DAOStarFinder and SEP are low-latency screening baselines; Photutils is a strong conditional-precision measurement branch; AstroCFR ePSF-deblend and especially spatial-ePSF joint fitting are recovery-oriented modes for high-density blends. Across the 11 ACSGGCT fields, spatial-ePSF joint fitting improves high-density recovery in every field and has lower magnitude RMS in 8/11 fields and lower position RMS in 7/11, but at higher computational cost. Photutils remains the precision-oriented choice in the remaining field/subset cases. The defensible claim is a conditional high-density advantage under the disclosed single-image protocol, not universal leadership over Photutils or multi-exposure DOLPHOT/ALLFRAME.")
    yan_conclusion = insert_after(pareto, style="Normal")
    plain_rewrite(yan_conclusion, "The CSST geometric-distortion solution is relevant to future registration-aware deployment, but it is not an additional AstroCFR baseline.")
    archive_note = pfind(doc, "The exact environment, package versions")
    archive_note._element.getparent().remove(archive_note._element)
    plain_rewrite(pfind(doc, "Funding:"), "Funding: None.")
    code_avail = pfind(doc, "Code availability:")
    distortion_context = insert_after(code_avail, style="Normal")
    linked_rewrite(distortion_context, "The companion geometric-distortion study by Yan et al., 2026b provides complementary context for CSST registration; it is not an additional AstroCFR baseline.", LINKS)
    plain_rewrite(code_avail, "Code availability: The manuscript-matched source is prepared for https://github.com/zxl1999/AstroCFR. The release contains reusable modules in src/wpdc, controlled HST and CSST experiment scripts, machine-readable summaries, environment locks, data provenance, manuscript builders, and the partition-isolated ePSF sensitivity audit. The prepared package version is 1.6.2; the tagged archive is https://github.com/zxl1999/AstroCFR/archive/refs/tags/v1.6.2.zip. No archival DOI is claimed before a public release exists.")
    # Geometric-distortion context is placed in the Introduction literature map,
    # not in Statements and Declarations.
    ref_note = pfind(doc, "For ACS/WFC calibration context, the manuscript also follows")
    plain_rewrite(ref_note, "For ACS/WFC calibration context, the manuscript also follows the effective-PSF and detector-photometry foundations of Anderson & King (2000) and Sirianni et al. (2005).")
    data_note = pfind(doc, "Data availability:")
    plain_rewrite(data_note, "Data availability: The primary observational benchmark uses the public MAST ACSGGCT v2 HST/ACS F606W images and quality-selected catalogues for the 11 globular-cluster fields. PHAT M31 F475W FLC exposures and their public catalogues are included only for the separately labelled DOLPHOT FakeStars reference audit and the independent-PSF pilot; they are not inputs to the 11-field headline matrix. CSST challenge data are not redistributed. Download scripts, field identifiers, hashes where available, and machine-readable summaries are provided in the reproducibility package; large FITS products remain external.")
    plain_rewrite(pfind(doc, "Validation-selected per-chip calibration was necessary"), "Validation-selected per-chip calibration was necessary for heterogeneous simulated chips. In the registered full-frame CSST audit, calibrated SExtractor and AstroCFR ePSF + residual deblend (registered full-frame branch) provide supplied-catalogue integration-audit values. The controlled-crop audit additionally runs all five image-only branches, but its purpose is to expose simulation-to-reality differences and verify implementation feasibility; it is not used to form an all-CSST ranking.")
    plain_rewrite(pfind(doc, "On the CSST-like simulations, calibrated SExtractor"), "On the CSST-like simulations, calibrated SExtractor and AstroCFR ePSF + residual deblend (registered full-frame branch) provide the registered supplied-catalogue integration audit. The new controlled-crop audit additionally provides a method-complete check: spatial-ePSF has 70.5% median recovery and 6.41 mas position RMS versus 63.3% and 9.35 mas for Photutils, while Photutils has lower magnitude RMS (0.0283 versus 0.0463 mag). These crop values are supplementary feasibility evidence, not a replacement for the full-frame CSST audit or a broad CSST ranking. For the 11-field HST headline, the corresponding 86.3% and 45.9% values are field medians across all fields, whereas the 87.6% value cited elsewhere is the single-field NGC 6752 dense-subset result (n=402); they are not competing estimates of the same denominator.")
    cap = pfind(doc, "Table 8. Full cross-field comparison")
    table_caption_rewrite(cap, "Compact field-by-field high-density comparison of AstroCFR spatial-ePSF joint fitting (S) and Photutils (P). Recovery uses V<=20 references with at least three neighbours within 10 pixels; RMS values are conditional on held-out matched references. Supplementary Table S26 contains the complete all-branch table.")
    rebuild_compact_main_table(doc, cross_field_table, base)
    refresh_embedded_figure(
        doc,
        "/word/media/image7.png",
        ROOT / "results" / "astrocfr_manuscript_figures" / "fig21_astrocfr_density_recovery.png",
    )
    doc.save(DST)
    subprocess.run([
        "powershell", "-ExecutionPolicy", "Bypass", "-File",
        str(ROOT / "tools" / "insert_native_word_equations.ps1"),
        "-DocumentPath", str(DST),
    ], check=True)
    print(DST)


if __name__ == "__main__":
    main()
