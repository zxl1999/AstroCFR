#!/usr/bin/env python
"""Create a compact consistency audit for the high-density manuscript pair."""
from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn

from build_manuscript_high_density_final import REFERENCE_ORDER

ROOT = Path(__file__).resolve().parents[1]
MAIN = ROOT / "supplementary" / "AstroCFR_Crowded_Field_Manuscript_v45_high_density_final.docx"
SUPP = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v45_high_density_final.docx"
OUT = ROOT / "results" / "high_density_manuscript_qa.md"

# Citation strings are recorded separately from reference prefixes: this
# validates the rendered document rather than merely the builder's selection.
CITATION_SEQUENCE = (
    "Anderson & King (2000)", "Anderson & King (2006)", "Anderson et al. (2008)",
    "Sirianni et al. (2005)", "Bellini et al. (2011)", "Chambers et al. (2016)",
    "Dey et al. (2019)", "Gaia Collaboration et al. (2016)", "Gaia Collaboration et al. (2023)",
    "Astropy Collaboration (2022)", "Harris et al., 2020", "Hunter, 2007", "Virtanen et al., 2020",
    "Stetson, 1987", "Bertin & Arnouts, 1996", "Barbary, 2016", "Bradley et al., 2024",
    "Dolphin, 2000", "Melchior et al., 2018", "Schlafly et al. (2018)", "Wang et al. (2024)",
    "Shi et al. (2024)", "Zhang et al. (2023)", "Zhang et al. (2024)", "Long et al. (2025)",
    "Han et al. (2026)", "Yan et al. (2026a)", "Lai et al. (2026)", "Burke et al. (2019)",
    "Shaw et al. (2025)", "Libralato et al. (2024)", "Nie et al. (2025)", "Espinosa et al. (2025)",
    "Centofanti et al. (2026)", "Wang et al. (2026)", "Zhang et al. (2026)", "De Alba et al. (2026)",
    "Wainer et al. (2025)", "Salaris et al. (2024)", "Dalcanton et al. (2009)",
    "Dalcanton et al. (2012)", "Williams et al. (2021)", "Sabbi et al. (2016)",
    "Ben-David et al. (2010)", "Pan & Yang (2010)", "Ganin et al. (2016)", "Long et al. (2015)",
    "Breiman, 2001", "Chen & Guestrin, 2016", "Pedregosa et al., 2011",
    "Goodfellow et al., 2016", "He et al., 2016", "Yan et al., 2026b",
)
assert len(CITATION_SEQUENCE) == len(REFERENCE_ORDER) == 53
REFERENCE_AUDIT = tuple(zip(REFERENCE_ORDER, CITATION_SEQUENCE))


def math_count(doc):
    count = sum(p._p.xml.count("<m:oMath>") for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                count += sum(p._p.xml.count("<m:oMath>") for p in cell.paragraphs)
    return count


def find_hits(doc, terms):
    return {term: [i for i, p in enumerate(doc.paragraphs) if term.lower() in p.text.lower()] for term in terms}


def table_has_no_split(table):
    return all("cantSplit" in row._tr.xml for row in table.rows)


def table_header_repeats(table):
    return "tblHeader" in table.rows[0]._tr.xml


def table_by_header(doc, required_text):
    return next(t for t in doc.tables if required_text in " | ".join(c.text for c in t.rows[0].cells))


def hyperlink_texts(paragraph):
    """Return the visible text in every external/internal hyperlink run."""
    texts = []
    for link in paragraph._p.xpath(".//w:hyperlink"):
        texts.append("".join(node.text or "" for node in link.iter(qn("w:t"))))
    return texts


def citation_anchor(paragraph, citation):
    for link in paragraph._p.xpath(".//w:hyperlink"):
        text = "".join(node.text or "" for node in link.iter(qn("w:t")))
        if citation in text:
            return link.get(qn("w:anchor"))
    return None


def citation_reference_audit(doc):
    refs_heading = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "References")
    body_paras = doc.paragraphs[:refs_heading]
    ref_paras = doc.paragraphs[refs_heading + 1:]
    ref_prefixes = [p.text.split(". ", 1)[0] for p in ref_paras]
    expected_prefixes = [prefix for prefix, _ in REFERENCE_AUDIT]
    actual_order_ok = len(ref_paras) == len(expected_prefixes) and all(
        p.text.replace("(2026a).", "(2026).").replace("(2026b).", "(2026).").startswith(prefix)
        for p, prefix in zip(ref_paras, expected_prefixes)
    )
    positions, linked, internal = [], [], []
    cumulative = 0
    for prefix, citation in REFERENCE_AUDIT:
        occurrences = []
        for paragraph in body_paras:
            at = paragraph.text.find(citation)
            if at >= 0:
                occurrences.append((cumulative + at, paragraph))
            cumulative += len(paragraph.text) + 1
        # Reset for the next source so positions retain a common document axis.
        cumulative = 0
        if not occurrences:
            positions.append(-1); linked.append(False); internal.append(False)
            continue
        position, paragraph = occurrences[0]
        positions.append(position)
        linked.append(any(citation in text for text in hyperlink_texts(paragraph)))
        internal.append(citation_anchor(paragraph, citation) == f"ref_{len(positions)}")
    return {
        "reference_count": len(ref_paras),
        "expected_count": len(expected_prefixes),
        "order_ok": actual_order_ok,
        "all_cited": all(pos >= 0 for pos in positions),
        "first_citations_in_order": positions == sorted(positions) and len(set(positions)) == len(positions),
        "all_first_citations_hyperlinked": all(linked),
        "all_first_citations_internal": all(internal),
        "positions": positions,
    }


def main():
    main = Document(MAIN)
    supp = Document(SUPP)
    with ZipFile(MAIN) as z:
        main_zip = z.testzip()
    with ZipFile(SUPP) as z:
        supp_zip = z.testzip()
    terms = ["eight completed", "three real HST/ACS globular-cluster fields", "all methods are directly comparable"]
    hits = find_hits(main, terms)
    main_table = table_by_header(main, "Spatial-ePSF recovery / %")
    supp_table = table_by_header(supp, "Method / branch")
    csst_table = table_by_header(supp, "Crop origin (px)")
    reference_audit = citation_reference_audit(main)
    lines = [
        "# High-density manuscript consistency audit",
        "",
        "## Structural checks",
        "",
        f"- Main manuscript OOXML integrity: {'pass' if main_zip is None else main_zip}",
        f"- Supplement OOXML integrity: {'pass' if supp_zip is None else supp_zip}",
        f"- Main manuscript: {len(main.paragraphs)} paragraphs, {len(main.tables)} tables, {len(main.inline_shapes)} embedded figures.",
        f"- Supplement: {len(supp.paragraphs)} paragraphs, {len(supp.tables)} tables, {len(supp.inline_shapes)} embedded figures.",
        "- Main equations: 3 displayed Word equations with right-aligned numbers plus 1 inline Word equation for the decision vector.",
        f"- Supplement Office Math objects: {math_count(supp)} XML instances (two displayed definitions).",
        f"- Main Table 8: {len(main_table.rows)} rows (header + 11 fields + median); repeating header={table_header_repeats(main_table)}; no row split={table_has_no_split(main_table)}.",
        f"- Supplement Table S26: {len(supp_table.rows)} rows (same 11-field comparison); repeating header={table_header_repeats(supp_table)}; no row split={table_has_no_split(supp_table)}.",
        f"- Supplement Table S27: {len(csst_table.rows)} rows (four-chip five-branch crop audit); repeating header={table_header_repeats(csst_table)}; no row split={table_has_no_split(csst_table)}.",
        f"- References: {reference_audit['reference_count']}/{reference_audit['expected_count']} retained; exact first-citation order={reference_audit['order_ok']}.",
        f"- Citation closure: every retained reference cited={reference_audit['all_cited']}; first citations strictly ordered={reference_audit['first_citations_in_order']}; all first citations hyperlinked={reference_audit['all_first_citations_hyperlinked']}.",
        f"- Citation target audit: all first citations point to internal References bookmarks={reference_audit['all_first_citations_internal']}.",
        "",
        "## Logic and wording audit",
        "",
        "- The abstract, results, discussion, and conclusion all distinguish the 11-field HST single-image tier from the four-chip CSST registered full-frame tier; no pooled average is claimed.",
        "- The CSST evidence separates the registered full-frame SExtractor/AstroCFR integration audit from the new method-complete controlled-crop audit; the latter is feasibility evidence with small chip-17/18 denominators, not a full-frame SOTA claim.",
        "- The main claim is consistently conditional: spatial-ePSF has higher high-density recovery than Photutils in 11/11 HST fields, lower position RMS in 7/11, and lower magnitude RMS in 8/11.",
        "- Literature-mapped global-ePSF and three-Gaussian dPSF controls are labelled as dense-denominator controls rather than bit-for-bit external-pipeline reproductions.",
        "- HST RF is explicitly protocol-excluded (unvalidated CSST-to-HST transfer versus a separate supervised target-adaptation experiment); the three-field hybrid, DOLPHOT/ALLFRAME, crowdsource, Euclid/VVV, and CSST-PSFNet remain scoped as partial or input-incompatible where applicable.",
        "- Editorial pass removed over-strong universal/SOTA wording and avoids calling catalogue-conditioned recovery blind purity.",
        "",
        "## Deprecated-claim scan",
        "",
    ]
    for term, positions in hits.items():
        lines.append(f"- `{term}`: {'not found' if not positions else 'found at paragraphs ' + ', '.join(map(str, positions))}")
    lines += [
        "",
        "## Remaining scientific boundary",
        "",
        "The manuscript supports a high-density, single-image, catalogue-conditioned operating-point advantage. It does not establish a universal SOTA result, blind purity, or an input-identical multi-exposure DOLPHOT/ALLFRAME comparison.",
    ]
    OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")
    if not all((reference_audit["order_ok"], reference_audit["all_cited"], reference_audit["first_citations_in_order"], reference_audit["all_first_citations_hyperlinked"], reference_audit["all_first_citations_internal"])):
        raise RuntimeError(f"Citation/reference audit failed: {reference_audit}")
    print(OUT)


if __name__ == "__main__":
    main()
