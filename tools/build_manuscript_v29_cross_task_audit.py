#!/usr/bin/env python
"""Document why cross-task Transformer repositories are not direct baselines."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v28_transformer_ablation.docx"
DEST = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v29_cross_task_audit.docx"


def font(run):
    run.font.name = "Times New Roman"; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5); run.font.color.rgb = RGBColor(0, 0, 0)


def main():
    doc = Document(SOURCE)
    heading = next(p for p in doc.paragraphs if p.text.strip() == "6 Discussion")
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Inches(0.35); p.paragraph_format.line_spacing = 1.08
    font(p.add_run("We also audited two recent Transformer repositories to avoid an inflated but unfair cross-task comparison. CSST-PSFNet accepts labelled 32 × 32 star stamps and reconstructs 64 × 64 PSFs, whereas the CSST-like FITS used here contains only the science IMAGE extension and no PSF-label extensions or released checkpoint. SwinBayesNet addresses five-band SDSS hot-subdwarf classification and its repository does not declare a usable open-source licence. We therefore report neither as a numerical WPDC baseline: the former is recorded as an interface audit, and the latter is represented by our independently implemented lightweight patch-Transformer ablation. This preserves task, data, and licensing comparability."))
    cursor = heading._p
    node = p._p; node.getparent().remove(node); cursor.addnext(node)
    doc.save(DEST); print(DEST)


if __name__ == "__main__": main()
