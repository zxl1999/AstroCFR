#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Apply the final manuscript typography and three-line table style."""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[1]
SOURCE=ROOT/'WPDC_Multimedia_Systems_SCI_manuscript_v16_github_ready.docx'
DEST=ROOT/'WPDC_Multimedia_Systems_SCI_manuscript_v17_sci_format.docx'

def set_border(parent, edge, val, sz='8', color='000000'):
    tag='w:'+edge
    el=parent.find(qn(tag))
    if el is None:
        el=OxmlElement(tag);parent.append(el)
    el.set(qn('w:val'),val)
    if val!='nil':
        el.set(qn('w:sz'),sz);el.set(qn('w:space'),'0');el.set(qn('w:color'),color)

def clear_borders(cell):
    tcPr=cell._tc.get_or_add_tcPr();b=tcPr.find(qn('w:tcBorders'))
    if b is not None:tcPr.remove(b)

def style_three_line(table):
    table.alignment=WD_TABLE_ALIGNMENT.CENTER;table.autofit=True
    tblPr=table._tbl.tblPr;old=tblPr.find(qn('w:tblBorders'))
    if old is not None:tblPr.remove(old)
    borders=OxmlElement('w:tblBorders')
    for edge,val,sz in [('top','single','8'),('bottom','single','8'),('insideH','nil','0'),('insideV','nil','0')]:set_border(borders,edge,val,sz)
    tblPr.append(borders)
    for ri,row in enumerate(table.rows):
        for cell in row.cells:
            clear_borders(cell)
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent=None;p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name='Times New Roman';run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体');run.font.size=Pt(10.5);run.font.color.rgb=RGBColor(0,0,0)
                    if ri==0:run.bold=True
        if ri==0:
            for cell in row.cells:
                tcPr=cell._tc.get_or_add_tcPr();b=OxmlElement('w:tcBorders');set_border(b,'bottom','single','8');tcPr.append(b)

def main():
    doc=Document(SOURCE)
    # Normalize and move captions that were appended after their table.
    for p in doc.paragraphs:
        m=re.match(r'^(Table\s+\d+)\.?\s+(.*)$',p.text.strip(),re.I)
        if m:
            p.text=m.group(1)+'. '+m.group(2)
            p.style='Caption'
    for table in list(doc.tables):
        nxt=table._tbl.getnext()
        if nxt is not None and nxt.tag.endswith('}p'):
            from docx.text.paragraph import Paragraph
            cap=Paragraph(nxt,doc)
            if cap.text.strip().startswith('Table '):table._tbl.addprevious(nxt)
    for table in doc.tables:style_three_line(table)

    # Times New Roman 10.5 pt, black headings and body; use first-line
    # indentation for formal paragraphs rather than inserting literal spaces.
    heading_active=False;first_after_heading=True;references=False
    for p in doc.paragraphs:
        style=p.style.name if p.style else ''
        is_heading=style.startswith('Heading')
        if is_heading:
            heading_active=True;first_after_heading=True
            for run in p.runs:
                run.font.name='Times New Roman';run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体');run.font.size=Pt(10.5);run.font.bold=True;run.font.color.rgb=RGBColor(0,0,0)
            p.paragraph_format.first_line_indent=None;p.paragraph_format.space_before=Pt(8);p.paragraph_format.space_after=Pt(3)
            if p.text.strip()=='References':references=True
            continue
        if not p.text.strip():continue
        for run in p.runs:
            run.font.name='Times New Roman';run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体');run.font.size=Pt(10.5);run.font.color.rgb=RGBColor(0,0,0)
        p.paragraph_format.line_spacing=1.08
        if style=='Caption' or references or p.text.startswith(('Funding:','Competing interests:','Data availability:','Code availability:','Author contributions:','AI-assisted drafting:')):
            p.paragraph_format.first_line_indent=None;p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        elif first_after_heading:
            p.paragraph_format.first_line_indent=None;first_after_heading=False
        else:
            p.paragraph_format.first_line_indent=Inches(0.35)
    doc.save(DEST);print(DEST)
if __name__=='__main__':main()
