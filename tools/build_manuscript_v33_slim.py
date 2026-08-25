#!/usr/bin/env python
"""Build a focused v33 main manuscript and a standalone supplement.

The v32 evidence and numerical values are preserved. The main manuscript is
reorganized into five Results themes. The nine-classifier comparison and the
RF/CNN/Transformer architecture ablation are moved to Supplementary S1-S2.
"""
from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_manuscript_v30_submission_fixes import set_font


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "AstroCFR_Multimedia_Systems_SCI_manuscript_v32.docx"
DEST = ROOT / "supplementary" / "AstroCFR_Multimedia_Systems_SCI_manuscript_v33_slim.docx"
SUPPLEMENT = ROOT / "supplementary" / "AstroCFR_Supplementary_Materials_v33.docx"
SUPP_FIG_DIR = ROOT / "results" / "astrocfr_supplementary_figures"
FIG_S1 = SUPP_FIG_DIR / "figS1_nine_classifier_comparison.png"


ABSTRACT = (
    "We present AstroCFR (Astronomical Crowded-Field Recovery System), a modular pipeline for candidate "
    "generation, lightweight target adaptation, empirical-PSF modelling, and crowded-field recovery in "
    "large-scale astronomical survey images. On four CSST-like challenge chips containing approximately "
    "4,000 reference sources, AstroCFR achieves a mean recall of 94.6% with 100% simulated-catalogue precision "
    "under the disclosed matching and per-chip calibration protocol. Controlled evaluation on three public "
    "HST/ACS crowded fields reveals distinct deployment operating points. On the dense NGC 6752 V≤20 subset, "
    "the AstroCFR ePSF+deblend branch recovers 87.6% of reference sources (95% CI: 84.0–90.4%), compared with "
    "57.0% for DAOStarFinder/Photutils proposals and 28.6% for SEP, whereas Photutils provides lower positional "
    "and photometric RMS. Lightweight adaptation is sample-efficient: one spatially isolated 200 × 200 pixel "
    "calibration tile raises held-out recall on NGC 6752 from 0.057 to 0.648. On 800 identical artificial-star "
    "injections, a fixed density-adaptive router achieves 50.00% recovery at an estimated 20.44 s/MPix, defining "
    "an intermediate operating point between Photutils-only processing (39.25%, 8.64 s/MPix) and the always-on "
    "spatial-ePSF branch (60.38%, 32.23 s/MPix). AstroCFR therefore provides explicit recovery–precision–cost "
    "trade-offs rather than universal superiority over specialized photometric tools."
)

INTRO_POSITIONING = (
    "AstroCFR is therefore positioned not as a universal replacement for specialized astronomical extraction "
    "or photometric tools, but as a modular deployment framework for heterogeneous survey streams. Its branches "
    "expose explicit operating points for conservative catalogue construction, high-crowding recovery, and "
    "density-aware resource allocation. The central research question is consequently how candidate recovery, "
    "measurement precision, target adaptation, and computational cost can be balanced under a common and "
    "reproducible system protocol."
)

ORGANIZATION = (
    "The manuscript is organized as follows. Section 2 describes the data and problem setting, Section 3 the "
    "AstroCFR pipeline, and Section 4 the common experimental protocol. Section 5 reports five focused result "
    "themes: CSST-like validation, target adaptation, controlled HST/ACS recovery, deployment operating points, "
    "and failure analysis. Section 6 discusses implications and limitations, and Section 7 concludes. The full "
    "nine-classifier comparison and RF/CNN/Transformer architecture ablation are provided in Supplementary S1–S2."
)

CLASSIFIER_SUMMARY = (
    "Classifier selection was treated as a system-design decision rather than a separate leaderboard. Under the "
    "registered candidate split, per-chip normalization, and threshold-selection protocol, the target-adapted "
    "RandomForest achieved the strongest accuracy–cost operating point among the nine evaluated classifiers and "
    "was retained as the fast conservative catalogue branch. The complete nine-method results are reported in "
    "Supplementary Table S1 and Fig. S1."
)

ARCHITECTURE_SUMMARY = (
    "A controlled RF/CNN/lightweight-Transformer ablation likewise retained the RandomForest as the strongest "
    "candidate-screening operating point under the pre-registered threshold rule. The Transformer improved over "
    "the original CNN but did not improve upon the RF and required substantially longer training. Architecture "
    "details and complete metrics are provided in Supplementary Table S2."
)


def replace_paragraph(paragraph, text: str, *, size: float = 10.5, bold: bool = False) -> None:
    paragraph.clear()
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)


def replace_text_nodes(paragraph, old: str, new: str) -> None:
    for node in paragraph._p.iter(qn("w:t")):
        if node.text:
            node.text = node.text.replace(old, new)


def paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def find_paragraph(doc: Document, text: str):
    return next(p for p in doc.paragraphs if p.text.strip() == text)


def find_starts(doc: Document, prefix: str):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def remove_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def body_range(doc: Document, start_text: str, end_text: str):
    children = list(doc._element.body.iterchildren())
    start = next(i for i, child in enumerate(children) if paragraph_text(child).strip() == start_text)
    end = next(i for i, child in enumerate(children[start + 1:], start + 1)
               if paragraph_text(child).strip() == end_text)
    return children[start:end]


def remove_body_range(doc: Document, start_text: str, end_text: str) -> None:
    for element in body_range(doc, start_text, end_text):
        element.getparent().remove(element)


def move_body_range_before(doc: Document, start_text: str, end_text: str, target_text: str) -> None:
    block = body_range(doc, start_text, end_text)
    target = next(child for child in doc._element.body.iterchildren()
                  if paragraph_text(child).strip() == target_text)
    for element in block:
        element.getparent().remove(element)
    for element in block:
        target.addprevious(element)


def insert_before(paragraph, text: str):
    new = OxmlElement("w:p")
    paragraph._p.addprevious(new)
    p = paragraph._parent.add_paragraph()
    p._p.getparent().remove(p._p)
    new.addnext(p._p)
    # Move the new python-docx paragraph into the exact position and remove the
    # empty marker paragraph.
    new.getparent().remove(new)
    replace_paragraph(p, text)
    p.paragraph_format.first_line_indent = Inches(.35)
    return p


def rename_heading(doc: Document, old: str, new: str, style: str = "Heading 1") -> None:
    paragraph = find_paragraph(doc, old)
    paragraph.style = style
    replace_paragraph(paragraph, new, bold=True)


def remove_heading(doc: Document, text: str) -> None:
    remove_paragraph(find_paragraph(doc, text))


def renumber_main(doc: Document) -> None:
    table_map = {n: n for n in range(1, 4)}
    table_map.update({n: n - 1 for n in range(5, 24)})
    table_map[25] = 23
    figure_map = {n: n for n in range(1, 6)}
    figure_map.update({n: n - 1 for n in range(7, 25)})

    table_rx = re.compile(r"\bTable\s+(\d+)\b")
    figure_rx = re.compile(r"\bFig\.\s+(\d+)\b")
    figure_word_rx = re.compile(r"\bFigure\s+(\d+)\b")
    for node in doc._element.body.iter(qn("w:t")):
        if not node.text:
            continue
        node.text = table_rx.sub(lambda m: f"Table {table_map.get(int(m.group(1)), int(m.group(1)))}", node.text)
        node.text = figure_rx.sub(lambda m: f"Fig. {figure_map.get(int(m.group(1)), int(m.group(1)))}", node.text)
        node.text = figure_word_rx.sub(lambda m: f"Figure {figure_map.get(int(m.group(1)), int(m.group(1)))}", node.text)

    for bookmark in doc._element.body.iter(qn("w:bookmarkStart")):
        name = bookmark.get(qn("w:name"))
        if not name:
            continue
        match = re.fullmatch(r"table_(\d+)", name)
        if match and int(match.group(1)) in table_map:
            bookmark.set(qn("w:name"), f"table_{table_map[int(match.group(1))]}")
        match = re.fullmatch(r"figure_(\d+)", name)
        if match and int(match.group(1)) in figure_map:
            bookmark.set(qn("w:name"), f"figure_{figure_map[int(match.group(1))]}")

    for link in doc._element.body.iter(qn("w:hyperlink")):
        anchor = link.get(qn("w:anchor"))
        if not anchor:
            continue
        match = re.fullmatch(r"table_(\d+)", anchor)
        if match and int(match.group(1)) in table_map:
            link.set(qn("w:anchor"), f"table_{table_map[int(match.group(1))]}")
        match = re.fullmatch(r"figure_(\d+)", anchor)
        if match and int(match.group(1)) in figure_map:
            link.set(qn("w:anchor"), f"figure_{figure_map[int(match.group(1))]}")


def border(cell, edge: str, value: str = "single", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    element = borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), value)
    element.set(qn("w:sz"), size)
    element.set(qn("w:color"), "000000")


def format_three_line(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            border(cell, "top", "single" if row_index == 0 else "nil", "14" if row_index == 0 else "0")
            border(cell, "bottom", "single" if row_index in (0, len(table.rows) - 1) else "nil",
                   "8" if row_index == 0 else ("14" if row_index == len(table.rows) - 1 else "0"))
            border(cell, "left", "nil", "0"); border(cell, "right", "nil", "0")


def add_source_table(target: Document, source_table, font_size: float = 8.0):
    table = target.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    for r, source_row in enumerate(source_table.rows):
        for c, source_cell in enumerate(source_row.cells):
            cell = table.cell(r, c)
            cell.text = source_cell.text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_font(run, size=font_size, bold=(r == 0))
    format_three_line(table)
    return table


def add_normal(doc: Document, text: str, indent: bool = True):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.first_line_indent = Inches(.35) if indent else None
    paragraph.paragraph_format.line_spacing = 1.08
    set_font(paragraph.add_run(text), size=10.5)
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.paragraph_format.first_line_indent = None
    set_font(paragraph.add_run(text), size=9.0)
    return paragraph


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(.8); section.bottom_margin = Inches(.8)
    section.left_margin = Inches(.85); section.right_margin = Inches(.85)
    for name, size, bold in (("Normal", 10.5, False), ("Heading 1", 12, True),
                             ("Heading 2", 10.5, True), ("Caption", 9, False)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"; style.font.size = Pt(size); style.font.bold = bold
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def build_supplement(source: Document) -> None:
    SUPP_FIG_DIR.mkdir(parents=True, exist_ok=True)
    with ZipFile(SOURCE) as archive:
        FIG_S1.write_bytes(archive.read("word/media/image6.png"))

    supplement = Document()
    configure_styles(supplement)
    title = supplement.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("Supplementary Materials for AstroCFR"), size=15, bold=True)
    subtitle = supplement.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run("Classifier selection and architecture ablations"), size=11, bold=True)
    add_normal(supplement,
               "This supplement contains the detailed classifier analyses moved from the focused main manuscript. "
               "Candidate generation, data partitions, association rules, and threshold protocols are unchanged.",
               indent=False)

    heading = supplement.add_paragraph("S1 Nine-method classifier comparison", style="Heading 1")
    for run in heading.runs: set_font(run, size=12, bold=True)
    for index in (73, 74, 75, 93):
        add_normal(supplement, source.paragraphs[index].text)
    add_caption(supplement, "Table S1. Nine-method classifier comparison under per-chip threshold calibration.")
    add_source_table(supplement, source.tables[3], font_size=7.7)
    add_caption(supplement, "Fig. S1. Recall and precision for nine classification methods.")
    picture = supplement.add_picture(str(FIG_S1), width=Inches(6.35))
    supplement.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading = supplement.add_paragraph("S2 Controlled classifier-architecture ablation", style="Heading 1")
    for run in heading.runs: set_font(run, size=12, bold=True)
    add_normal(supplement, source.paragraphs[265].text)
    add_normal(supplement, source.paragraphs[266].text.replace("Table 24", "Table S2"))
    add_caption(supplement,
                "Table S2. Controlled simulation-domain candidate-classifier ablation. Values use the untouched "
                "20% test partition; thresholds are selected on the validation partition to retain at least 90% "
                "validation recall. GPU figures apply to the CNN and Transformer classifier stages only.")
    add_source_table(supplement, source.tables[23], font_size=8.0)
    add_normal(supplement,
               "CSST-PSFNet and SwinBayesNet are not added to Table S2 because their tasks, required labels, "
               "available checkpoints, and licensing conditions do not define a common candidate-classification "
               "experiment. The interface audit is archived separately in the repository.")
    supplement.save(SUPPLEMENT)


def build_main(source: Document) -> None:
    doc = source
    abstract_index = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Abstract")
    replace_paragraph(doc.paragraphs[abstract_index + 1], ABSTRACT)
    replace_paragraph(find_starts(doc, "We present AstroCFR as a deployable system"), INTRO_POSITIONING)
    replace_paragraph(find_starts(doc, "The manuscript is organized as follows."), ORGANIZATION)
    replace_text_nodes(find_starts(doc, "Code availability:"), "tag v1.1.1", "tag v1.2.0")

    method = find_starts(doc, "For isolated classifier benchmarking")
    replace_paragraph(method,
        "The classifier comparison uses identical exported candidate pools, per-chip normalization, stratified "
        "train/validation/test separation, and independently calibrated validation thresholds. The main text "
        "retains the resulting RandomForest selection; complete nine-method and neural-architecture protocols are "
        "reported in Supplementary S1–S2.")
    remove_paragraph(find_starts(doc, "The same exported candidate pools support"))
    remove_paragraph(find_starts(doc, "Feature normalization is performed per chip"))

    move_body_range_before(doc, "5.22 Density-adaptive deployment routing", "6 Discussion",
                           "5.20 Failure analysis and stratified recovery")
    remove_body_range(doc, "5.2 Nine-method classifier comparison", "5.3 Image-feature fusion")
    remove_body_range(doc, "5.21 Classifier architecture ablation", "6 Discussion")
    # The cross-task audit is preserved in the supplement and repository, not
    # repeated at the start of the focused Discussion.
    remove_paragraph(find_starts(doc, "We also audited two recent Transformer repositories"))

    target = find_paragraph(doc, "5.3 Image-feature fusion")
    insert_before(target, CLASSIFIER_SUMMARY)

    rename_heading(doc, "5.1 Per-chip threshold calibration", "5.1 CSST-like simulation validation")
    for old in ("5.3 Image-feature fusion", "5.4 Full-pipeline classifier integration",
                "5.5 Comparison with SExtractor", "5.6 Final chip-level catalog quality",
                "5.7 Candidate funnel and correction-chain behavior"):
        remove_heading(doc, old)

    rename_heading(doc, "5.8 Simulated-to-real zero-shot generalization",
                   "5.2 Simulation-to-real target adaptation")
    remove_heading(doc, "5.9 Few-shot target-domain adaptation")

    rename_heading(doc, "5.10 HST/ACS real-data benchmark and contextual comparison",
                   "5.3 HST/ACS crowded-field recovery")
    for old in ("5.11 Empirical-PSF fitting, residual deblending, and artificial-star recovery",
                "5.12 Controlled baseline, efficiency, and injection-recovery evaluation",
                "5.13 Spatially varying ePSF and joint-fitting operating point",
                "5.14 Measurement uncertainty and runtime repeatability",
                "5.15 Independent third-cluster validation: NGC 1851",
                "5.16 Expanded artificial-star validation",
                "5.17 Target-adaptation budget learning curve"):
        remove_heading(doc, old)

    rename_heading(doc, "5.18 Accuracy–cost operating points",
                   "5.4 Deployment operating points and resource cost")
    remove_heading(doc, "5.19 CNN accelerator resource profile")
    routing = find_paragraph(doc, "5.22 Density-adaptive deployment routing")
    insert_before(routing, ARCHITECTURE_SUMMARY)
    rename_heading(doc, "5.22 Density-adaptive deployment routing",
                   "5.4.1 Density-adaptive routing", style="Heading 2")
    rename_heading(doc, "5.20 Failure analysis and stratified recovery",
                   "5.5 Failure modes and stratified recovery")

    renumber_main(doc)
    doc.save(DEST)


def main() -> None:
    source_for_supplement = Document(SOURCE)
    build_supplement(source_for_supplement)
    build_main(Document(SOURCE))
    print(DEST)
    print(SUPPLEMENT)


if __name__ == "__main__":
    main()
