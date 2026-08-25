#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Add 2026 literature, rebuild linked references, and replace Fig. 1."""
from __future__ import annotations

import importlib.util
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v22_figure_layout.docx"
DEST = ROOT / "supplementary" / "WPDC_Multimedia_Systems_SCI_manuscript_v24_2026_refs_fig1_abstract.docx"
FIG1 = ROOT / "results" / "figures" / "fig1_wpdc_architecture_redrawn.png"
NEW_REFS = {
    "Han, Z.": "Han, Z., Zhang, T., Liu, C., & Ling, C. (2026). A multi-modal fusion network for star-galaxy classification from CSST simulated datasets. Astronomy and Computing, 56, 101112. https://doi.org/10.1016/j.ascom.2026.101112",
    "De Alba, K.": "De Alba, K., De Alba, E., Cabello, A., Gazak, J. Z., & Fletcher, J. (2026). Aperture-X: Physics-informed aperture feature learning for robust photometry. Astronomy and Computing, 57, 101165. https://doi.org/10.1016/j.ascom.2026.101165",
    "Wang, P.": "Wang, P., Wei, P., Liu, C., Wang, R., Wang, F., & Zhang, X. (2026). CSST-PSFNet: A point-spread function reconstruction model for the CSST based on deep learning. The Astrophysical Journal Supplement Series, 283(2), 77. https://doi.org/10.3847/1538-4365/ae5053",
    "Zhang, S., Wang, L.": "Zhang, S., Wang, L., Diao, Y., Yan, Z., Peng, X., Liu, Y., Shan, H., Wang, G., Liu, F., Wei, C., Nie, L., Chen, X., Ding, H., & Zheng, Z. (2026). Aperture photometric accuracy of point-spread-function-deconvolved astronomical images. The Astronomical Journal, 172(1), 67. https://doi.org/10.3847/1538-3881/ae7718",
}


def load_v20_module():
    path = ROOT / "tools" / "build_manuscript_v20_introduction_hyperlinks.py"
    spec = importlib.util.spec_from_file_location("wpdc_v20_links", path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def bookmark(paragraph, name, identifier):
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(identifier)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(identifier))
    paragraph._p.insert(0, start); paragraph._p.append(end)


def replace_media(source, destination, replacement):
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=destination.parent) as temporary:
        temp = Path(temporary.name)
    try:
        with ZipFile(source) as inp, ZipFile(temp, "w", ZIP_DEFLATED) as out:
            for info in inp.infolist():
                data = replacement.read_bytes() if info.filename == "word/media/image1.png" else inp.read(info.filename)
                if info.filename == "word/document.xml":
                    # Match the embedded figure's width and update its height to
                    # the two-row asset aspect ratio (avoids Word distortion).
                    data = data.replace(b'cx="6035040" cy="3095422"', b'cx="6035040" cy="3159000"')
                out.writestr(info, data)
        temp.replace(destination)
    finally:
        if temp.exists(): temp.unlink()


def insert_after(items, prefix, value):
    index = next(i for i, text in enumerate(items) if text.startswith(prefix))
    items.insert(index + 1, value)


def main():
    if not FIG1.exists(): raise FileNotFoundError(FIG1)
    base = load_v20_module()
    intro = list(base.INTRO)
    intro[2] = (
        "Our positioning follows established measurement and deployment practice: empirical ACS ePSFs and crowded-field catalogues (Anderson & King, 2006; Anderson et al., 2008; Bellini et al., 2011), survey-scale reference systems (Chambers et al., 2016; Dey et al., 2019; Gaia Collaboration et al., 2016, 2023), and reproducible scientific-computing components (Astropy Collaboration, 2022; Harris et al., 2020; Hunter, 2007; Virtanen et al., 2020; Barbary, 2016; Bradley et al., 2024; Pedregosa et al., 2011). Deep-learning foundations and residual networks provide the image-model context (Goodfellow et al., 2016; He et al., 2016). Recent context includes CSST photometric preparation (Shi et al., 2024), deep-feature point-source detection (Long et al., 2025), CSST multi-modal star-galaxy classification (Han et al., 2026), Rubin crowded-field photometry (Wainer et al., 2025), PSF-fitting uncertainty analysis (Espinosa et al., 2025), physics-informed aperture feature learning (De Alba et al., 2026), image-based PSF recovery (Centofanti et al., 2026; Wang et al., 2026), aperture photometry after PSF deconvolution (Zhang et al., 2026), and recent HST/Euclid cluster studies (Libralato et al., 2024; Salaris et al., 2024). Transfer-learning claims are framed using domain-shift theory and adaptation methods (Ben-David et al., 2010; Pan & Yang, 2010; Ganin et al., 2016; Long et al., 2015), while the classifier baselines are standard Random Forest and gradient-boosting references (Breiman, 2001; Chen & Guestrin, 2016)."
    )
    cite = {
        "Anderson & King, 2006":"Anderson, J., & King", "Anderson et al., 2008":"Anderson, J., et al.", "Bellini et al., 2011":"Bellini,", "Chambers et al., 2016":"Chambers,", "Dey et al., 2019":"Dey,", "Gaia Collaboration et al., 2016, 2023":"SPECIAL_GAIA", "Astropy Collaboration, 2022":"Astropy Collaboration,", "Harris et al., 2020":"Harris,", "Hunter, 2007":"Hunter,", "Virtanen et al., 2020":"Virtanen,", "Barbary, 2016":"Barbary,", "Bradley et al., 2024":"Bradley,", "Pedregosa et al., 2011":"Pedregosa,", "Goodfellow et al., 2016":"Goodfellow,", "He et al., 2016":"He,", "Shi et al., 2024":"Shi,", "Long et al., 2025":"Long, M., Xin", "Han et al., 2026":"Han, Z.", "Wainer et al., 2025":"Wainer,", "Espinosa et al., 2025":"Espinosa,", "De Alba et al., 2026":"De Alba, K.", "Centofanti et al., 2026":"Centofanti,", "Wang et al., 2026":"Wang, P.", "Zhang et al., 2026":"Zhang, S., Wang, L.", "Libralato et al., 2024":"Libralato,", "Salaris et al., 2024":"Salaris,", "Ben-David et al., 2010":"Ben-David,", "Pan & Yang, 2010":"Pan,", "Ganin et al., 2016":"Ganin,", "Long et al., 2015":"Long, M., Cao", "Breiman, 2001":"Breiman,", "Chen & Guestrin, 2016":"Chen,", "Stetson, 1987":"Stetson,", "Bertin & Arnouts, 1996":"Bertin,", "Dolphin, 2000":"Dolphin,", "Schlafly et al., 2018":"Schlafly,", "Melchior et al., 2018":"Melchior,", "Wang et al., 2024":"Wang, Y.", "Zhang et al., 2024a":"Zhang, Y.,", "Zhang et al., 2024b":"Zhang, S., Fang"
    }
    doc = Document(SOURCE)
    abstract = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Abstract")
    abstract_text = ("Dense stellar-field imaging is a demanding multimedia-processing problem that requires transforming high-resolution visual data into calibrated catalogues despite blending, background variation, and bright-star artifacts. We present WPDC as a deployable system that combines adaptive background modelling, candidate generation, target-domain RandomForest adaptation, empirical-PSF measurement, and residual deblending, and evaluate it against DAOStarFinder, SEP/SExtractor-style extraction, and Photutils PSFPhotometry on identical HST/ACS crops, spatial test splits, a 2-pixel association radius, and fixed artificial-star scenes. We formulate deployment in two stages: simulation development on CSST-like images, followed by lightweight target adaptation using an image-only PSF estimate and a small, spatially disjoint labelled calibration region. On three HST/ACS fields, one 200 x 200 pixel calibration tile raises held-out recall from 0.183 to 0.828 +/- 0.036 on NGC 6397 and from 0.057 to 0.648 +/- 0.021 on NGC 6752; the difficult NGC 1851 field requires six tiles to reach 0.166 +/- 0.002, near its single-image proposal ceiling of 0.173. In the controlled dense NGC 6752 subset, WPDC ePSF+deblend recovers 87.6% (95% CI: 84.0-90.4%) versus 57.0% for DAO, 28.6% for SEP, and 52.7% for WPDC-RF, at 28.0 s/MPix compared with 0.12 s/MPix for DAO; Photutils yields lower position and magnitude RMS. WPDC therefore exposes a fast, high-match-rate RF catalogue branch and a slower recovery-oriented ePSF+deblend branch. The evidence supports a bounded crowded-field recovery advantage and an auditable low-cost adaptation strategy, not universal leadership in astrometry, photometry, or throughput.")
    p_abs = doc.paragraphs[abstract + 1]; p_abs.clear(); p_abs.paragraph_format.first_line_indent = None; p_abs.paragraph_format.line_spacing = 1.08
    run = p_abs.add_run(abstract_text); base.font_run(run)
    ref_index = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "References")
    existing = [p.text.strip() for p in doc.paragraphs[ref_index + 1:] if p.text.strip()]
    insert_after(existing, "Long, M., Xin", NEW_REFS["Han, Z."])
    insert_after(existing, "Espinosa,", NEW_REFS["De Alba, K."])
    insert_after(existing, "Centofanti,", NEW_REFS["Wang, P."])
    insert_after(existing, "Wang, P.", NEW_REFS["Zhang, S., Wang, L."])
    for p in list(doc.paragraphs[ref_index + 1:]): p._element.getparent().remove(p._element)
    heading = doc.paragraphs[ref_index]; cursor = heading
    anchors = {}
    for number, text in enumerate(existing, 1):
        p = doc.add_paragraph(text, style="Normal")
        p.paragraph_format.left_indent = Inches(.20); p.paragraph_format.first_line_indent = Inches(-.20); p.paragraph_format.line_spacing = 1.0
        for run in p.runs: base.font_run(run)
        bookmark(p, f"ref_{number}", 1000 + number)
        cursor._p.addnext(p._p); cursor = p
        for prefix in cite.values():
            if prefix != "SPECIAL_GAIA" and text.startswith(prefix): anchors[prefix] = f"ref_{number}"
        if text.startswith("Gaia Collaboration, et al. (2016)"): anchors["GAIA2016"] = f"ref_{number}"
        if text.startswith("Gaia Collaboration, et al. (2023)"): anchors["GAIA2023"] = f"ref_{number}"
    intro_index = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "1 Introduction")
    targets = [p for p in doc.paragraphs[intro_index + 1:] if p.text.strip()][:6]
    import re
    rx = re.compile("|".join(re.escape(x) for x in sorted(cite, key=len, reverse=True)))
    for i, (paragraph, text) in enumerate(zip(targets, intro)):
        paragraph.clear(); paragraph.style = "Normal"; paragraph.paragraph_format.first_line_indent = None if i == 0 else Inches(.35); paragraph.paragraph_format.line_spacing = 1.08
        at = 0
        for match in rx.finditer(text):
            if match.start() > at:
                run = paragraph.add_run(text[at:match.start()]); base.font_run(run)
            label = match.group(0)
            if cite[label] == "SPECIAL_GAIA":
                base.link(paragraph, "Gaia Collaboration et al., 2016", anchors["GAIA2016"])
                run = paragraph.add_run(", "); base.font_run(run)
                base.link(paragraph, "2023", anchors["GAIA2023"])
            else:
                base.link(paragraph, label, anchors[cite[label]])
            at = match.end()
        if at < len(text):
            run = paragraph.add_run(text[at:]); base.font_run(run)
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=DEST.parent) as temporary:
        intermediate = Path(temporary.name)
    doc.save(intermediate)
    replace_media(intermediate, DEST, FIG1)
    intermediate.unlink(missing_ok=True)
    print(DEST, "references", len(existing))


if __name__ == "__main__":
    main()
